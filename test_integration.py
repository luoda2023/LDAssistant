#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
工程助手 LDAssistant — 端到端集成测试
测试完整流程：打开文件 → 提取文本 → 识别规范 → 查询数据库
"""
import sys
import os
import time
import traceback
import re

PROJECT_DIR = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, PROJECT_DIR)

# 测试统计
PASS = 0
FAIL = 0
RESULTS = []

def log(name, ok, detail=""):
    global PASS, FAIL
    status = "✅" if ok else "❌"
    if ok:
        PASS += 1
    else:
        FAIL += 1
    RESULTS.append((name, ok, detail))
    print(f"  {status} {name} — {detail}")


# ═══════════════════════════════════════════════
# 1. 文本提取函数（不依赖 GUI）
# ═══════════════════════════════════════════════

def test_pdf_text_extraction():
    """测试 PDF 文本提取"""
    try:
        import fitz  # pymupdf
        path = os.path.join(PROJECT_DIR, "test_files", "test_document.pdf")
        doc = fitz.open(path)
        full_text = ""
        for page in doc:
            full_text += page.get_text()
        doc.close()

        # 验证提取到文本
        assert len(full_text) > 0, "提取的文本为空"

        # 验证包含规范编号
        codes = re.findall(r'[A-Z]{1,5}[0-9]*(?:/[A-Z]{1,10})?\s*\d+(?:\.\d+)?-\d{4}', full_text)
        assert len(codes) > 0, f"未提取到规范编号，文本: {full_text[:200]}"

        return True, f"提取 {len(full_text)} 字符, {len(codes)} 个规范编号: {codes[:3]}"
    except Exception as e:
        return False, str(e)


def test_docx_text_extraction():
    """测试 Word (.docx) 文本提取"""
    try:
        from docx import Document
        path = os.path.join(PROJECT_DIR, "test_files", "test_document.docx")
        doc = Document(path)
        full_text = '\n'.join([p.text for p in doc.paragraphs])

        assert len(full_text) > 0, "提取的文本为空"

        codes = re.findall(r'[A-Z]{1,5}[0-9]*(?:/[A-Z]{1,10})?\s*\d+(?:\.\d+)?-\d{4}', full_text)
        assert len(codes) > 0, f"未提取到规范编号"

        return True, f"提取 {len(full_text)} 字符, {len(codes)} 个规范编号: {codes[:3]}"
    except Exception as e:
        return False, str(e)


def test_txt_text_extraction():
    """测试 TXT 文本提取"""
    try:
        path = os.path.join(PROJECT_DIR, "test_files", "test_document.txt")
        with open(path, 'r', encoding='utf-8') as f:
            full_text = f.read()

        assert len(full_text) > 0, "提取的文本为空"

        codes = re.findall(r'[A-Z]{1,5}[0-9]*(?:/[A-Z]{1,10})?\s*\d+(?:\.\d+)?-\d{4}', full_text)
        assert len(codes) > 0, f"未提取到规范编号"

        return True, f"提取 {len(full_text)} 字符, {len(codes)} 个规范编号: {codes[:3]}"
    except Exception as e:
        return False, str(e)


def test_dxf_render():
    """测试 DXF 文件渲染为图像"""
    try:
        import ezdxf
        from ezdxf.addons.drawing import RenderContext, Frontend
        from ezdxf.addons.drawing.matplotlib import MatplotlibBackend
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt

        path = os.path.join(PROJECT_DIR, "test_files", "test_drawing.dxf")
        doc = ezdxf.readfile(path)
        msp = doc.modelspace()

        fig, ax = plt.subplots(figsize=(10, 8))
        ax.set_aspect('equal')

        ctx = RenderContext(doc)
        backend = MatplotlibBackend(ax)
        frontend = Frontend(ctx, backend)
        # ezdxf 1.4+ 使用 draw_layout 替代 draw
        if hasattr(frontend, 'draw_layout'):
            frontend.draw_layout(msp)
        else:
            frontend.draw(msp)

        import tempfile
        img_path = tempfile.mktemp(suffix='.png')
        fig.savefig(img_path, dpi=100, facecolor='white')
        plt.close(fig)

        assert os.path.exists(img_path), "渲染图像未生成"
        assert os.path.getsize(img_path) > 1000, "渲染图像太小"

        # 清理
        os.remove(img_path)

        return True, f"DXF 渲染成功, 实体数: {len(list(msp))}"
    except Exception as e:
        return False, str(e)


# ═══════════════════════════════════════════════
# 2. 规范编号提取
# ═══════════════════════════════════════════════

def test_code_extraction():
    """测试从文本中提取规范编号"""
    try:
        # 导入正则模式
        from standard_checker import CODE_PATTERN, normalize_for_matching

        # 测试文本
        test_text = """
        GB/T 50430-2017 工程建设施工企业质量管理规范
        GB 50010-2010(2015版) 混凝土结构设计规范
        CJJ 1-2008 城镇道路工程施工与质量验收规范
        JGJ 1-2014 装配式混凝土结构技术规程
        DG/TJ 08-2002 上海市工程建设规范
        """

        codes = CODE_PATTERN.findall(test_text)
        unique_codes = list(set(normalize_for_matching(c) for c in codes))

        # 验证提取到足够的规范编号
        assert len(codes) >= 5, f"提取到 {len(codes)} 个编号，期望 >= 5"
        # 验证规范化结果
        assert 'GB/T50430-2017' in unique_codes, f"规范化结果缺少 GB/T50430-2017"
        assert 'CJJ1-2008' in unique_codes, f"规范化结果缺少 CJJ1-2008"

        return True, f"提取 {len(codes)} 个编号，规范化后 {len(unique_codes)} 个唯一编号"
    except Exception as e:
        return False, str(e)


# ═══════════════════════════════════════════════
# 3. 数据库查询
# ═══════════════════════════════════════════════

def test_database_query():
    """测试数据库查询功能"""
    try:
        from standard_db import StandardChecker

        checker = StandardChecker()
        total = checker.total_count if hasattr(checker, 'total_count') else 0

        # 测试多个规范编号查询
        test_codes = [
            ("GB/T 50430-2017", True, "现行"),
            ("CJJ 1-2008", True, None),
            ("JGJ 1-2014", True, None),
        ]

        found_count = 0
        for code, expect_found, expect_status in test_codes:
            result = checker.check_code(code)
            if result.get('found', False):
                found_count += 1
                # 如果期望有状态，检查
                if expect_status:
                    status = result.get('status', '')
                    if expect_status not in status:
                        return False, f"{code} 状态异常: 期望'{expect_status}' 实际'{status}'"

        checker.close()

        if found_count < 2:
            return False, f"仅找到 {found_count}/{len(test_codes)} 个规范"

        return True, f"数据库加载 {total} 条, 查询 {found_count}/{len(test_codes)} 成功"
    except Exception as e:
        return False, str(e)


# ═══════════════════════════════════════════════
# 4. OCR 预处理函数
# ═══════════════════════════════════════════════

def test_ocr_preprocessing():
    """测试 OCR 文本预处理函数"""
    try:
        from standard_checker import fullwidth_to_halfwidth, normalize_for_matching

        # 全角 → 半角测试
        tests = [
            ("１２３４５", "12345"),
            ("ＧＢ／Ｔ", "GB/T"),
            ("（）", "()"),
        ]
        for inp, expected in tests:
            result = fullwidth_to_halfwidth(inp)
            if expected not in result:
                return False, f"全角转换失败: {inp} → {result}, 期望包含 {expected}"

        # normalize 测试
        norm_tests = [
            ("CJJ J 1-2014", "CJJ1-2014"),
            ("DGJ 08-2002", "DG/TJ08-2002"),
            ("GB/T  50430-2017", "GB/T50430-2017"),
        ]
        for inp, expected in norm_tests:
            result = normalize_for_matching(inp)
            if result != expected:
                return False, f"规范化失败: {inp} → {result}, 期望 {expected}"

        return True, f"全角转换 {len(tests)} 项, 规范化 {len(norm_tests)} 项全部通过"
    except Exception as e:
        return False, str(e)


# ═══════════════════════════════════════════════
# 5. 完整端到端流程（模拟 App 操作）
# ═══════════════════════════════════════════════

def test_e2e_pdf_to_database():
    """端到端：PDF → 文本提取 → 规范识别 → 数据库查询"""
    try:
        from standard_checker import CODE_PATTERN, normalize_for_matching
        from standard_db import StandardChecker
        import fitz

        # Step 1: 读取 PDF
        path = os.path.join(PROJECT_DIR, "test_files", "test_document.pdf")
        doc = fitz.open(path)
        full_text = ""
        for page in doc:
            full_text += page.get_text()
        doc.close()
        assert len(full_text) > 100, f"PDF 文本提取失败: {len(full_text)} 字符"

        # Step 2: 提取规范编号
        codes = CODE_PATTERN.findall(full_text)
        unique_codes = list(set(codes))
        assert len(unique_codes) > 0, "未提取到规范编号"

        # Step 3: 查询数据库
        checker = StandardChecker()
        query_results = []
        for code in unique_codes[:5]:  # 取前5个查询
            result = checker.check_code(code)
            query_results.append((code, result))

        checker.close()

        # Step 4: 统计
        found = sum(1 for _, r in query_results if r.get('found', False))

        return True, f"PDF → {len(full_text)}字符 → {len(unique_codes)}编号 → 查询{len(query_results)}条, 找到{found}条"
    except Exception as e:
        return False, str(e)


def test_e2e_docx_to_database():
    """端到端：Word → 文本提取 → 规范识别 → 数据库查询"""
    try:
        from standard_checker import CODE_PATTERN, normalize_for_matching
        from standard_db import StandardChecker
        from docx import Document

        # Step 1: 读取 Word
        path = os.path.join(PROJECT_DIR, "test_files", "test_document.docx")
        doc = Document(path)
        full_text = '\n'.join([p.text for p in doc.paragraphs])
        assert len(full_text) > 100, f"Word 文本提取失败: {len(full_text)} 字符"

        # Step 2: 提取规范编号
        codes = CODE_PATTERN.findall(full_text)
        unique_codes = list(set(codes))
        assert len(unique_codes) > 0, "未提取到规范编号"

        # Step 3: 查询数据库
        checker = StandardChecker()
        query_results = []
        for code in unique_codes[:5]:
            result = checker.check_code(code)
            query_results.append((code, result))
        checker.close()

        found = sum(1 for _, r in query_results if r.get('found', False))

        return True, f"Word → {len(full_text)}字符 → {len(unique_codes)}编号 → 查询{len(query_results)}条, 找到{found}条"
    except Exception as e:
        return False, str(e)


def test_e2e_txt_to_database():
    """端到端：TXT → 文本提取 → 规范识别 → 数据库查询"""
    try:
        from standard_checker import CODE_PATTERN, normalize_for_matching
        from standard_db import StandardChecker

        # Step 1: 读取 TXT
        path = os.path.join(PROJECT_DIR, "test_files", "test_document.txt")
        with open(path, 'r', encoding='utf-8') as f:
            full_text = f.read()
        assert len(full_text) > 50, f"TXT 文本提取失败: {len(full_text)} 字符"

        # Step 2: 提取规范编号
        codes = CODE_PATTERN.findall(full_text)
        unique_codes = list(set(codes))
        assert len(unique_codes) > 0, "未提取到规范编号"

        # Step 3: 查询数据库
        checker = StandardChecker()
        query_results = []
        for code in unique_codes[:5]:
            result = checker.check_code(code)
            query_results.append((code, result))
        checker.close()

        found = sum(1 for _, r in query_results if r.get('found', False))

        return True, f"TXT → {len(full_text)}字符 → {len(unique_codes)}编号 → 查询{len(query_results)}条, 找到{found}条"
    except Exception as e:
        return False, str(e)


# ═══════════════════════════════════════════════
# 主程序
# ═══════════════════════════════════════════════

def main():
    global PASS, FAIL, RESULTS

    print("=" * 60)
    print("  工程助手 LDAssistant — 端到端集成测试")
    print("=" * 60)

    tests = [
        # 1. 单文件提取
        ("PDF 文本提取", test_pdf_text_extraction),
        ("Word 文本提取", test_docx_text_extraction),
        ("TXT 文本提取", test_txt_text_extraction),
        ("DXF 图形渲染", test_dxf_render),

        # 2. 规范识别
        ("规范编号提取", test_code_extraction),

        # 3. 数据库
        ("数据库查询", test_database_query),

        # 4. OCR 预处理
        ("OCR 预处理函数", test_ocr_preprocessing),

        # 5. 完整端到端
        ("端到端 PDF→数据库", test_e2e_pdf_to_database),
        ("端到端 Word→数据库", test_e2e_docx_to_database),
        ("端到端 TXT→数据库", test_e2e_txt_to_database),
    ]

    print()
    for name, func in tests:
        try:
            ok, detail = func()
        except Exception as e:
            ok, detail = False, f"异常: {traceback.format_exc()[:200]}"
        log(name, ok, detail)

    # 汇总
    print()
    print("=" * 60)
    total = PASS + FAIL
    print(f"  总计: {total} 项测试")
    print(f"  通过: {PASS}/{total}  ✅")
    print(f"  失败: {FAIL}/{total}  {'❌' if FAIL else ''}")
    print("=" * 60)

    if FAIL == 0:
        print("\n🎉 全部通过！功能集成测试 OK！")
    else:
        print(f"\n⚠️ 有 {FAIL} 项失败，需要检查：")
        for name, ok, detail in RESULTS:
            if not ok:
                print(f"   - {name}: {detail}")

    return 0 if FAIL == 0 else 1


if __name__ == "__main__":
    sys.exit(main())
