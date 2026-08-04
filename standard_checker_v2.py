#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
工程助手 LDAssistant (增强版 v2)
...
"""
import matplotlib
matplotlib.use('Agg')  # 非交互后端，PyInstaller 打包后更稳定

VERSION = "2.0.0"
APP_NAME = "LDAssistant"
APP_TITLE = f"{APP_NAME} v{VERSION}"

import sqlite3
import os
import sys
import json
import re
import subprocess
import tempfile
import threading
import concurrent.futures
import time
import urllib.request
import urllib.parse
import io
from pathlib import Path
from datetime import datetime

import tkinter as tk
from tkinter import ttk, filedialog, messagebox

try:
    import fitz
    HAS_FITZ = True
except Exception:
    HAS_FITZ = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except Exception:
    HAS_DOCX = False


try:
    from PIL import Image, ImageDraw, ImageFilter, ImageTk, ImageFont
    HAS_PIL = True
except Exception:
    HAS_PIL = False

try:
    from standard_db import StandardChecker as SQLiteStandardChecker, normalize_for_matching as db_normalize_for_matching
    USE_SQLITE = True
except Exception:
    USE_SQLITE = False

try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    HAS_OPENPYXL = True
except Exception:
    HAS_OPENPYXL = False

# CAD support
try:
    import ezdxf
    from ezdxf.addons.drawing import Frontend, RenderContext
    from ezdxf.addons.drawing.matplotlib import MatplotlibBackend
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    HAS_CAD = True
except Exception as _cad_err:
    HAS_CAD = False
    print(f"[WARN] ezdxf/matplotlib 导入失败，DXF 渲染不可用: {_cad_err}", flush=True)

# Fix blurry text on high-DPI Windows displays
try:
    import ctypes
    ctypes.windll.shcore.SetProcessDpiAwareness(2)
except Exception:
    try:
        ctypes.windll.user32.SetProcessDPIAware()
    except Exception:
        pass

# Paths
if getattr(sys, 'frozen', False) and hasattr(sys, '_MEIPASS'):
    _APP_DIR = Path(sys._MEIPASS).resolve()
else:
    _APP_DIR = Path(__file__).parent.resolve()


def _find_ocr_path():
    """查找 PaddleOCR-json.exe 路径（打包目录内 ocr/ 或 UmiOCR 安装路径）"""
    # 1) 打包目录内的 ocr/（PyInstaller --add-data 或本地开发拷贝）
    bundled = _APP_DIR / "ocr" / "PaddleOCR-json.exe"
    if bundled.exists():
        return bundled, _APP_DIR / "ocr"
    # 2) 本地开发时的 UmiOCR 安装路径
    umi = Path(r"D:/Program Files/图片文字识别/UmiOCR-data/plugins/win7_x64_PaddleOCR-json")
    exe = umi / "PaddleOCR-json.exe"
    if exe.exists():
        return exe, umi
    return None, None


def _find_data_file():
    """查找标准数据库 JSON 文件"""
    for p in [
        _APP_DIR / "data" / "all_standards_merged_20260629_092235.json",
        _APP_DIR / "all_standards_merged_20260629_092235.json",
    ]:
        if p.exists():
            return p
    return None


PADDLE_OCR_EXE, OCR_DIR = _find_ocr_path()
DATA_FILE = _find_data_file()

# Patterns
CODE_PATTERN = re.compile(
    r'[A-Z]{1,5}[0-9]*(?:/[A-Z]{1,10})?\s*\d+(?:\.\d+)?-\d{4}',
    re.IGNORECASE)
NAME_PATTERN = re.compile(
    r'(?:[A-Z]{1,5}(?:/[A-Z]{1,2})?)\s*\d+(?:\.\d+)?-\d{4}\s+([\u4e00-\u9fff]{2,60})')

OBSOLETE_KEYWORDS = ['废止', '作废', '代替', '被代替', '被...代替']

IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.tif', '.webp'}
CAD_EXTENSIONS = {'.dwg', '.dxf'}
SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt'} | IMAGE_EXTENSIONS | CAD_EXTENSIONS


def fullwidth_to_halfwidth(text):
    """全角转半角 — 覆盖所有全角字符：数字、英文字母、标点符号"""
    result = []
    for ch in text:
        code = ord(ch)
        # 全角 ASCII 字母/数字/符号 (U+FF01 ~ U+FF5E)
        if 0xFF01 <= code <= 0xFF5E:
            result.append(chr(code - 0xFEE0))
        # 全角空格
        elif code == 0x3000:
            result.append(' ')
        # 全角括号、引号等 CJK 符号 (U+3000~U+303F, U+FE30~U+FE4F)
        elif code == 0x3002:  # 。
            result.append('.')
        elif code == 0x3001:  # 、
            result.append(',')
        elif code == 0x301C:  # 〜
            result.append('~')
        elif code == 0x3008:  # 〈
            result.append('<')
        elif code == 0x3009:  # 〉
            result.append('>')
        elif code == 0x300A:  # 《
            result.append('<')
        elif code == 0x300B:  # 》
            result.append('>')
        elif code == 0x3010:  # 【
            result.append('[')
        elif code == 0x3011:  # 】
            result.append(']')
        elif code == 0x3014:  # 〔
            result.append('[')
        elif code == 0x3015:  # 〕
            result.append(']')
        elif code == 0x2018 or code == 0x2019:  # '' 单引号
            result.append("'")
        elif code == 0x201C or code == 0x201D:  # "" 双引号
            result.append('"')
        elif code == 0x2013 or code == 0x2014:  # –—
            result.append('-')
        elif code == 0x2026:  # … 省略号
            result.append('...')
        # 全角数字兼容区域 (U+2460~U+24FF, 圈数字等) — 映射为普通数字
        elif 0x2460 <= code <= 0x2468:  # ①-⑨
            result.append(str(code - 0x245F))
        elif 0x2474 <= code <= 0x247C:  # ⑴-⑼
            result.append(str(code - 0x2473))
        elif 0x2488 <= code <= 0x2490:  # ⒈-⒐
            result.append(str(code - 0x2487))
        # 全角数字 (又一种) U+FF10-U+FF19 已在上面 U+FF01~U+FF5E 范围中处理
        else:
            result.append(ch)
    return ''.join(result)


def normalize_for_matching(text):
    """统一格式用于匹配 — 全角转半角 + 中文标点转英文 + 空格 + OCR 修正"""
    if not text:
        return ''
    result = fullwidth_to_halfwidth(text)
    punct_map = {
        # 中文标点 → 英文
        '\u3002': '.', '\u3001': ',', '\u301C': '~',
        '\u2014': '-', '\u2013': '-', '\u2026': '...',
        '\u201C': '"', '\u201D': '"', '\u2018': "'", '\u2019': "'",
        '\u00D7': 'x',        # 乘号 × → x
        '\u00F7': '/',        # 除号 ÷ → /
        '\u00B7': '-',        # 中间点 · → -
        '\u2022': '-',        # 圆点 • → -
        '\u2032': "'",        # 分 ′ → '
        '\u2033': '"',        # 秒 ″ → "
        '\u3008': '<',        # 〈
        '\u3009': '>',        # 〉
        '\u300A': '<',        # 《
        '\u300B': '>',        # 》
        '\u3010': '[',        # 【
        '\u3011': ']',        # 】
        '\u3014': '[',        # 〔
        '\u3015': ']',        # 〕
        '\uFF08': '(',        # 全角 (
        '\uFF09': ')',        # 全角 )
        '\uFF1A': ':',        # 全角 :
        '\uFF1B': ';',        # 全角 ;
        '\uFF0C': ',',        # 全角 ，
        '\u3000': ' ',        # 全角空格
    }
    for cn, en in punct_map.items():
        result = result.replace(cn, en)
    result = re.sub(r'\s+', '', result)
    # OCR 常见识别错误修正
    result = re.sub(r'CJJJ', 'CJJ', result, flags=re.IGNORECASE)
    result = re.sub(r'DGJ(?=\d)', 'DG/TJ', result, flags=re.IGNORECASE)
    result = re.sub(r'[LlI](?=[A-Z\d])', '1', result)  # 字母 l/I 在数字/字母前→1
    result = re.sub(r'(?<=[A-Z\d])[LlI]', '1', result)  # 字母 l/I 在数字/字母后→1
    result = re.sub(r'[Oo](?=\d)', '0', result)  # 字母 O 在数字前→0
    result = re.sub(r'(?<=\d)[Oo]', '0', result)  # 字母 O 在数字后→0
    result = result.upper()  # 统一大写
    return result


def ocr_image_standalone(image_path):
    """调用 PaddleOCR-json.exe 识别图片中的文字（不依赖 StandardChecker 实例）"""
    if PADDLE_OCR_EXE is None:
        return "OCR_ERROR: PaddleOCR 未找到", []
    cmd = [str(PADDLE_OCR_EXE), f"-image_path={image_path}"]
    try:
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30, cwd=str(OCR_DIR))
        ansi_escape = re.compile(r'\x1b\[[0-9;]*m')
        for line in result.stdout.split('\n'):
            line = ansi_escape.sub('', line).strip()
            if line.startswith('{'):
                try:
                    ocr_result = json.loads(line)
                    blocks = []
                    for item in ocr_result.get('data', []):
                        text = item.get('text', '')
                        box = item.get('box', [])
                        if box and len(box) == 4:
                            xs = [p[0] for p in box]
                            ys = [p[1] for p in box]
                            bbox = (min(xs), min(ys), max(xs), max(ys))
                        else:
                            bbox = (0, 0, 0, 0)
                        blocks.append((text, bbox))
                    text = ' '.join([b[0] for b in blocks])
                    return text, blocks
                except Exception as e:
                    print(f"OCR parse error: {e}, line: {line[:200]}")
        cleaned = ansi_escape.sub('', result.stdout).strip()
        return cleaned, []
    except Exception as e:
        return f"OCR_ERROR: {e}", []


class StandardChecker:
    """标准规范检查器"""

    def __init__(self):
        self.data = []
        self.code_index = {}
        self.name_index = {}
        self._sqlite_checker = None
        if USE_SQLITE:
            try:
                self._sqlite_checker = SQLiteStandardChecker()
                print("[StandardChecker] 已启用 SQLite + FTS5 加速")
            except Exception as e:
                print(f"[StandardChecker] SQLite 初始化失败，回退到 JSON: {e}")
                self._sqlite_checker = None
        if self._sqlite_checker is None:
            self.load_data()

    def load_data(self):
        if not DATA_FILE.exists():
            print(f"Data file not found: {DATA_FILE}")
            return
        print(f"Loading data from {DATA_FILE}...")
        data = None
        # 1) Try JSON with multiple encodings
        for enc in ['utf-8', 'utf-8-sig', 'gbk', 'gb18030']:
            try:
                with open(DATA_FILE, 'r', encoding=enc) as f:
                    data = json.load(f)
                print(f"  Loaded as JSON with encoding: {enc}")
                break
            except (UnicodeDecodeError, json.JSONDecodeError):
                continue
        # 2) Fallback: treat as SQLite database
        if data is None:
            try:
                import sqlite3
                conn = sqlite3.connect(str(DATA_FILE))
                try:
                    conn.row_factory = sqlite3.Row
                    cur = conn.cursor()
                    tables = [r['name'] for r in cur.execute(
                        "SELECT name FROM sqlite_master WHERE type='table'").fetchall()]
                    print(f"  SQLite tables found: {tables}")
                    target = 'standards' if 'standards' in tables else (tables[0] if tables else None)
                    if target:
                        cur.execute(f"SELECT * FROM {target}")
                        data = [dict(row) for row in cur.fetchall()]
                        print(f"  Loaded {len(data)} records from SQLite table '{target}'")
                finally:
                    conn.close()
            except Exception as e:
                print(f"  SQLite fallback failed: {e}")
        if data is None:
            raise RuntimeError(f"Cannot decode data file: {DATA_FILE}")
        self.data = data
        for r in self.data:
            code = normalize_for_matching(r.get('code', ''))
            if code:
                self.code_index[code] = r
            name = r.get('name', '').strip()
            if name:
                norm_name = normalize_for_matching(name)
                self.name_index[norm_name] = r
        print(f"Loaded {len(self.data)} records, indexed {len(self.code_index)} codes, {len(self.name_index)} names")

    def check_code(self, code, name=''):
        if self._sqlite_checker is not None:
            return self._sqlite_checker.check_code(code, name=name)
        normalized = normalize_for_matching(code)
        result = {'found': False, 'status': '未找到', 'replacement_raw': '', 'publisher': '', 'implement_date': ''}
        if normalized in self.code_index:
            r = self.code_index[normalized]
            result.update({'found': True, 'status': r.get('status', ''), 'replacement_raw': r.get('replacement_raw', ''),
                          'publisher': r.get('publisher', ''), 'implement_date': r.get('implement_date', ''),
                          'matched_name': r.get('name', '')})
            if name:
                norm_name = normalize_for_matching(name).strip()
                db_name = normalize_for_matching(r.get('name', '')).strip()
                if norm_name and db_name and (norm_name in db_name or db_name in norm_name):
                    result['dual_match'] = True
            return result
        if name:
            norm_name = normalize_for_matching(name).strip()
            if norm_name and norm_name in self.name_index:
                r = self.name_index[norm_name]
                result.update({'found': True, 'status': r.get('status', ''), 'replacement_raw': r.get('replacement_raw', ''),
                              'publisher': r.get('publisher', ''), 'implement_date': r.get('implement_date', ''),
                              'matched_name': r.get('name', '')})
                if code:
                    norm_code = normalize_for_matching(code).strip()
                    db_code = normalize_for_matching(r.get('code', '')).strip()
                    if norm_code and db_code and (norm_code in db_code or db_code in norm_code):
                        result['dual_match'] = True
                return result
            if norm_name and len(norm_name) >= 4:
                for k, v in self.name_index.items():
                    if norm_name in k or k in norm_name:
                        result.update({'found': True, 'status': v.get('status', ''),
                                      'replacement_raw': v.get('replacement_raw', ''),
                                      'publisher': v.get('publisher', ''),
                                      'implement_date': v.get('implement_date', ''),
                                      'matched_name': v.get('name', '')})
                        if code:
                            norm_code = normalize_for_matching(code).strip()
                            db_code = normalize_for_matching(v.get('code', '')).strip()
                            if norm_code and db_code and (norm_code in db_code or db_code in norm_code):
                                result['dual_match'] = True
                        return result
        best_match = None
        best_score = 0
        for k, v in self.code_index.items():
            if normalized in k or k in normalized:
                score = len(normalized) / max(len(k), len(normalized))
                if score > best_score:
                    best_score = score
                    best_match = v
            elif len(normalized) > 3 and len(k) > 3:
                matches = 0
                n_idx = 0
                k_idx = 0
                while n_idx < len(normalized) and k_idx < len(k):
                    if normalized[n_idx] == k[k_idx]:
                        matches += 1
                        n_idx += 1
                        k_idx += 1
                    else:
                        k_idx += 1
                similarity = matches / max(len(normalized), len(k))
                if similarity > 0.8 and similarity > best_score:
                    best_score = similarity
                    best_match = v
        if best_match:
            result.update({'found': True, 'status': best_match.get('status', ''),
                          'replacement_raw': best_match.get('replacement_raw', ''),
                          'publisher': best_match.get('publisher', ''),
                          'implement_date': best_match.get('implement_date', ''),
                          'matched_code': best_match.get('code', ''),
                          'matched_name': best_match.get('name', '')})
            if name:
                norm_name = normalize_for_matching(name).strip()
                db_name = normalize_for_matching(best_match.get('name', '')).strip()
                if norm_name and db_name and (norm_name in db_name or db_name in norm_name):
                    result['dual_match'] = True
            return result
        return result

    def find_similar_codes(self, code, limit=5):
        if self._sqlite_checker is not None:
            norm_code = db_normalize_for_matching(code)
            raw_code = code.strip()
            rows = []
            cur = None
            try:
                cur = self._sqlite_checker.conn.cursor()
                try:
                    cur.execute("SELECT code, name, status FROM standards_fts WHERE standards_fts MATCH ? LIMIT ?",
                                (norm_code, limit))
                    rows = cur.fetchall()
                except sqlite3.OperationalError:
                    rows = []
                if not rows and raw_code:
                    try:
                        cur.execute("SELECT code, name, status FROM standards WHERE code LIKE ? OR name LIKE ? LIMIT ?",
                                    (f'%{raw_code}%', f'%{raw_code}%', limit))
                        rows = cur.fetchall()
                    except sqlite3.OperationalError:
                        rows = []
            except Exception:
                rows = []
            finally:
                if cur:
                    cur.close()
            results = []
            for r in rows:
                results.append((r['code'], r['code'], r['name'], 'sqlite'))
            return results
        normalized = normalize_for_matching(code)
        results = []
        for k, v in self.code_index.items():
            if normalized in k or k in normalized:
                results.append((k, v.get('code', ''), v.get('name', ''), 'substring'))
            elif len(normalized) > 3 and len(k) > 3:
                matches = 0
                n_idx = 0
                k_idx = 0
                while n_idx < len(normalized) and k_idx < len(k):
                    if normalized[n_idx] == k[k_idx]:
                        matches += 1
                        n_idx += 1
                        k_idx += 1
                    else:
                        k_idx += 1
                similarity = matches / max(len(normalized), len(k))
                if similarity > 0.6:
                    results.append((k, v.get('code', ''), v.get('name', ''), f'similar:{similarity:.2f}'))
        results.sort(key=lambda x: x[3], reverse=True)
        return results[:limit]

    
    def close(self):
        if getattr(self, '_sqlite_checker', None) is not None:
            self._sqlite_checker.close()


class RegionSelector:
    """拖拽选择识别区域的辅助类"""

    def __init__(self, canvas, image_item_id, on_selected):
        self.canvas = canvas
        self.image_item_id = image_item_id
        self.on_selected = on_selected
        self.start_x = None
        self.start_y = None
        self.rect_id = None
        self.active = False

    def enable(self):
        self.active = True
        self.canvas.config(cursor="cross")
        self.canvas.bind('<ButtonPress-1>', self.on_press)
        self.canvas.bind('<B1-Motion>', self.on_drag)
        self.canvas.bind('<ButtonRelease-1>', self.on_release)

    def disable(self):
        self.active = False
        self.canvas.config(cursor="")
        self.canvas.unbind('<ButtonPress-1>')
        self.canvas.unbind('<B1-Motion>')
        self.canvas.unbind('<ButtonRelease-1>')
        if self.rect_id:
            self.canvas.delete(self.rect_id)
            self.rect_id = None

    def on_press(self, event):
        if not self.active:
            return
        self.start_x = self.canvas.canvasx(event.x)
        self.start_y = self.canvas.canvasy(event.y)
        if self.rect_id:
            self.canvas.delete(self.rect_id)
        self.rect_id = self.canvas.create_rectangle(
            self.start_x, self.start_y, self.start_x, self.start_y,
            outline='red', width=2, dash=(4, 2))

    def on_drag(self, event):
        if not self.active or self.rect_id is None:
            return
        cur_x = self.canvas.canvasx(event.x)
        cur_y = self.canvas.canvasy(event.y)
        self.canvas.coords(self.rect_id, self.start_x, self.start_y, cur_x, cur_y)

    def on_release(self, event):
        if not self.active or self.rect_id is None:
            return
        end_x = self.canvas.canvasx(event.x)
        end_y = self.canvas.canvasy(event.y)
        x1 = min(self.start_x, end_x)
        y1 = min(self.start_y, end_y)
        x2 = max(self.start_x, end_x)
        y2 = max(self.start_y, end_y)
        if abs(x2 - x1) < 10 or abs(y2 - y1) < 10:
            self.canvas.delete(self.rect_id)
            self.rect_id = None
            return
        if self.on_selected:
            self.on_selected((x1, y1, x2, y2))
        self.disable()


def mask_seals_pil(image_path, out_path=None):
    """红色印章遮盖"""
    if not HAS_PIL:
        return image_path
    try:
        img = Image.open(image_path).convert("RGB")
        w, h = img.size
        max_side = 1600
        if max(w, h) > max_side:
            scale = max_side / max(w, h)
            img_small = img.resize((int(w * scale), int(h * scale)), Image.Resampling.LANCZOS)
        else:
            img_small = img
            scale = 1.0
        hsv = img_small.convert("HSV")
        pixels = hsv.load()
        mask = Image.new("L", img_small.size, 0)
        mask_pixels = mask.load()
        for y in range(img_small.size[1]):
            for x in range(img_small.size[0]):
                h_val, s_val, v_val = pixels[x, y]
                if (h_val < 12 or h_val > 160) and s_val > 60 and v_val > 60:
                    mask_pixels[x, y] = 255
        mask = mask.filter(ImageFilter.MaxFilter(3))
        red_count = sum(1 for p in mask.getdata() if p > 128)
        total = mask.width * mask.height
        should_mask = red_count > total * 0.08
        if should_mask:
            if scale != 1.0:
                mask = mask.resize((w, h), Image.Resampling.NEAREST)
            white = Image.new("RGB", (w, h), (255, 255, 255))
            img = Image.composite(white, img, mask)
            if out_path:
                img.save(out_path)
                return out_path
            fd, tmp_path = tempfile.mkstemp(suffix=".png"); os.close(fd)
            try:
                img.save(tmp_path)
                return tmp_path
            except Exception:
                try:
                    os.unlink(tmp_path)
                except Exception:
                    pass
                raise
    except Exception as e:
        print(f"mask_seals error: {e}")
        return image_path


def render_cad_to_image(dxf_path, dpi=300):
    """将 CAD 文件（DWG/DXF）渲染为 PNG 图片"""
    if not HAS_CAD:
        return None
    try:
        # 设置 matplotlib 中文字体，避免 DXF 中文文字显示为方框
        import matplotlib
        matplotlib.rcParams['font.family'] = 'sans-serif'
        matplotlib.rcParams['font.sans-serif'] = [
            'SimHei', 'Microsoft YaHei', 'SimSun',
            'DejaVu Sans', 'Arial'
        ]
        matplotlib.rcParams['axes.unicode_minus'] = False

        doc = ezdxf.readfile(dxf_path)
        fig = plt.figure(figsize=(12, 9), dpi=dpi)
        ax = fig.add_axes([0, 0, 1, 1])
        ax.set_aspect('equal')
        ax.axis('off')
        backend = MatplotlibBackend(ax)
        Frontend(RenderContext(doc), backend).draw_layout(doc.modelspace())
        fd, tmp = tempfile.mkstemp(suffix=".png"); os.close(fd)
        try:
            fig.savefig(tmp, dpi=dpi, bbox_inches="tight", pad_inches=0.1,
                         facecolor="white", edgecolor="none")
            plt.close(fig)
            return tmp
        except Exception:
            plt.close(fig)
            try:
                os.unlink(tmp)
            except Exception:
                pass
            raise
    except Exception as e:
        print(f"CAD render error: {e}")
        return None


class StandardSearchDialog(tk.Toplevel):
    """规范搜索与推荐弹窗"""

    def __init__(self, parent, checker, code='', name=''):
        super().__init__(parent)
        self.checker = checker
        self.code = code
        self.name = name
        self.title("规范搜索与推荐")
        self.geometry("750x550")
        self.minsize(600, 450)
        self.transient(parent)
        self.grab_set()
        self._setup_ui()
        self._search_recommend()

    def _setup_ui(self):
        """暗色主题搜索对话框"""
        C = {'bg': "#1E293B", 'bg_dark': "#0F172A", 'card': "#334155",
             'text': "#E2E8F0", 'text_muted': "#94A3B8",
             'primary': "#3B82F6", 'primary_hover': "#2563EB"}
        self.configure(bg=C['bg_dark'])

        # 搜索行
        search_frame = tk.Frame(self, bg=C['bg_dark'])
        search_frame.pack(side=tk.TOP, fill=tk.X, padx=10, pady=(10, 0))
        tk.Label(search_frame, text="🔍 搜索规范:",
                 font=("Microsoft YaHei UI", 10),
                 bg=C['bg_dark'], fg=C['text']).pack(side=tk.LEFT, padx=(0, 8))
        self.search_var = tk.StringVar()
        search_entry = tk.Entry(search_frame, textvariable=self.search_var, width=50,
                                font=("Microsoft YaHei UI", 10),
                                bg=C['card'], fg=C['text'],
                                insertbackground=C['text'],
                                borderwidth=0, highlightthickness=0)
        search_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 5), ipady=4)
        search_entry.bind('<Return>', lambda e: self._do_search())
        # 搜索按钮
        btn_search = tk.Label(search_frame, text="搜索",
                              font=("Microsoft YaHei UI", 9, "bold"),
                              bg=C['primary'], fg="#FFFFFF", cursor="hand2", padx=12, pady=2)
        btn_search.pack(side=tk.LEFT, padx=(0, 3))
        btn_search.bind('<Button-1>', lambda e: self._do_search())
        btn_search.bind('<Enter>', lambda e: btn_search.config(bg=C['primary_hover']))
        btn_search.bind('<Leave>', lambda e: btn_search.config(bg=C['primary']))
        btn_recommend = tk.Label(search_frame, text="推荐相近",
                                 font=("Microsoft YaHei UI", 9),
                                 bg=C['card'], fg=C['text'], cursor="hand2", padx=10, pady=2)
        btn_recommend.pack(side=tk.LEFT)
        btn_recommend.bind('<Button-1>', lambda e: self._search_recommend())
        btn_recommend.bind('<Enter>', lambda e: btn_recommend.config(bg=C['primary']))
        btn_recommend.bind('<Leave>', lambda e: btn_recommend.config(bg=C['card']))

        # 信息标签
        if self.code:
            info_text = f"当前规范: {self.code}" + (f"  《{self.name}》" if self.name else "")
        else:
            info_text = "请输入编号或名称搜索规范"
        self.info_label = tk.Label(self, text=info_text,
                                   font=("Microsoft YaHei UI", 9),
                                   bg=C['bg_dark'], fg=C['text_muted'], anchor=tk.W)
        self.info_label.pack(side=tk.TOP, anchor=tk.W, padx=10, pady=(4, 0))

        # 结果列表
        results_frame = tk.Frame(self, bg=C['bg'])
        results_frame.pack(side=tk.TOP, fill=tk.BOTH, expand=True, padx=10, pady=(10, 0))
        columns = ('code', 'name', 'status', 'action')
        self.result_tree = ttk.Treeview(results_frame, columns=columns, show='headings', height=15)
        self.result_tree.heading('code', text='规范编号')
        self.result_tree.heading('name', text='规范名称')
        self.result_tree.heading('status', text='状态')
        self.result_tree.heading('action', text='操作')
        self.result_tree.column('code', width=160)
        self.result_tree.column('name', width=350)
        self.result_tree.column('status', width=100)
        self.result_tree.column('action', width=120)
        self.result_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scroll = ttk.Scrollbar(results_frame, orient=tk.VERTICAL, command=self.result_tree.yview)
        self.result_tree.configure(yscrollcommand=scroll.set)
        scroll.pack(side=tk.RIGHT, fill=tk.Y)

        # 底部按钮
        btn_frame = tk.Frame(self, bg=C['bg_dark'])
        btn_frame.pack(side=tk.BOTTOM, fill=tk.X, padx=10, pady=10)
        btn_close = tk.Label(btn_frame, text="关闭",
                             font=("Microsoft YaHei UI", 9),
                             bg=C['card'], fg=C['text'], cursor="hand2", padx=16, pady=2)
        btn_close.pack(side=tk.RIGHT)
        btn_close.bind('<Button-1>', lambda e: self.destroy())
        btn_close.bind('<Enter>', lambda e: btn_close.config(bg=C['primary']))
        btn_close.bind('<Leave>', lambda e: btn_close.config(bg=C['card']))

    def _do_search(self):
        query = self.search_var.get().strip()
        if not query:
            messagebox.showwarning("提示", "请输入搜索内容")
            return
        self._perform_search(query)

    def _search_recommend(self):
        if self.code:
            self._perform_search(self.code)

    def _perform_search(self, query):
        self.result_tree.delete(*self.result_tree.get_children())
        norm_query = normalize_for_matching(query)
        results = []
        for k, v in self.checker.code_index.items():
            code = v.get('code', '')
            name = v.get('name', '')
            norm_code = normalize_for_matching(code)
            norm_name = normalize_for_matching(name)
            score = 0
            if norm_query in norm_code or norm_code in norm_query:
                score = max(score, len(norm_query) / max(len(norm_code), 1))
            if norm_query in norm_name or norm_name in norm_query:
                score = max(score, len(norm_query) / max(len(norm_name), 1))
            if score > 0:
                results.append((score, v))
        results.sort(key=lambda x: x[0], reverse=True)
        results = results[:50]
        for _, v in results:
            code = v.get('code', '')
            name = v.get('name', '')
            status = v.get('status', '')
            is_obsolete = '废止' in status or '作废' in status
            item_id = self.result_tree.insert('', tk.END, values=(code, name, status, ''))
            if is_obsolete:
                self.result_tree.item(item_id, tags=('obsolete',))
            else:
                self.result_tree.item(item_id, tags=('active',))
        self.result_tree.tag_configure('obsolete', foreground='red')
        self.result_tree.tag_configure('active', foreground='green')
        self.result_tree.bind('<Double-Button-1>', self._on_result_double_click)

    def _on_result_double_click(self, event):
        item = self.result_tree.selection()
        if not item:
            return
        values = self.result_tree.item(item[0], 'values')
        code, name, status = values[0], values[1], values[2]
        copy_win = tk.Toplevel(self)
        copy_win.title("复制规范信息")
        copy_win.geometry("500x200")
        copy_win.transient(self)
        copy_win.grab_set()
        ttk.Label(copy_win, text=f"规范编号: {code}", font=("SimSun", 11)).pack(anchor=tk.W, padx=10, pady=(10, 5))
        ttk.Label(copy_win, text=f"规范名称: {name}", font=("SimSun", 11)).pack(anchor=tk.W, padx=10, pady=5)
        ttk.Label(copy_win, text=f"状态: {status}",
                  foreground="red" if '废止' in status or '作废' in status else "green").pack(
            anchor=tk.W, padx=10, pady=5)
        btn_frame = ttk.Frame(copy_win, padding=10)
        btn_frame.pack(side=tk.BOTTOM, fill=tk.X)

        def copy_code():
            self.clipboard_clear()
            self.clipboard_append(code)
            messagebox.showinfo("完成", "规范编号已复制", parent=copy_win)

        def copy_name():
            self.clipboard_clear()
            self.clipboard_append(name)
            messagebox.showinfo("完成", "规范名称已复制", parent=copy_win)

        ttk.Button(btn_frame, text=f"复制编号: {code[:30]}...", command=copy_code).pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(btn_frame, text=f"复制名称: {name[:30]}...", command=copy_name).pack(side=tk.LEFT)
        ttk.Button(btn_frame, text="关闭", command=copy_win.destroy).pack(side=tk.RIGHT)


# ===== AI 配置持久化 =====
_CONFIG_FILE = Path.home() / ".ldassistant_config.json"
_USAGE_FILE = Path.home() / ".ldassistant_usage.json"

# 免费模型配置（无需用户设置即可使用，100次免费额度）
_FREE_MODEL_CONFIG = {
    "api_url": "https://open.bigmodel.cn/api/paas/v4/chat/completions",
    "api_key": "1d0f8e5f7c3a4b5e9d8f2a1c6b7e3d5f",
    "model": "glm-4-flash",
}
_FREE_QUOTA = 100  # 免费额度100次

def _load_ai_config():
    """加载 AI API 配置"""
    default = {
        "api_url": _FREE_MODEL_CONFIG["api_url"],
        "api_key": _FREE_MODEL_CONFIG["api_key"],
        "model": _FREE_MODEL_CONFIG["model"],
        "use_free_model": True,  # 默认使用免费模型
    }
    try:
        if _CONFIG_FILE.exists():
            with open(_CONFIG_FILE, 'r', encoding='utf-8') as f:
                return {**default, **json.load(f)}
    except Exception:
        pass
    return default

def _load_usage():
    """加载免费使用次数"""
    try:
        if _USAGE_FILE.exists():
            with open(_USAGE_FILE, 'r', encoding='utf-8') as f:
                return json.load(f).get("free_uses", 0)
    except Exception:
        pass
    return 0

def _save_usage(count):
    """保存免费使用次数"""
    try:
        with open(_USAGE_FILE, 'w', encoding='utf-8') as f:
            json.dump({"free_uses": count}, f)
    except Exception:
        pass

def _save_ai_config(config):
    """保存 AI API 配置"""
    try:
        with open(_CONFIG_FILE, 'w', encoding='utf-8') as f:
            json.dump(config, f, ensure_ascii=False, indent=2)
        return True
    except Exception as e:
        print(f"保存配置失败: {e}")
        return False

# ===== AI 聊天悬浮窗 =====
class AIChatFloatingWindow:
    """AI 聊天悬浮窗"""

    def __init__(self, master, config=None):
        self.master = master
        self.config = config or _load_ai_config()
        self._free_uses = _load_usage()
        self.window = tk.Toplevel(master)
        self.window.title("AI 助手")
        self.window.geometry("420x560")
        self.window.minsize(300, 350)
        # 设置 AI 窗口图标
        self._set_ai_icon()
        self._floating = True
        self._pin_icon = "📌"
        self._float_icon = "💠"
        self._offset_x = 0
        self._offset_y = 0
        self._messages = []
        self._links = []
        self._executor = concurrent.futures.ThreadPoolExecutor(max_workers=1)
        # 颜色常量 & 字体
        self._C = {
            'bg': "#1E293B", 'bg_dark': "#0F172A", 'card': "#334155",
            'text': "#E2E8F0", 'text_muted': "#94A3B8",
            'primary': "#3B82F6", 'primary_hover': "#2563EB",
            'select': "#1E40AF", 'success': "#22C55E", 'danger': "#EF4444",
            'font': ("Microsoft YaHei UI", 10),
            'font_md': ("Microsoft YaHei UI", 10, "bold"),
            'font_sm': ("Microsoft YaHei UI", 9),
            'font_bold': ("Microsoft YaHei UI", 9, "bold"),
        }
        self._setup_ui()
        self._setup_drag()
        remaining = max(0, _FREE_QUOTA - self._free_uses)
        self.add_message("ai",
            f"你好！我是标准查询 AI 助手。\n"
            f"发送标准号或关键词，我可以帮你查询国家标准。\n\n"
            f"🎁 当前使用免费模型（glm-4-flash），剩余 {remaining}/{_FREE_QUOTA} 次。\n"
            f"用完后可在「⚙️ 配置」中设置自己的 API Key。\n\n"
            f"OCR 识别的结果会自动显示在这里。")

    def _set_ai_icon(self):
        """给 AI 聊天窗口设置图标（兼容 PyInstaller onedir 路径）"""
        for p in [_APP_DIR / 'app_icon.ico', _APP_DIR / 'app_icon.png']:
            if p.exists():
                try:
                    if p.suffix == '.ico':
                        self.window.iconbitmap(str(p))
                        return
                    else:
                        img = tk.PhotoImage(file=str(p))
                        self.window.iconphoto(True, img)
                        self._icon_image = img
                        return
                except Exception:
                    continue

    def _setup_ui(self):
        """暗色即时通讯风格 AI 对话 UI"""
        C = self._C
        # ── 标题栏 ──
        titlebar = tk.Frame(self.window, bg=C['bg_dark'], height=38)
        titlebar.pack(side=tk.TOP, fill=tk.X)
        titlebar.pack_propagate(False)
        titlebar.bind('<ButtonPress-1>', self._start_drag)
        titlebar.bind('<B1-Motion>', self._on_drag)
        self._title_label = tk.Label(titlebar, text="🤖 AI 助手",
                                     font=C['font_md'], bg=C['bg_dark'], fg=C['text'])
        self._title_label.pack(side=tk.LEFT, padx=12)
        self._title_label.bind('<ButtonPress-1>', self._start_drag)
        self._title_label.bind('<B1-Motion>', self._on_drag)

        # 标题栏右侧按钮（扁平暗色）
        btn_frame = tk.Frame(titlebar, bg=C['bg_dark'])
        btn_frame.pack(side=tk.RIGHT, padx=6)
        for txt, tooltip, cmd in [
            (self._pin_icon, "置顶", self._toggle_pin),
            ("⚙️", "设置", self._open_config),
            ("🗑", "清空", self._clear_chat),
            ("💾", "导出", self._export_chat),
            ("—", "最小化", self._minimize),
            ("✕", "关闭", self._close),
        ]:
            b = tk.Label(btn_frame, text=txt, font=C['font_sm'],
                         bg=C['bg_dark'], fg=C['text_muted'], cursor="hand2", padx=4)
            b.pack(side=tk.LEFT, padx=1)
            b.bind('<Button-1>', lambda e, c=cmd: c())
            b.bind('<Enter>', lambda e, b=b: b.config(fg=C['text']))
            b.bind('<Leave>', lambda e, b=b: b.config(fg=C['text_muted']))
        self._pin_btn = btn_frame.winfo_children()[0]

        # ── 快捷查询行 ──
        quick_row = tk.Frame(self.window, bg=C['bg'], height=34)
        quick_row.pack(side=tk.TOP, fill=tk.X)
        quick_row.pack_propagate(False)
        quick_frame = tk.Frame(quick_row, bg=C['bg'])
        quick_frame.pack(expand=True, padx=8)
        for txt, q in [("🏗️ 混凝土", "混凝土结构"), ("🔥 防火", "建筑防火"),
                       ("📋 GB 50068", "GB 50068"), ("🔍 最新", "最新发布")]:
            btn = tk.Label(quick_frame, text=txt, font=C['font_sm'],
                           bg=C['card'], fg=C['text'], padx=8, pady=2,
                           cursor="hand2")
            btn.pack(side=tk.LEFT, padx=3)
            btn.bind('<Button-1>', lambda e, query=q: self._send_query(query))
            btn.bind('<Enter>', lambda e, b=btn: b.config(bg=C['primary']))
            btn.bind('<Leave>', lambda e, b=btn: b.config(bg=C['card']))

        # ── 消息区域 ──
        msg_frame = tk.Frame(self.window, bg=C['bg'])
        msg_frame.pack(side=tk.TOP, fill=tk.BOTH, expand=True)
        self._msg_canvas = tk.Canvas(msg_frame, bg=C['bg'], highlightthickness=0, bd=0)
        msg_scroll = tk.Scrollbar(msg_frame, orient=tk.VERTICAL, command=self._msg_canvas.yview,
                                  bg=C['card'], fg=C['text_muted'],
                                  troughcolor=C['bg_dark'],
                                  activebackground=C['primary'],
                                  highlightthickness=0, bd=0)
        self._msg_canvas.configure(yscrollcommand=msg_scroll.set)
        msg_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self._msg_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self._msg_inner = tk.Frame(self._msg_canvas, bg=C['bg'])
        self._msg_window = self._msg_canvas.create_window(
            (0, 0), window=self._msg_inner, anchor='nw', width=380)
        self._msg_inner.bind('<Configure>', self._on_msg_configure)

        # ── 输入区域 ──
        input_frame = tk.Frame(self.window, bg=C['bg_dark'])
        input_frame.pack(side=tk.TOP, fill=tk.X, padx=0, pady=0)
        # 圆角输入框容器
        entry_bg = tk.Frame(input_frame, bg=C['card'], bd=0, highlightthickness=0)
        entry_bg.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(8, 4), pady=8)
        self._input_entry = tk.Text(entry_bg, height=2, font=C['font'],
                                    wrap=tk.WORD, bd=0, padx=10, pady=6,
                                    bg=C['card'], fg=C['text'],
                                    insertbackground=C['text'],
                                    highlightthickness=0)
        self._input_entry.pack(fill=tk.BOTH, expand=True)
        self._input_entry.bind('<Return>', self._on_input_enter)
        self._input_entry.bind('<KeyRelease>', self._on_input_keyrelease)
        # 发送按钮
        send_frame = tk.Frame(input_frame, bg=C['bg_dark'])
        send_frame.pack(side=tk.RIGHT, padx=(0, 8), pady=8)
        self._send_btn = tk.Label(send_frame, text="➤ 发送",
                                  font=C['font_bold'],
                                  bg=C['primary'], fg="#FFFFFF",
                                  padx=14, pady=5, cursor="hand2")
        self._send_btn.pack()
        self._send_btn.bind('<Button-1>', lambda e: self._send_input())
        self._send_btn.bind('<Enter>', lambda e: self._send_btn.config(bg=C['primary_hover']))
        self._send_btn.bind('<Leave>', lambda e: self._send_btn.config(bg=C['primary']))

        # ── 状态栏 ──
        remaining = max(0, _FREE_QUOTA - self._free_uses)
        self._status_label = tk.Label(self.window, text=f"就绪 ✓  |  免费剩余: {remaining}/{_FREE_QUOTA} 次",
            font=C['font_sm'], fg=C['text_muted'],
            bg=C['bg_dark'], anchor=tk.W)
        self._status_label.pack(side=tk.BOTTOM, fill=tk.X, padx=10, pady=3)

        self.window.protocol("WM_DELETE_WINDOW", self._close)
        self.window.bind('<Escape>', lambda e: self._minimize())

    def _setup_drag(self):
        self.window.bind('<ButtonPress-1>', self._start_drag)
        self.window.bind('<B1-Motion>', self._on_drag)

    def _start_drag(self, event):
        if self._floating:
            self._offset_x = event.x
            self._offset_y = event.y

    def _on_drag(self, event):
        if self._floating and hasattr(self, '_offset_x'):
            x = self.window.winfo_x() + event.x - self._offset_x
            y = self.window.winfo_y() + event.y - self._offset_y
            self.window.geometry(f"+{x}+{y}")

    def _toggle_pin(self):
        self._floating = not self._floating
        if self._floating:
            self._pin_btn.config(text=self._pin_icon)
            self.window.attributes('-topmost', True)
            self._title_label.config(text="🤖 AI 助手 (浮动)")
        else:
            self._pin_btn.config(text=self._float_icon)
            self.window.attributes('-topmost', False)
            self._title_label.config(text="🤖 AI 助手")

    def _minimize(self):
        self.window.withdraw()

    def _close(self):
        self.window.withdraw()

    def show(self):
        self.window.deiconify()
        self.window.lift()

    def _clear_chat(self):
        self._messages = []
        self._links = []
        self._inline_images = []
        for w in self._msg_inner.winfo_children():
            w.destroy()
        self.add_message("ai", "对话已清空，可以重新开始提问。")

    def _export_chat(self):
        """导出 AI 对话记录为 Word 或 PDF（暗色主题）"""
        C = self._C
        dialog = tk.Toplevel(self.window)
        dialog.title("导出对话记录")
        dialog.configure(bg=C['bg_dark'])
        dialog.geometry("360x150")
        dialog.transient(self.window)
        dialog.grab_set()
        if hasattr(self.master, '_set_icon_for_toplevel'):
            self.master._set_icon_for_toplevel(dialog)
        tk.Label(dialog, text="请选择导出格式:", font=("Microsoft YaHei UI", 11),
                 bg=C['bg_dark'], fg=C['text']).pack(anchor=tk.W, padx=16, pady=(14, 10))

        def do_export_docx():
            dialog.destroy()
            self._do_export_chat_docx()

        def do_export_pdf():
            dialog.destroy()
            self._do_export_chat_pdf()

        btn_frame = tk.Frame(dialog, bg=C['bg_dark'])
        btn_frame.pack(fill=tk.X, padx=16, pady=8)

        for txt, cmd in [("📄 导出 Word", do_export_docx), ("📕 导出 PDF", do_export_pdf)]:
            b = tk.Label(btn_frame, text=txt, font=C['font_bold'],
                         bg=C['primary'], fg="#FFFFFF", cursor="hand2", padx=12, pady=3)
            b.pack(side=tk.LEFT, padx=4)
            b.bind('<Button-1>', lambda e, c=cmd: c())
            b.bind('<Enter>', lambda e, b=b: b.config(bg=C['primary_hover']))
            b.bind('<Leave>', lambda e, b=b: b.config(bg=C['primary']))

        cancel = tk.Label(btn_frame, text="取消", font=C['font_sm'],
                          bg=C['card'], fg=C['text'], cursor="hand2", padx=12, pady=3)
        cancel.pack(side=tk.RIGHT, padx=4)
        cancel.bind('<Button-1>', lambda e: dialog.destroy())
        cancel.bind('<Enter>', lambda e: cancel.config(bg=C['primary_hover']))
        cancel.bind('<Leave>', lambda e: cancel.config(bg=C['card']))

    def _build_chat_text(self):
        """将对话历史拼接为文本（返回 (title_lines, body_lines)）"""
        from datetime import datetime
        title = f"LDAssistant AI 对话记录 — {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        body = []
        for m in self._messages:
            role = "🤖 AI" if m.get("role") == "ai" else "👤 用户"
            body.append(f"{role}：")
            body.append(m.get("content", "") + "\n")
            body.append("")
        return title, body

    def _do_export_chat_docx(self):
        """导出对话为 Word 文档"""
        if not HAS_DOCX:
            messagebox.showwarning("提示", "需要安装 python-docx 库才能导出 Word", parent=self.window)
            return
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        try:
            title, body = self._build_chat_text()
            doc = Document()
            p = doc.add_heading(title, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for line in body:
                para = doc.add_paragraph(line)
                for run in para.runs:
                    run.font.name = "SimSun"
                    run.font.size = Pt(11)
            path = filedialog.asksaveasfilename(
                title="保存对话记录 (Word)",
                defaultextension=".docx",
                filetypes=[("Word 文档", "*.docx"), ("所有文件", "*.*")])
            if path:
                doc.save(path)
                messagebox.showinfo("完成", f"对话已导出:\n{path}", parent=self.window)
        except Exception as e:
            messagebox.showerror("错误", f"导出失败: {e}", parent=self.window)

    def _do_export_chat_pdf(self):
        """导出对话为 PDF 文档"""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import cm
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib import colors
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        except Exception:
            messagebox.showwarning("提示", "需要安装 reportlab 库才能导出 PDF，请在命令行执行:\npip install reportlab",
                                   parent=self.window)
            return
        try:
            pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
            title, body = self._build_chat_text()
            path = filedialog.asksaveasfilename(
                title="保存对话记录 (PDF)",
                defaultextension=".pdf",
                filetypes=[("PDF 文档", "*.pdf"), ("所有文件", "*.*")])
            if not path:
                return
            doc = SimpleDocTemplate(path, pagesize=A4,
                                    leftMargin=2*cm, rightMargin=2*cm,
                                    topMargin=2*cm, bottomMargin=2*cm)
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle('Title', parent=styles['Title'], fontName='STSong-Light', fontSize=16, spaceAfter=20, alignment=1)
            story = [Paragraph(title, title_style), Spacer(1, 0.3*cm)]
            ai_style = ParagraphStyle('AI', parent=styles['Normal'], fontName='STSong-Light', fontSize=10, spaceAfter=8, textColor=colors.HexColor('#1a73e8'))
            user_style = ParagraphStyle('User', parent=styles['Normal'], fontName='STSong-Light', fontSize=10, spaceAfter=8, textColor=colors.HexColor('#333333'))
            for m in self._messages:
                role = m.get("role")
                role_label = "🤖 AI" if role == "ai" else "👤 用户"
                style = ai_style if role == "ai" else user_style
                content = m.get("content", "").replace("\n", "<br/>")
                content = content.replace("**", "").replace("•", "- ")
                story.append(Paragraph(f"<b>{role_label}：</b>", style))
                story.append(Paragraph(content, style))
                story.append(Spacer(1, 0.2*cm))
            doc.build(story)
            messagebox.showinfo("完成", f"对话已导出:\n{path}", parent=self.window)
        except Exception as e:
            messagebox.showerror("错误", f"导出 PDF 失败: {e}", parent=self.window)

    def _open_config(self):
        """打开AI API配置对话框（暗色主题）"""
        C = self._C
        cfg = tk.Toplevel(self.window)
        cfg.title("AI API 配置")
        cfg.configure(bg=C['bg_dark'])
        cfg.geometry("480x340")
        cfg.transient(self.window)
        cfg.grab_set()

        # 标题
        tk.Label(cfg, text="⚙️ API 配置", font=("Microsoft YaHei UI", 12, "bold"),
            bg=C['bg_dark'], fg=C['text']).pack(anchor=tk.W, padx=14, pady=(12, 2))
        remaining = max(0, _FREE_QUOTA - _load_usage())
        tk.Label(cfg, text=f"免费模型可用，剩余 {remaining}/{_FREE_QUOTA} 次", font=("Microsoft YaHei UI", 9),
            bg=C['bg_dark'], fg=C['success']).pack(anchor=tk.W, padx=14, pady=(0, 8))

        frame = tk.Frame(cfg, bg=C['bg'], padx=14, pady=8)
        frame.pack(fill=tk.BOTH, expand=True)

        # API 地址
        tk.Label(frame, text="API 地址（留空默认使用免费模型）", font=("Microsoft YaHei UI", 9),
            bg=C['bg'], fg=C['text']).pack(anchor=tk.W, pady=(4, 2))
        api_url_var = tk.StringVar(value=self.config.get("api_url", _FREE_MODEL_CONFIG["api_url"]))
        url_entry = tk.Entry(frame, textvariable=api_url_var, width=50,
                             font=("Microsoft YaHei UI", 9),
                             bg=C['card'], fg=C['text'],
                             insertbackground=C['text'],
                             borderwidth=0, highlightthickness=0)
        url_entry.pack(fill=tk.X, ipady=4)

        # API Key
        key_frame = tk.Frame(frame, bg=C['bg'])
        key_frame.pack(fill=tk.X, pady=(8, 0))
        tk.Label(key_frame, text="API Key（可选）", font=("Microsoft YaHei UI", 9),
                 bg=C['bg'], fg=C['text']).pack(anchor=tk.W)
        key_row = tk.Frame(key_frame, bg=C['bg'])
        key_row.pack(fill=tk.X, pady=(2, 0))
        api_key_var = tk.StringVar(value=self.config.get("api_key", ""))
        key_entry = tk.Entry(key_row, textvariable=api_key_var, width=45, show="*",
                             font=("Microsoft YaHei UI", 9),
                             bg=C['card'], fg=C['text'],
                             insertbackground=C['text'],
                             borderwidth=0, highlightthickness=0)
        key_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, ipady=4)
        def toggle_key_show():
            key_entry.config(show='' if key_entry.cget('show') == '*' else '*')
        btn_eye = tk.Label(key_row, text="👁", font=("Microsoft YaHei UI", 9),
                           bg=C['card'], fg=C['text'], cursor="hand2", padx=6)
        btn_eye.pack(side=tk.LEFT, padx=(4, 0))
        btn_eye.bind('<Button-1>', lambda e: toggle_key_show())

# 模型 ID
        tk.Label(frame, text="模型 ID", font=("Microsoft YaHei UI", 9),
                 bg=C['bg'], fg=C['text']).pack(anchor=tk.W, pady=(8, 2))
        model_var = tk.StringVar(value=self.config.get("model", _FREE_MODEL_CONFIG["model"]))
        # 暗色自定义下拉框（替代 ttk.Combobox）
        model_row = tk.Frame(frame, bg=C['bg'])
        model_row.pack(fill=tk.X, pady=(2, 0))
        model_entry = tk.Entry(model_row, textvariable=model_var, width=40,
                               font=("Microsoft YaHei UI", 9),
                               bg=C['card'], fg=C['text'],
                               insertbackground=C['text'],
                               borderwidth=0, highlightthickness=0)
        model_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, ipady=4)
        model_btn = tk.Label(model_row, text="▼", font=("Microsoft YaHei UI", 9),
                             bg=C['card'], fg=C['text_muted'], cursor="hand2", padx=8)
        model_btn.pack(side=tk.LEFT, padx=(0, 0))
        model_btn.bind('<Button-1>', lambda e: _show_model_menu())
        model_values = ('glm-4-flash', 'glm-4', 'glm-4-air', 'gpt-4o-mini', 'gpt-4o', 'deepseek-chat', 'qwen-plus')

        def _show_model_menu():
            menu = tk.Menu(model_row, tearoff=0,
                           bg=C['card'], fg=C['text'],
                           activebackground=C['primary'], activeforeground="#FFFFFF",
                           font=("Microsoft YaHei UI", 9))
            for val in model_values:
                menu.add_command(label=val, command=lambda v=val: model_var.set(v))
            menu.tk_popup(model_btn.winfo_rootx(), model_btn.winfo_rooty() + model_btn.winfo_height())

        # 状态标签
        status_lbl = tk.Label(frame, text="", font=("Microsoft YaHei UI", 9),
                              bg=C['bg'], fg=C['success'])
        status_lbl.pack(anchor=tk.W, pady=6)

        def test_connection():
            status_lbl.config(text="测试中...", fg=C['primary'])
            cfg.update()
            def do_test():
                try:
                    url = api_url_var.get().strip()
                    data = json.dumps({"message": "你好", "stream": False}).encode('utf-8')
                    req = urllib.request.Request(url, data=data,
                                                 headers={'Content-Type': 'application/json'}, method='POST')
                    resp = urllib.request.urlopen(req, timeout=10)
                    result = json.loads(resp.read().decode('utf-8'))
                    reply = result.get('reply', '') or result.get('content', '') or 'OK'
                    cfg.after(0, lambda: status_lbl.config(text=f"✅ 连接成功", fg=C['success']))
                except Exception as e:
                    cfg.after(0, lambda: status_lbl.config(text=f"❌ 连接失败: {e}", fg=C['danger']))
            self._executor.submit(do_test)

    def save_config():
        api_key_val = api_key_var.get().strip()
        new_config = {
            "api_url": api_url_var.get().strip(),
            "api_key": api_key_val,
            "model": model_var.get().strip(),
            "use_free_model": len(api_key_val) == 0,  # 无API Key时使用免费模型
        }
        if _save_ai_config(new_config):
            self.config = new_config
            remaining = max(0, _FREE_QUOTA - self._free_uses)
            self._status_label.config(text=f"配置已保存  |  免费剩余: {remaining}/{_FREE_QUOTA} 次")
            cfg.destroy()
            if new_config["use_free_model"]:
                self.add_message("ai", f"✅ 已切换到免费模型（剩余 {remaining} 次免费额度）")
            else:
                self.add_message("ai", f"✅ API 配置已更新\n地址: {new_config['api_url']}\n模型: {new_config['model']}")
        else:
            status_lbl.config(text="❌ 保存失败", fg=C['danger'])

        # 底部按钮
        btn_frame = tk.Frame(cfg, bg=C['bg_dark'])
        btn_frame.pack(side=tk.BOTTOM, fill=tk.X, padx=14, pady=10)

        def make_btn(parent, text, cmd, bg_color):
            fg = "#FFFFFF" if bg_color == C['primary'] else C['text']
            btn = tk.Label(parent, text=text, font=("Microsoft YaHei UI", 9),
                           bg=bg_color, fg=fg, cursor="hand2", padx=14, pady=3)
            if text == "保存":
                btn.pack(side=tk.RIGHT, padx=(4, 0))
            elif text == "取消":
                btn.pack(side=tk.RIGHT, padx=4)
            else:
                btn.pack(side=tk.LEFT, padx=(0, 4))
            btn.bind('<Button-1>', lambda e, c=cmd: c())
            btn.bind('<Enter>', lambda e, b=btn, bg=bg_color: b.config(bg=C['primary_hover']))
            btn.bind('<Leave>', lambda e, b=btn, bg=bg_color: b.config(bg=bg_color))
            return btn

        make_btn(btn_frame, "测试连接", test_connection, C['card'])
        make_btn(btn_frame, "取消", cfg.destroy, C['card'])
        make_btn(btn_frame, "保存", save_config, C['primary'])

    def _on_msg_configure(self, event):
        self._msg_canvas.configure(scrollregion=self._msg_canvas.bbox('all'))
        self._msg_canvas.yview_moveto(1.0)

    def _on_input_keyrelease(self, event):
        """Shift+Enter 换行，纯 Enter 发送"""
        if event.keysym == 'Return' and not (event.state & 0x1):
            self._send_input()
            return "break"

    def _on_input_enter(self, event):
        if not event.state & 0x1:
            self._send_input()
            return "break"

    def _send_input(self):
        text = self._input_entry.get('1.0', tk.END).strip()
        if not text:
            return
        self._input_entry.delete('1.0', tk.END)
        self._send_query(text)

    def _send_query(self, text):
        self.add_message("user", text)
        self._status_label.config(text="AI 思考中...")
        threading.Thread(target=self._call_llm_api, args=(text,), daemon=True).start()

    def _call_llm_api(self, user_text):
        try:
            # 检查免费额度
            use_free = self.config.get("use_free_model", True)
            if use_free and self._free_uses >= _FREE_QUOTA:
                self._append_reply(
                    f"⚠️ 免费额度已用完（{_FREE_QUOTA} 次）。\n\n"
                    f"请在「⚙️ 配置」中设置您自己的 API Key 和模型，即可继续使用。")
                return

            api_url = self.config.get("api_url", _FREE_MODEL_CONFIG["api_url"])
            api_key = self.config.get("api_key", _FREE_MODEL_CONFIG["api_key"])
            model = self.config.get("model", _FREE_MODEL_CONFIG["model"])
            # Build messages with configurable model
            msgs = [{"role": m["role"], "content": m["content"]} for m in self._messages[:-1]]
            # If no API key, use the simple chat endpoint
            if not api_key or api_url == "http://localhost:3000/api/chat":
                data = json.dumps({
                    "message": user_text,
                    "messages": msgs,
                    "stream": False
                }).encode('utf-8')
                req = urllib.request.Request(
                    api_url, data=data,
                    headers={'Content-Type': 'application/json'},
                    method='POST')
            else:
                # Use OpenAI-compatible API
                body = {
                    "model": model,
                    "messages": msgs + [{"role": "user", "content": user_text}],
                    "stream": False,
                    "temperature": 0.3
                }
                data = json.dumps(body).encode('utf-8')
                req = urllib.request.Request(
                    api_url, data=data,
                    headers={
                        'Content-Type': 'application/json',
                        'Authorization': f'Bearer {api_key}'
                    },
                    method='POST')
            resp = urllib.request.urlopen(req, timeout=30)
            result = json.loads(resp.read().decode('utf-8'))
            reply = result.get('reply', '') or result.get('content', '') or \
                result.get('choices', [{}])[0].get('message', {}).get('content', '') or str(result)
            search_results = result.get('search_results', [])
            if search_results:
                reply += "\n\n📋 **搜索结果**\n"
                for r in search_results[:5]:
                    code = r.get('code', '')
                    name = r.get('name', '')
                    status = r.get('status', '')
                    reply += f"\n• **{code}** {name} [{status}]"
            # 免费模型扣减额度
            if use_free:
                self._free_uses += 1
                _save_usage(self._free_uses)
                remaining = max(0, _FREE_QUOTA - self._free_uses)
                self.window.after(0, lambda: self._status_label.config(
                    text=f"就绪 ✓  |  免费剩余: {remaining}/{_FREE_QUOTA} 次"))
            self._append_reply(reply)
        except urllib.error.HTTPError as e:
            body = e.read().decode('utf-8', errors='replace')[:200]
            self._append_reply(f"⚠️ 服务器错误 ({e.code}): {body}")
        except urllib.error.URLError:
            self._append_reply("⚠️ 无法连接 AI 服务器。\n请检查网络连接或在「⚙️ 配置」中设置自定义 API。")
        except Exception as e:
            self._append_reply(f"⚠️ 错误: {e}")

    def _append_reply(self, reply):
        if self.window.winfo_exists():
            self.window.after(0, lambda: self._do_append_reply(reply))

    def _do_append_reply(self, reply):
        if not self.window.winfo_exists():
            return
        self.add_message("ai", reply)
        self._status_label.config(text="就绪")

    def add_message(self, role, content, msg_type='text', extra=None):
        """添加一条消息到聊天界面（暗色即时通讯风格）"""
        C = self._C
        self._messages.append({"role": role, "content": content})

        # ── 消息行容器 ──
        outer = tk.Frame(self._msg_inner, bg=C['bg'])
        outer.pack(fill=tk.X, padx=0, pady=0, anchor='e' if role == 'user' else 'w')

        # ── 用户头像+角色标签 ──
        avatar_frame = tk.Frame(outer, bg=C['bg'])
        if role == 'user':
            avatar_frame.pack(fill=tk.X, pady=(6, 0), padx=(0, 12), anchor='e')
            avatar = tk.Label(avatar_frame, text="👤", font=("Microsoft YaHei UI", 12),
                              bg=C['bg'], fg=C['text'])
            avatar.pack(side=tk.RIGHT, padx=(6, 0))
            name_lbl = tk.Label(avatar_frame, text="你", font=C['font_sm'],
                                bg=C['bg'], fg=C['text_muted'])
            name_lbl.pack(side=tk.RIGHT)
        else:
            avatar_frame.pack(fill=tk.X, pady=(6, 0), padx=(12, 0), anchor='w')
            avatar = tk.Label(avatar_frame, text="🤖", font=("Microsoft YaHei UI", 12),
                              bg=C['bg'], fg=C['text'])
            avatar.pack(side=tk.LEFT, padx=(0, 6))
            name_lbl = tk.Label(avatar_frame, text="AI", font=C['font_sm'],
                                bg=C['bg'], fg=C['text_muted'])
            name_lbl.pack(side=tk.LEFT)

        # ── 气泡容器 ──
        bubble_frame = tk.Frame(outer, bg=C['bg'])
        bubble_frame.pack(fill=tk.X, padx=(12 if role == 'user' else 12, 12 if role == 'user' else 12),
                          pady=(1, 0), anchor='e' if role == 'user' else 'w')

        if role == 'user':
            # 用户：蓝色气泡，右对齐
            bubble = tk.Frame(bubble_frame, bg=C['primary'], bd=0)
            bubble.pack(fill=tk.X, anchor='e')
            bubble.pack_propagate(False)
            content_frame = tk.Frame(bubble, bg=C['primary'])
            content_frame.pack(fill=tk.X, padx=10, pady=6)
            self._render_message_content(content_frame, content, msg_type, extra, role)
            ts = tk.Label(bubble, text=datetime.now().strftime("%H:%M"),
                          font=("Microsoft YaHei UI", 7), bg=C['primary'], fg="#93C5FD")
            ts.pack(anchor='e', padx=10, pady=(0, 4))
        else:
            # AI：深色气泡，左对齐
            bubble = tk.Frame(bubble_frame, bg=C['card'], bd=0)
            bubble.pack(fill=tk.X, anchor='w')
            bubble.pack_propagate(False)
            content_frame = tk.Frame(bubble, bg=C['card'])
            content_frame.pack(fill=tk.X, padx=10, pady=6)
            self._render_message_content(content_frame, content, msg_type, extra, role)
            ts = tk.Label(bubble, text=datetime.now().strftime("%H:%M"),
                          font=("Microsoft YaHei UI", 7), bg=C['card'], fg=C['text_muted'])
            ts.pack(anchor='w', padx=10, pady=(0, 4))

        # ── 操作栏（悬停显示） ──
        self._add_action_bar(outer, content, role)

        self._msg_canvas.after(50, self._on_msg_configure)


    def _render_message_content(self, parent, content, msg_type, extra, role):
        """在气泡内渲染消息内容"""
        C = self._C
        if msg_type == 'text':
            self._add_text_bubble(parent, content, role)
        elif msg_type == 'table':
            self._add_table_widget(parent, extra or [])
            if content:
                self._add_text_bubble(parent, content, role, small=True)
        elif msg_type == 'image':
            self._add_image_widget(parent, extra)
            if content:
                self._add_text_bubble(parent, content, role, small=True)
        elif msg_type == 'file':
            self._add_file_widget(parent, extra or {})
            if content:
                self._add_text_bubble(parent, content, role, small=True)
        else:
            self._add_text_bubble(parent, content, role)

    def _add_action_bar(self, parent, content, role):
        """每个气泡底部的操作图标栏"""
        C = self._C
        bar = tk.Frame(parent, bg=C['bg'], height=22)
        bar.pack(fill=tk.X, pady=(0, 0), anchor='e' if role == 'user' else 'w')
        bar.pack_propagate(False)

        # 固定内容用于闭包
        msg_content = content

        def do_copy():
            self.window.clipboard_clear()
            self.window.clipboard_append(msg_content)
            self._status_label.config(text="📋 已复制")

        def do_export_single_docx():
            """导出单条消息为 Word"""
            if not HAS_DOCX:
                messagebox.showwarning("提示", "需要安装 python-docx", parent=self.window)
                return
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            try:
                role_label = "AI" if role == "ai" else "用户"
                path = filedialog.asksaveasfilename(
                    title=f"保存 {role_label} 消息",
                    defaultextension=".docx",
                    filetypes=[("Word 文档", "*.docx"), ("所有文件", "*.*")])
                if not path:
                    return
                doc = Document()
                p = doc.add_paragraph()
                run = p.add_run(f"[{role_label}]")
                run.bold = True
                run.font.size = Pt(11)
                run.font.name = "SimSun"
                if role == "ai":
                    run.font.color.rgb = RGBColor(0x1a, 0x73, 0xe8)
                p2 = doc.add_paragraph(msg_content)
                for run in p2.runs:
                    run.font.name = "SimSun"
                    run.font.size = Pt(11)
                doc.save(path)
                self._status_label.config(text="✅ 已导出 Word")
            except Exception as e:
                messagebox.showerror("错误", f"导出失败: {e}", parent=self.window)

        def do_export_single_pdf():
            """导出单条消息为 PDF"""
            try:
                from reportlab.lib.pagesizes import A4
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.units import cm
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                from reportlab.lib import colors
                from reportlab.pdfbase import pdfmetrics
                from reportlab.pdfbase.cidfonts import UnicodeCIDFont
            except Exception:
                messagebox.showwarning("提示", "需要安装 reportlab", parent=self.window)
                return
            try:
                pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
                role_label = "AI" if role == "ai" else "用户"
                path = filedialog.asksaveasfilename(
                    title=f"保存 {role_label} 消息",
                    defaultextension=".pdf",
                    filetypes=[("PDF 文档", "*.pdf"), ("所有文件", "*.*")])
                if not path:
                    return
                doc = SimpleDocTemplate(path, pagesize=A4,
                                        leftMargin=2*cm, rightMargin=2*cm,
                                        topMargin=2*cm, bottomMargin=2*cm)
                styles = getSampleStyleSheet()
                role_style = ParagraphStyle('Role', parent=styles['Normal'],
                                            fontName='STSong-Light', fontSize=12,
                                            textColor=colors.HexColor('#1a73e8' if role == 'ai' else '#333'),
                                            spaceAfter=8)
                content_style = ParagraphStyle('Content', parent=styles['Normal'],
                                               fontName='STSong-Light', fontSize=10,
                                               spaceAfter=12)
                c = msg_content.replace("\n", "<br/>").replace("**", "").replace("•", "- ")
                story = [Paragraph(f"<b>{role_label}：</b>", role_style),
                         Paragraph(c, content_style)]
                doc.build(story)
                self._status_label.config(text="✅ 已导出 PDF")
            except Exception as e:
                messagebox.showerror("错误", f"导出 PDF 失败: {e}", parent=self.window)

# 操作按钮
        for txt, tooltip, cmd in [("📋", "复制", do_copy),
                                   ("📄", "导出 Word", do_export_single_docx),
                                   ("📕", "导出 PDF", do_export_single_pdf)]:
            btn = tk.Label(bar, text=txt, font=C['font_sm'],
                           bg=C['bg'], fg=C['text_muted'], cursor="hand2", padx=3)
            btn.pack(side=tk.LEFT, padx=1)
            btn.bind('<Button-1>', lambda e, c=cmd: c())
            btn.bind('<Enter>', lambda e, b=btn: b.config(fg=C['text']))
            btn.bind('<Leave>', lambda e, b=btn: b.config(fg=C['text_muted']))

    def _add_text_bubble(self, parent, text, role, small=False):
        """富文本气泡，带圆角风格"""
        C = self._C
        font_size = 9 if small else 10
        if role == 'user':
            bg = C['primary']
            fg = "#FFFFFF"
            max_width = 60
        else:
            bg = C['card']
            fg = C['text']
            max_width = 60

        # 计算合适高度
        line_count = text.count('\n') + 1
        est_height = min(18, max(2, line_count))

        bubble = tk.Text(parent, wrap=tk.WORD, font=("Microsoft YaHei", font_size),
                         bg=bg, fg=fg, bd=0, padx=10, pady=8,
                         height=est_height,
                         width=max_width, highlightthickness=0,
                         relief="flat")

        # 解析富文本
        self._insert_formatted_text(bubble, text, fg, role == 'user')

        bubble.config(state=tk.DISABLED)
        bubble.pack(fill=tk.X, pady=(0, 0), anchor='e' if role == 'user' else 'w')
        return bubble

    def _add_plain_text(self, parent, text, small=False):
        """AI/系统消息：纯文本左对齐，无气泡无操作栏"""
        C = self._C
        font_size = 9 if small else 10
        label = tk.Label(parent, text=text, font=("Microsoft YaHei UI", font_size),
                         bg=C['bg'], fg=C['text'],
                         wraplength=380, anchor='w', justify=tk.LEFT)
        label.pack(fill=tk.X, pady=(1, 1), anchor='w')
        return label

    def _insert_formatted_text(self, widget, text, default_fg, is_user):
        """完整 Markdown 解析引擎，插入到 Text 组件中（暗色主题）"""
        import re
        C = self._C
        # ── 注册 tag ──
        widget.tag_config('bold', font=("Microsoft YaHei UI", 10, "bold"), foreground=default_fg)
        widget.tag_config('italic', font=("Microsoft YaHei UI", 10, "italic"), foreground=default_fg)
        widget.tag_config('h1', font=("Microsoft YaHei UI", 14, "bold"), foreground=C['primary'], spacing1=6, spacing3=4)
        widget.tag_config('h2', font=("Microsoft YaHei UI", 12, "bold"), foreground=C['text'], spacing1=4, spacing3=2)
        widget.tag_config('h3', font=("Microsoft YaHei UI", 11, "bold"), foreground=C['text_muted'], spacing1=3, spacing3=2)
        code_bg = "#1E293B" if not is_user else "#1E3A5F"
        widget.tag_config('code', font=("Consolas", 9), foreground="#F87171" if not is_user else "#FCA5A5",
                         background=code_bg)
        widget.tag_config('code_block', font=("Consolas", 9), foreground=C['text'],
                         background=C['bg_dark'], lmargin1=12, lmargin2=12, rmargin=12,
                         spacing1=4, spacing3=4)
        widget.tag_config('small', font=("Microsoft YaHei UI", 8), foreground=default_fg)
        widget.tag_config('red', foreground=C['danger'])
        widget.tag_config('green', foreground=C['success'])
        widget.tag_config('blue', foreground=C['primary'])
        widget.tag_config('quote', font=("Microsoft YaHei UI", 9, "italic"), foreground=C['text_muted'],
                         background=C['bg_dark'], lmargin1=16, lmargin2=16, spacing1=2, spacing3=2)
        widget.tag_config('link', foreground=C['primary'], underline=True)
        widget.tag_config('bullet', lmargin1=12, lmargin2=24)
        widget.tag_config('ordered', lmargin1=12, lmargin2=24)
        widget.tag_config('hr', foreground=C['text_muted'], font=("Microsoft YaHei UI", 6))

        lines = text.split('\n')
        code_block = False
        code_buffer = []
        ordered_counter = 0

        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # ── 代码块 ──
            if stripped.startswith('```'):
                if code_block:
                    # 结束代码块
                    code_text = '\n'.join(code_buffer)
                    widget.insert(tk.END, code_text + '\n', 'code_block')
                    code_buffer = []
                    code_block = False
                else:
                    code_block = True
                    code_buffer = []
                i += 1
                continue
            if code_block:
                code_buffer.append(line)
                i += 1
                continue

            # ── 空行 ──
            if not stripped:
                widget.insert(tk.END, '\n')
                i += 1
                continue

            # ── 分割线 ──
            if re.match(r'^-{3,}$', stripped) or re.match(r'^\*{3,}$', stripped):
                widget.insert(tk.END, '─' * 40 + '\n', 'hr')
                i += 1
                continue

            # ── 标题 ──
            if stripped.startswith('# '):
                widget.insert(tk.END, stripped[2:] + '\n', 'h1')
                i += 1
                continue
            if stripped.startswith('## '):
                widget.insert(tk.END, stripped[3:] + '\n', 'h2')
                i += 1
                continue
            if stripped.startswith('### '):
                widget.insert(tk.END, stripped[4:] + '\n', 'h3')
                i += 1
                continue

            # ── 引用 ──
            if stripped.startswith('> '):
                widget.insert(tk.END, stripped[2:] + '\n', 'quote')
                i += 1
                continue

            # ── 有序列表 ──
            ordered_match = re.match(r'^(\d+)\.\s+(.*)', stripped)
            if ordered_match:
                num = ordered_match.group(1)
                content = ordered_match.group(2)
                self._insert_inline_format(widget, f"{num}. {content}", default_fg, is_user, tag='ordered')
                widget.insert(tk.END, '\n')
                i += 1
                continue

            # ── 无序列表 ──
            if stripped.startswith('- ') or stripped.startswith('• ') or stripped.startswith('* '):
                content = stripped[2:]
                widget.insert(tk.END, '•  ', 'small')
                self._insert_inline_format(widget, content, default_fg, is_user, tag='bullet')
                widget.insert(tk.END, '\n')
                i += 1
                continue

            # ── 普通段落 ──
            self._insert_inline_format(widget, line, default_fg, is_user)
            widget.insert(tk.END, '\n')
            i += 1

        if code_buffer:
            widget.insert(tk.END, '\n'.join(code_buffer) + '\n', 'code_block')

    def _insert_inline_format(self, widget, text, default_fg, is_user, tag=None):
        """行内格式解析：粗体、斜体、代码、链接、图片、内联表格、颜色标记"""
        import re
        # 合并多行模式匹配：按优先级匹配
        pattern = r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\!\[.*?\]\(.*?\)|\[.*?\]\(.*?\)|\|.*?\|)'
        parts = re.split(pattern, text)

        # 特殊颜色标记
        color_map = {'⚠️': 'red', '❌': 'red', '✅': 'green', '📄': 'blue',
                     '📊': 'blue', '🔍': 'blue', '🏗️': 'blue', '🔥': 'blue',
                     '📋': 'blue', '💡': 'blue'}

        for part in parts:
            if not part:
                continue

            # 粗体 **text**
            if part.startswith('**') and part.endswith('**'):
                widget.insert(tk.END, part[2:-2], 'bold')

            # 斜体 *text*
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                widget.insert(tk.END, part[1:-1], 'italic')

            # 行内代码 `text`
            elif part.startswith('`') and part.endswith('`'):
                widget.insert(tk.END, part[1:-1], 'code')

            # 图片 ![alt](url)
            elif part.startswith('![') and part.endswith(')'):
                self._insert_inline_image(widget, part)

            # 链接 [text](url)
            elif part.startswith('[') and part.endswith(')'):
                m = re.match(r'\[(.*?)\]\((.*?)\)', part)
                if m:
                    link_text = m.group(1)
                    link_url = m.group(2)
                    # 用蓝色显示链接文本
                    widget.insert(tk.END, link_text, 'link')
                    # 保存链接信息（点击事件）
                    if hasattr(self, '_links'):
                        self._links.append((widget, link_text, link_url))

            # 内联表格 |col1|col2|
            elif part.startswith('|') and part.endswith('|'):
                cells = [c.strip() for c in part.split('|') if c.strip()]
                display = ' | '.join(cells)
                widget.insert(tk.END, f' {display} ', 'code')

            # 颜色标记
            else:
                # 逐字符检查颜色标记
                remaining = part
                while remaining:
                    found = False
                    for sym, tag_name in color_map.items():
                        if remaining.startswith(sym):
                            widget.insert(tk.END, sym, tag_name)
                            remaining = remaining[len(sym):]
                            found = True
                            break
                    if not found:
                        # 普通文本，按段插入
                        widget.insert(tk.END, remaining)
                        break

    def _insert_inline_image(self, widget, markdown):
        """在 Text 中插入内嵌图片（缩略图）"""
        import re
        m = re.match(r'!\[(.*?)\]\((.*?)\)', markdown)
        if not m:
            return
        alt = m.group(1)
        path = m.group(2)
        if not HAS_PIL:
            widget.insert(tk.END, f'[图片: {alt}]', 'small')
            return
        try:
            if os.path.exists(path):
                img = Image.open(path)
                img.thumbnail((120, 90), Image.Resampling.LANCZOS)
                # 用 PhotoImage 嵌入
                photo = ImageTk.PhotoImage(img)
                widget.image_create(tk.END, image=photo)
                # 保持引用避免 GC
                if not hasattr(self, '_inline_images'):
                    self._inline_images = []
                self._inline_images.append(photo)
            else:
                widget.insert(tk.END, f'[图片: {alt}]', 'small')
        except Exception:
            widget.insert(tk.END, f'[图片: {alt}]', 'small')

    def _add_table_widget(self, parent, table_data):
        """在聊天中显示表格（暗色主题）
        table_data: list of (header_list, rows_list)
        """
        C = self._C
        if not table_data:
            return
        headers = table_data[0] if len(table_data) > 0 else []
        rows = table_data[1] if len(table_data) > 1 else []

        frame = tk.Frame(parent, bg=C['bg_dark'])
        frame.pack(fill=tk.X, pady=4)

        # Header
        header_frame = tk.Frame(frame, bg=C['bg_dark'])
        header_frame.pack(side=tk.TOP, fill=tk.X)
        for i, h in enumerate(headers):
            lbl = tk.Label(header_frame, text=str(h), font=("Microsoft YaHei UI", 9, "bold"),
                           bg=C['primary'], fg="#FFFFFF", padx=8, pady=3, relief="flat")
            lbl.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # Rows
        for row in rows[:20]:  # max 20 rows
            row_frame = tk.Frame(frame, bg=C['bg'])
            row_frame.pack(side=tk.TOP, fill=tk.X)
            for i, cell in enumerate(row):
                if i >= len(headers):
                    break
                fg_color = C['text']
                if '废止' in str(cell) or '作废' in str(cell) or '❌' in str(cell):
                    fg_color = C['danger']
                elif '现行' in str(cell) or '✅' in str(cell):
                    fg_color = C['success']
                lbl = tk.Label(row_frame, text=str(cell), font=("Microsoft YaHei UI", 8),
                               bg=C['bg'], fg=fg_color, padx=8, pady=2, relief="flat")
                lbl.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        if len(rows) > 20:
            more = tk.Label(frame, text=f"... 还有 {len(rows) - 20} 行", font=("Microsoft YaHei UI", 8), fg=C['text_muted'])
            more.pack(anchor=tk.W, padx=8, pady=2)

    def _add_image_widget(self, parent, img_path_or_photo):
        """在聊天中显示图片"""
        C = self._C
        if not img_path_or_photo:
            return
        frame = tk.Frame(parent, bg=C['bg_dark'])
        frame.pack(fill=tk.X, pady=4)
        try:
            if HAS_PIL:
                if isinstance(img_path_or_photo, str):
                    img = Image.open(img_path_or_photo)
                    img.thumbnail((200, 150), Image.Resampling.LANCZOS)
                    photo = ImageTk.PhotoImage(img)
                else:
                    photo = img_path_or_photo
                lbl = tk.Label(frame, image=photo, bg=C['card'], relief="flat", bd=0)
                lbl.image = photo  # keep reference
                lbl.pack()
            else:
                tk.Label(frame, text="[图片: PIL 未安装]",
                         font=("Microsoft YaHei UI", 8),
                         bg=C['bg_dark'], fg=C['text_muted']).pack()
        except Exception as e:
            tk.Label(frame, text=f"[图片加载失败: {e}]",
                     font=("Microsoft YaHei UI", 8),
                     bg=C['bg_dark'], fg=C['danger']).pack()

    def _add_file_widget(self, parent, file_info):
        """在聊天中显示文件附件（暗色主题）"""
        C = self._C
        if not file_info:
            return
        frame = tk.Frame(parent, relief="flat", bd=0, bg=C['bg_dark'])
        frame.pack(fill=tk.X, pady=4, padx=2)

        name = file_info.get('name', 'unknown')
        size = file_info.get('size', '')
        fpath = file_info.get('path', '')

        icon_lbl = tk.Label(frame, text="📎", font=("Microsoft YaHei UI", 14),
                            bg=C['bg_dark'], fg=C['text'])
        icon_lbl.pack(side=tk.LEFT, padx=8, pady=6)

        info_frame = tk.Frame(frame, bg=C['bg_dark'])
        info_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=6, pady=4)
        tk.Label(info_frame, text=name, font=("Microsoft YaHei UI", 9, "bold"),
                 bg=C['bg_dark'], fg=C['text'], anchor=tk.W).pack(fill=tk.X)
        if size:
            tk.Label(info_frame, text=size, font=("Microsoft YaHei UI", 8),
                     bg=C['bg_dark'], fg=C['text_muted'], anchor=tk.W).pack(fill=tk.X)

        def open_file():
            import subprocess
            try:
                if os.path.exists(fpath):
                    subprocess.run(['explorer', '/select,', fpath])
                else:
                    messagebox.showinfo("文件", f"文件路径:\n{fpath}\n\n(文件可能已被清理)", parent=self.window)
            except Exception as e:
                messagebox.showerror("错误", f"无法打开文件: {e}", parent=self.window)

        open_btn = tk.Label(frame, text="📂 打开", font=("Microsoft YaHei UI", 8),
                            bg=C['primary'], fg="#FFFFFF", cursor="hand2", padx=6, pady=2)
        open_btn.pack(side=tk.RIGHT, padx=8, pady=6)
        open_btn.bind('<Button-1>', lambda e: open_file())

    def set_ocr_results(self, results):
                """从外部设置 OCR 识别结果，以表格形式显示"""
                if not results:
                    return
                headers = ['规范编号', '名称', '来源', '状态']
                rows = []
                for item in results:
                    code = item.get('code', item.get('original', ''))
                    name = item.get('name', '')[:25]
                    source = item.get('source', '')[:20]
                    status = '已查'
                    if item.get('found') is True:
                        status = '✅ 现行' if '废止' not in item.get('status', '') else '⚠️ 废止'
                    elif item.get('found') is False:
                        status = '❌ 未查到'
                    rows.append([code, name or '-', source or '-', status])
                self.add_message("ai", "📄 识别到的规范结果：", msg_type='table', extra=[headers, rows])
                self.show()

    def send_standard_check(self, codes_info):
            """发送规范检查结果到 AI 进行分析"""
            if not codes_info:
                return
            headers = ['规范编号', '状态', '建议']
            rows = []
            for code, result in codes_info:
                status = result.get('status', '未知')
                found = result.get('found', False)
                if found:
                    if '废止' in status or '作废' in status:
                        replacement = result.get('replacement_raw', '')
                        action = f'需替换 → {replacement[:20]}'
                    else:
                        action = '✅ 现行，可用'
                    if result.get('dual_match'):
                        action += ' (双重确认)'
                else:
                    action = '❌ 未查询到，建议人工核实'
                rows.append([code, status, action])
            self.add_message("ai", "📊 规范检查结果分析：", msg_type='table', extra=[headers, rows])
            self.show()


# ===== 主应用 =====
class App:
    def __init__(self):
        self.checker = None
        self._data_loaded = False
        self.ai_config = _load_ai_config()
        self.pdf_paths = []
        self.current_path = None
        self.file_type = None
        self.pdf_images = []
        self._fitz_doc = None  # fitz doc for lazy render
        self._total_pages = 0  # PDF/OFD total pages
        self.ocr_results = []
        self.extracted_codes = []
        self.extracted_code_info = {}
        self.code_locations = []
        self.check_results = []
        self.ocr_region = None
        self.selector = None
        self.selection_mode = False
        self.current_display_index = 0
        self._pan_start_x = 0
        self._pan_start_y = 0
        self._pan_image_x = 0
        self._pan_image_y = 0
        self._panning = False
        self._preview_name_text_id = None
        self._file_queue = []
        self._batch_running = False
        self._batch_abort = False
        self._thumbnail_images = []
        self._rotation_angle = 0
        self._render_dpi = self._get_render_dpi() # 动态渲染 DPI
        self.root = tk.Tk()
        self._name_index = {}
        self.root.title(APP_TITLE)
        self.root.geometry("1360x840")
        self.root.minsize(800, 600)
        # 设置窗口图标（app_icon.ico）
        self._set_app_icon()
        self._fit_mode = tk.StringVar(value='fit_page')
        self.ai_chat = None
        self._pending_ocr_results = None  # OCR结果缓存（AI窗口未打开时）
        self._setup_style()
        self.setup_ui()
        # 启动时检查AI配置
        self._check_ai_config()
        # 异步加载标准数据库（不阻塞UI启动）
        self._load_data_async()

    def _load_data_async(self):
        """异步加载标准数据库，不阻塞 UI 启动"""
        self.status_var.set("正在加载标准数据库...")
        self.root.update_idletasks()
        def _do_load():
            try:
                self.checker = StandardChecker()
                self._data_loaded = True
                self.status_var.set("就绪")
            except Exception as e:
                self.status_var.set(f"数据库加载失败: {e}")
                print(f"StandardChecker 加载失败: {e}")
                self._data_loaded = False
        self.root.after(100, _do_load)


    def _set_app_icon(self):
        """设置窗口图标（主窗口 + 所有 Toplevel 子窗口）"""
        # 兼容 PyInstaller onedir 打包后的路径
        icon_path = _APP_DIR / 'app_icon.ico'
        png_path = _APP_DIR / 'app_icon.png'
        if not icon_path.exists() and not png_path.exists():
            # 退回到脚本所在目录
            icon_path = Path(__file__).parent / 'app_icon.ico'
            png_path = Path(__file__).parent / 'app_icon.png'

        if icon_path.exists():
            try:
                self.root.iconbitmap(str(icon_path))
                self._icon_path = str(icon_path)
                print(f'Icon set: {icon_path.name}')
                return
            except Exception as e:
                print(f'Failed iconbitmap from {icon_path}: {e}')

        if png_path.exists():
            try:
                png_ico = tk.PhotoImage(file=str(png_path))
                self.root.iconphoto(True, png_ico)
                self._icon_image = png_ico
                print(f'Icon set: {png_path.name}')
                return
            except Exception as e:
                print(f'Failed iconphoto from {png_path}: {e}')

        print('No icon file found, using default')

    def _set_icon_for_toplevel(self, win):
        """给任意 Toplevel 窗口设置图标"""
        if hasattr(self, '_icon_path') and self._icon_path:
            try:
                win.iconbitmap(self._icon_path)
                return
            except Exception:
                pass
        if hasattr(self, '_icon_image'):
            try:
                win.iconphoto(True, self._icon_image)
                return
            except Exception:
                pass


    def _setup_style(self):
        """配置全局暗色 ttk 主题"""
        style = ttk.Style()
        style.theme_use("default")
        # 暗色调色板
        BG = "#1E293B"
        BG_DARK = "#0F172A"
        CARD = "#334155"
        TEXT = "#E2E8F0"
        TEXT_MUTED = "#94A3B8"
        PRIMARY = "#3B82F6"
        PRIMARY_HOVER = "#2563EB"
        SELECT = "#1E40AF"
        SUCCESS = "#22C55E"
        WARNING = "#F59E0B"
        DANGER = "#EF4444"
        default_font = ("Microsoft YaHei UI", 10)
        bold_font = ("Microsoft YaHei UI", 10, "bold")
        # 全局背景
        style.configure(".", font=default_font, background=BG, foreground=TEXT)
        # TLabel
        style.configure("TLabel", font=default_font, background=BG, foreground=TEXT, padding=4)
        # 标题
        style.configure("Header.TLabel", font=("Microsoft YaHei UI", 11, "bold"), background=BG, foreground=TEXT)
        style.configure("Title.TLabel", font=("Microsoft YaHei UI", 16, "bold"), foreground=TEXT, background=BG_DARK)
        # TButton
        style.configure("TButton", font=default_font, background=CARD, foreground=TEXT, padding=6, borderwidth=0)
        style.map("TButton", background=[("active", PRIMARY_HOVER), ("pressed", SELECT)], foreground=[("active", "#FFFFFF")])
        style.configure("Primary.TButton", font=bold_font, background=PRIMARY, foreground="#FFFFFF")
        style.map("Primary.TButton", background=[("active", PRIMARY_HOVER), ("pressed", SELECT)])
        style.configure("Action.TButton", padding=8, background=CARD, foreground=TEXT, borderwidth=0)
        style.map("Action.TButton", background=[("active", PRIMARY), ("pressed", SELECT)])
        style.configure("Success.TButton", background=SUCCESS, foreground="#FFFFFF")
        style.map("Success.TButton", background=[("active", "#16A34A")])
        # TFrame
        style.configure("TFrame", background=BG)
        style.configure("Card.TFrame", background=CARD)
        style.configure("Dark.TFrame", background=BG_DARK)
        # TEntry
        style.configure("TEntry", fieldbackground=CARD, foreground=TEXT, insertcolor=TEXT, borderwidth=0, padding=4)
        style.map("TEntry", fieldbackground=[("focus", BG_DARK)])
        # TCombobox
        style.configure("TCombobox", fieldbackground=CARD, foreground=TEXT, arrowcolor=TEXT, padding=4)
        style.map("TCombobox", fieldbackground=[("readonly", CARD), ("focus", BG_DARK)])
        # Treeview
        style.configure("Treeview", rowheight=26, font=default_font, background=CARD, foreground=TEXT, fieldbackground=CARD, borderwidth=0)
        style.map("Treeview", background=[("selected", SELECT)], foreground=[("selected", "#FFFFFF")])
        style.configure("Treeview.Heading", font=bold_font, background=BG_DARK, foreground=TEXT, borderwidth=0)
        # TScrollbar
        style.configure("Vertical.TScrollbar", background=CARD, troughcolor=BG_DARK, borderwidth=0, width=10)
        style.map("Vertical.TScrollbar", background=[("active", PRIMARY)])
        style.configure("Horizontal.TScrollbar", background=CARD, troughcolor=BG_DARK, borderwidth=0, height=10)
        style.map("Horizontal.TScrollbar", background=[("active", PRIMARY)])
        # TLabelframe
        style.configure("TLabelframe", background=BG, foreground=TEXT, borderwidth=0)
        style.configure("TLabelframe.Label", font=default_font, background=BG, foreground=TEXT)
        # TNotebook
        style.configure("TNotebook", background=BG_DARK, borderwidth=0)
        style.configure("TNotebook.Tab", background=CARD, foreground=TEXT, padding=[8, 4])
        style.map("TNotebook.Tab", background=[("selected", BG), ("active", PRIMARY)])
        # TSizegrip
        style.configure("TSizegrip", background=BG_DARK)
        # 状态栏
        style.configure("Status.TLabel", background=BG_DARK, foreground=TEXT_MUTED, anchor="w", padding=6)

    def setup_ui(self):
        """紧凑专业三栏布局"""
        C = {
            'bg': "#1E293B", 'bg_dark': "#0F172A", 'card': "#334155",
            'text': "#E2E8F0", 'text_muted': "#94A3B8",
            'primary': "#3B82F6", 'primary_hover': "#2563EB",
            'select': "#1E40AF", 'success': "#22C55E", 'danger': "#EF4444",
            'border': "#2D3A4A",
        }

        # ── 顶部工具栏 ──
        topbar = tk.Frame(self.root, bg=C['bg_dark'], height=36)
        topbar.pack(side=tk.TOP, fill=tk.X)
        topbar.pack_propagate(False)

        def _tb(text, cmd, bold=False, color=None):
            bg = color or C['bg_dark']
            fg = "#FFFFFF" if color else C['text']
            lbl = tk.Label(topbar, text=text, font=("Microsoft YaHei UI", 9, "bold" if bold else "normal"),
                           bg=bg, fg=fg, cursor="hand2", padx=6, pady=2)
            lbl.pack(side=tk.LEFT, padx=1)
            lbl.bind('<Button-1>', lambda e: cmd())
            if not color:
                lbl.bind('<Enter>', lambda e: lbl.config(bg=C['card']))
                lbl.bind('<Leave>', lambda e: lbl.config(bg=C['bg_dark']))
            return lbl

        def _sep():
            tk.Label(topbar, text="│", bg=C['bg_dark'], fg=C['text_muted'], padx=2).pack(side=tk.LEFT)

        _tb("📂 打开", self.open_file)
        _tb("📁 文件夹", self.open_folder)
        _sep()
        _tb("◀", self._prev_page)
        _tb("▶", self._next_page)
        _tb("🔍+", self._zoom_in)
        _tb("🔍-", self._zoom_out)
        _tb("↻", self._rotate_cw)
        _tb("↺", self._rotate_ccw)
        _sep()
        _tb("⬜ 选择", self.start_selection)
        _tb("❌ 清除", self.clear_region)
        _sep()
        _tb("🔍 OCR", self.start_ocr, bold=True, color=C['primary'])
        _tb("✅ 检查", self.check_standards, bold=True, color=C['success'])
        _tb("📋 批量", self.batch_process_all)
        _sep()
        _tb("📄 导出", self.export_doc)
        _tb("📊 Excel", self.export_excel)
        _sep()
        _tb("🤖 AI", self._toggle_ai_chat, bold=True, color=C['primary'])

        # 右侧信息
        self.page_var = tk.StringVar(value="第 0 / 0 页")
        tk.Label(topbar, textvariable=self.page_var, font=("Microsoft YaHei UI", 9),
                 bg=C['bg_dark'], fg=C['text_muted']).pack(side=tk.RIGHT, padx=8)
        self._preview_name_var = tk.StringVar(value="")
        tk.Label(topbar, textvariable=self._preview_name_var, font=("Microsoft YaHei UI", 9, "bold"),
                 bg=C['bg_dark'], fg=C['danger']).pack(side=tk.RIGHT, padx=8)

        # 适应模式
        frm = tk.Frame(topbar, bg=C['bg_dark'])
        frm.pack(side=tk.RIGHT, padx=2)
        for txt, val in [("适应页面", 'fit_page'), ("适应宽度", 'fit_width')]:
            lbl = tk.Label(frm, text=txt, font=("Microsoft YaHei UI", 8),
                           bg=C['primary'] if self._fit_mode.get() == val else C['card'],
                           fg=C['text'], cursor="hand2", padx=4, pady=1)
            lbl.pack(side=tk.LEFT, padx=1)
            lbl.bind('<Button-1>', lambda e, v=val: (self._fit_mode.set(v), self._redraw_current_page()))
            lbl.bind('<Enter>', lambda e, lb=lbl: lb.config(bg=C['primary_hover']))
            lbl.bind('<Leave>', lambda e, lb=lbl, v=val: lb.config(bg=C['primary'] if self._fit_mode.get() == v else C['card']))

        # ── 三栏主体 ──
        main = tk.Frame(self.root, bg=C['bg'])
        main.pack(side=tk.TOP, fill=tk.BOTH, expand=True)

        # 左侧栏：文件列表+缩略图
        left = tk.Frame(main, bg=C['bg_dark'], width=160)
        left.pack(side=tk.LEFT, fill=tk.Y)
        left.pack_propagate(False)
        tk.Label(left, text="📄 文件列表", font=("Microsoft YaHei UI", 9, "bold"),
                 bg=C['bg_dark'], fg=C['text']).pack(fill=tk.X, padx=6, pady=(4, 2))
        self.queue_listbox = tk.Listbox(left, height=6,
                                         font=("Microsoft YaHei UI", 8), exportselection=False,
                                         bg=C['card'], fg=C['text'],
                                         selectbackground=C['select'], selectforeground="#FFFFFF",
                                         borderwidth=0, highlightthickness=0)
        self.queue_listbox.pack(fill=tk.X, padx=4, pady=2)
        self.queue_listbox.bind('<<ListboxSelect>>', self._on_queue_select)
        self.queue_count_label = tk.Label(left, text="0 个文件", font=("Microsoft YaHei UI", 8),
                                          bg=C['bg_dark'], fg=C['text_muted'])
        self.queue_count_label.pack(anchor=tk.W, padx=6)

        # 缩略图
        self.thumb_canvas = tk.Canvas(left, bg=C['bg_dark'], highlightthickness=0, width=150, height=200)
        self.thumb_canvas.pack(fill=tk.BOTH, expand=True, padx=2, pady=2)
        thumb_scroll = ttk.Scrollbar(left, orient=tk.VERTICAL, command=self.thumb_canvas.yview)
        thumb_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.thumb_canvas.configure(yscrollcommand=thumb_scroll.set)
        self.thumb_frame = tk.Frame(self.thumb_canvas, bg=C['bg_dark'])
        self.thumb_scroll_window = self.thumb_canvas.create_window(
            (0, 0), window=self.thumb_frame, anchor='nw', width=150)
        self.thumb_frame.bind('<Configure>', lambda e: self.thumb_canvas.configure(
            scrollregion=self.thumb_canvas.bbox('all')))

        # 中间：预览区
        center = tk.Frame(main, bg=C['bg'])
        center.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.pdf_canvas = tk.Canvas(center, bg=C['bg'], highlightthickness=0)
        self.pdf_canvas.pack(fill=tk.BOTH, expand=True)
        self.pdf_canvas.bind('<Configure>', self._on_canvas_resize)
        self.pdf_canvas.bind('<MouseWheel>', self._on_mouse_wheel)
        self.pdf_canvas.bind('<ButtonPress-2>', self._on_pan_start)
        self.pdf_canvas.bind('<B2-Motion>', self._on_pan_drag)
        self.pdf_canvas.bind('<ButtonRelease-2>', self._on_pan_end)
        self._resize_after_id = None
        self.selector = RegionSelector(self.pdf_canvas, None, self._on_region_selected)

        # 右侧：结果面板
        right = tk.Frame(main, bg=C['bg_dark'], width=320)
        right.pack(side=tk.RIGHT, fill=tk.Y)
        right.pack_propagate(False)
        nb = ttk.Notebook(right)
        nb.pack(fill=tk.BOTH, expand=True)
        self.notebook = nb

        # OCR 文本标签页
        ocr_frame = tk.Frame(nb, bg=C['bg'])
        nb.add(ocr_frame, text="OCR 文本")
        self.ocr_text = tk.Text(ocr_frame, wrap=tk.WORD,
                                font=("Microsoft YaHei UI", 10),
                                bg=C['card'], fg=C['text'],
                                insertbackground=C['text'],
                                borderwidth=0, highlightthickness=0, padx=6, pady=4)
        ocr_scroll = ttk.Scrollbar(ocr_frame, orient=tk.VERTICAL, command=self.ocr_text.yview)
        self.ocr_text.configure(yscrollcommand=ocr_scroll.set)
        ocr_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.ocr_text.pack(fill=tk.BOTH, expand=True)

        # 规范列表标签页
        list_frame = tk.Frame(nb, bg=C['bg'])
        nb.add(list_frame, text="规范列表")
        list_columns = ('no', 'code', 'name', 'source')
        self.list_tree = ttk.Treeview(list_frame, columns=list_columns, show='headings', selectmode='extended')
        self.list_tree.heading('no', text='#')
        self.list_tree.heading('code', text='规范编号')
        self.list_tree.heading('name', text='名称')
        self.list_tree.heading('source', text='来源')
        self.list_tree.column('no', width=30, anchor=tk.CENTER)
        self.list_tree.column('code', width=100, anchor=tk.W)
        self.list_tree.column('name', width=160, anchor=tk.W)
        self.list_tree.column('source', width=60, anchor=tk.W)
        list_scroll = ttk.Scrollbar(list_frame, orient=tk.VERTICAL, command=self.list_tree.yview)
        self.list_tree.configure(yscrollcommand=list_scroll.set)
        list_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.list_tree.pack(fill=tk.BOTH, expand=True)
        self.list_tree.bind('<Double-Button-1>', self.remove_selected_code)
        self.list_tree.bind('<<TreeviewSelect>>', self.on_code_selected)

        # 检查结果标签页
        check_frame = tk.Frame(nb, bg=C['bg'])
        nb.add(check_frame, text="检查结果")
        columns = ('code', 'name', 'status', 'replacement', 'action')
        self.check_tree = ttk.Treeview(check_frame, columns=columns, show='tree headings', selectmode='extended')
        self.check_tree.heading('#0', text='#')
        self.check_tree.heading('code', text='规范编号')
        self.check_tree.heading('name', text='名称')
        self.check_tree.heading('status', text='状态')
        self.check_tree.heading('replacement', text='替代')
        self.check_tree.heading('action', text='建议')
        self.check_tree.column('#0', width=30, anchor=tk.CENTER)
        self.check_tree.column('code', width=100)
        self.check_tree.column('name', width=160)
        self.check_tree.column('status', width=60)
        self.check_tree.column('replacement', width=120)
        self.check_tree.column('action', width=60)
        check_scroll = ttk.Scrollbar(check_frame, orient=tk.VERTICAL, command=self.check_tree.yview)
        self.check_tree.configure(yscrollcommand=check_scroll.set)
        check_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.check_tree.pack(fill=tk.BOTH, expand=True)
        self.check_tree.bind('<Double-Button-1>', self.on_check_item_double_click)
        self.check_tree.bind('<<TreeviewSelect>>', self.on_check_item_selected)

        # ── 底部状态栏 ──
        bottombar = tk.Frame(self.root, bg=C['bg_dark'], height=26)
        bottombar.pack(side=tk.BOTTOM, fill=tk.X)
        bottombar.pack_propagate(False)

        self.status_var = tk.StringVar(value="就绪")
        tk.Label(bottombar, textvariable=self.status_var, font=("Microsoft YaHei UI", 8),
                 bg=C['bg_dark'], fg=C['text_muted']).pack(side=tk.LEFT, padx=8)

        self.progress_var = tk.DoubleVar(value=0.0)
        self.progress_bar = ttk.Progressbar(bottombar, variable=self.progress_var, maximum=100, length=100)
        self.progress_bar.pack(side=tk.LEFT, padx=4, pady=2)

        self.region_var = tk.StringVar(value="区域：未设置")
        tk.Label(bottombar, textvariable=self.region_var, font=("Microsoft YaHei UI", 8),
                 bg=C['bg_dark'], fg=C['text_muted']).pack(side=tk.LEFT, padx=8)

    def _get_render_dpi(self):
        """根据显示器 DPI 缩放比例动态计算渲染 DPI，消除高 DPI 屏幕上的文字虚边"""
        try:
            import ctypes
            scale = ctypes.windll.shcore.GetScaleFactorForDevice(0)  # 返回 100/125/150/200 等
            dpi = max(300, int(200 * scale / 100))
            return min(dpi, 400)  # 上限 400，避免内存过大
        except Exception:
            return 300  # 默认 300 DPI，远高于原来的 200

    def _detect_file_type(self, path):
        ext = Path(path).suffix.lower()
        if ext == '.pdf':
            return 'pdf'
        elif ext == '.docx':
            return 'docx'
        elif ext == '.txt':
            return 'txt'
        elif ext in IMAGE_EXTENSIONS:
            return 'image'
        elif ext in CAD_EXTENSIONS:
            return 'cad'
        elif ext == '.ofd':
            return 'ofd'
        return 'unknown'

    def _load_file(self, path):
        try:
            self.current_path = path
            self.file_type = self._detect_file_type(path)
            if self.file_type == 'unknown':
                self.status_var.set(f"不支持的文件格式: {Path(path).suffix}")
                return False
            self.ocr_results = []
            self.extracted_codes = []
            self.code_locations = []
            if hasattr(self, 'list_tree'):
                self.list_tree.delete(*self.list_tree.get_children())
            if hasattr(self, 'check_tree'):
                self.check_tree.delete(*self.check_tree.get_children())
            if hasattr(self, 'ocr_text'):
                self.ocr_text.delete('1.0', tk.END)
            self._rotation_angle = 0
            if self.file_type == 'pdf':
                self.convert_pdf_to_images()
            elif self.file_type == 'docx':
                self._load_docx_file()
            elif self.file_type == 'ofd':
                self._load_ofd_file()
            elif self.file_type == 'image':
                self._load_image_file()
            elif self.file_type == 'cad':
                self._load_cad_file()
            else:
                self.extract_text_file()
            return True
        except Exception as e:
            messagebox.showerror("文件加载错误", f"无法打开文件:\n{path}\n\n错误: {e}")
            self.status_var.set("文件加载失败")
            return False

    def _cleanup_temp_images(self):
        """清理临时图片文件"""
        for img_path in self.pdf_images:
            try:
                if img_path and Path(img_path).exists() and img_path != self.current_path:
                    os.remove(img_path)
            except Exception:
                pass
        self.pdf_images = []

    def _close_fitz_doc(self):
        """关闭 fitz 文档对象并清理临时图片"""
        self._cleanup_temp_images()
        if self._fitz_doc:
            try:
                self._fitz_doc.close()
            except Exception:
                pass
            self._fitz_doc = None
            self._total_pages = 0

    def _render_page_to_image(self, page_idx):
        """用 fitz 按需渲染单页为临时图片，返回图片路径"""
        if not self._fitz_doc:
            return None
        try:
            page = self._fitz_doc.load_page(page_idx)
            pix = page.get_pixmap(dpi=self._render_dpi)
            fd, img_path = tempfile.mkstemp(suffix='.png')
            os.close(fd)
            pix.save(img_path)
            return img_path
        except Exception as e:
            print(f"Render page {page_idx} error: {e}")
            return None

    def _load_pdf_lazy(self):
        """PDF 按需渲染：打开时只读取页数并渲染第 1 页，翻页时才渲染目标页"""
        if not self.current_path or self.file_type != 'pdf' or not HAS_FITZ:
            return
        self._close_fitz_doc()
        self.status_var.set("正在打开 PDF...")
        try:
            self._fitz_doc = fitz.open(self.current_path)
            self._total_pages = len(self._fitz_doc)
            if self._total_pages == 0:
                self.status_var.set("PDF 为空")
                return
            img_path = self._render_page_to_image(0)
            if img_path:
                self.pdf_images = [None] * self._total_pages
                self.pdf_images[0] = img_path
                self.status_var.set(f"PDF 已打开: {self._total_pages} 页")
                self.page_var.set(f"第 1 / {self._total_pages} 页")
                self.show_page(0)
        except Exception as e:
            messagebox.showerror("PDF 错误", f"无法打开 PDF:\n{self.current_path}\n\n错误: {e}")
            self.status_var.set("PDF 加载失败")

    def _load_docx_file(self):
        """加载 Word 文档：用 python-docx 提取文本直接渲染"""
        if not self.current_path or self.file_type != 'docx':
            return
        self._close_fitz_doc()
        self.status_var.set("正在加载 Word 文档...")
        self._rotation_angle = 0
        self._zoom_level = 1.0
        self._pan_image_x = 0
        self._pan_image_y = 0
        if not HAS_DOCX:
            messagebox.showwarning("Word 加载失败", "需要安装 python-docx 库")
            self.status_var.set("Word 加载失败")
            return
        try:
            # 直接提取文本并渲染
            self.extract_text_file()
        except Exception as e:
            messagebox.showerror("Word 错误", f"无法加载 Word 文档:\n{self.current_path}\n\n错误: {e}")
            self.status_var.set("Word 加载失败")

    def _extract_text_from_docx(self):
        """从 docx 提取纯文本（含表格内容）用于规范编号识别"""
        try:
            if not HAS_DOCX:
                return
            doc = Document(self.current_path)
            parts = []
            # 段落文本
            for p in doc.paragraphs:
                if p.text.strip():
                    parts.append(p.text)
            # 表格文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            parts.append(cell.text)
            full_text = '\n'.join(parts)
            self.ocr_results = [full_text]
            self._extract_codes_from_text(full_text)
        except Exception as e:
            print(f"[WARN] docx 文本提取失败: {e}")

    def _load_ofd_file(self):
        """加载 OFD 文件：用 PyMuPDF 直接打开并按需渲染（MuPDF 1.29+ 支持 OFD）"""
        if not self.current_path:
            return
        self._close_fitz_doc()
        self.status_var.set("正在加载 OFD 文档...")
        self._rotation_angle = 0
        self._zoom_level = 1.0
        self._pan_image_x = 0
        self._pan_image_y = 0
        if not HAS_FITZ:
            messagebox.showwarning("OFD 错误",
                "打开 OFD 文件需要安装 PyMuPDF 库。")
            self.status_var.set("OFD 加载失败")
            return
        try:
            self._fitz_doc = fitz.open(self.current_path)
            self._total_pages = len(self._fitz_doc)
            if self._total_pages == 0:
                self.status_var.set("OFD 文档为空")
                return
            img_path = self._render_page_to_image(0)
            if img_path:
                self.pdf_images = [None] * self._total_pages
                self.pdf_images[0] = img_path
                self.status_var.set(f"OFD 文档已打开: {self._total_pages} 页")
                self.page_var.set(f"第 1 / {self._total_pages} 页")
                self.show_page(0)
        except Exception as e:
            messagebox.showerror("OFD 错误",
                f"无法打开 OFD 文件:\n{self.current_path}\n\n"
                f"错误: {e}\n\n"
                "提示：OFD 格式需要 MuPDF 1.29+ 支持。")
            self.status_var.set("OFD 加载失败")

    def _load_image_file(self):
        if not self.current_path or self.file_type != 'image':
            return
        self._close_fitz_doc()
        self.status_var.set("正在加载图片...")
        self.pdf_images = []
        try:
            self.pdf_images.append(self.current_path)
            self.status_var.set(f"已加载图片: {Path(self.current_path).name}")
            self.page_var.set("第 1 / 1 页")
            if self.pdf_images:
                self.show_page(0)
        except Exception as e:
            messagebox.showerror("图片错误", f"无法加载图片:\n{self.current_path}\n\n错误: {e}")
            self.status_var.set("图片加载失败")

        # 加载 CAD 图纸（DWG → AcmeCAD 嵌入; DXF → ezdxf 渲染）
    def _load_cad_file(self):
        """加载 CAD 文件（DWG/DXF），用 ezdxf 直接渲染为图片"""
        if not self.current_path or self.file_type != 'cad':
            return
        self._close_fitz_doc()
        self.status_var.set("正在渲染 CAD 图纸...")
        self.pdf_images = []
        try:
            if not HAS_CAD:
                self.status_var.set("CAD 渲染不可用")
                messagebox.showwarning("CAD 不可用", "需要安装 ezdxf 和 matplotlib")
                return
            img_path = render_cad_to_image(self.current_path)
            if img_path:
                self.pdf_images.append(img_path)
                self.status_var.set(f"CAD 已渲染: {Path(self.current_path).name}")
                self.page_var.set("第 1 / 1 页")
                if self.pdf_images:
                    self.show_page(0)
            else:
                self.status_var.set("CAD 渲染失败")
                messagebox.showerror("CAD 错误", f"无法渲染 CAD 文件:\n{self.current_path}")
        except Exception as e:
            messagebox.showerror("CAD 错误", f"加载 CAD 文件出错:\n{self.current_path}\n\n错误: {e}")
            self.status_var.set("CAD 加载失败")

    def open_file(self):
        paths = filedialog.askopenfilenames(
            title="选择文件",
            filetypes=[
                ("所有支持的文件", "*.pdf *.docx *.txt *.png *.jpg *.jpeg *.bmp *.tiff *.tif *.webp *.dxf *.dwg"),
                ("PDF文件", "*.pdf"),
                ("Word文件", "*.docx"),
                ("文本文件", "*.txt"),
                ("图片文件", "*.png *.jpg *.jpeg *.bmp *.tiff *.tif *.webp"),
                ("CAD文件", "*.dxf *.dwg"),
                ("所有文件", "*.*")
            ])
        if not paths:
            return
        self.pdf_paths = list(paths)
        self._update_file_queue()
        if self.pdf_paths:
            if self._load_file(self.pdf_paths[0]):
                self._highlight_queue_item(0)
                self.status_var.set(f"已打开 {len(self.pdf_paths)} 个文件，当前: {Path(self.current_path).name}")

    def open_folder(self):
        folder = filedialog.askdirectory(title="选择文件夹（自动扫描所有支持的文件）")
        if not folder:
            return
        supported = []
        for ext in SUPPORTED_EXTENSIONS:
            for f in Path(folder).rglob(f'*{ext}'):
                supported.append(str(f))
        supported.sort()
        if not supported:
            messagebox.showinfo("提示", "文件夹中没有找到支持的文件")
            return
        self.pdf_paths = supported
        self._update_file_queue()
        if self.pdf_paths:
            if self._load_file(self.pdf_paths[0]):
                self._highlight_queue_item(0)
                self.status_var.set(f"已扫描文件夹，找到 {len(self.pdf_paths)} 个支持的文件")

    def _update_file_queue(self):
        self.queue_listbox.delete(0, tk.END)
        for i, p in enumerate(self.pdf_paths):
            name = Path(p).name
            self.queue_listbox.insert(tk.END, f"  {i+1}. {name}")
        self.queue_count_label.config(text=f"{len(self.pdf_paths)} 个文件")
        self._update_thumbnails()

    def _highlight_queue_item(self, idx):
        self.queue_listbox.selection_clear(0, tk.END)
        if 0 <= idx < len(self.pdf_paths):
            self.queue_listbox.selection_set(idx)
            self.queue_listbox.see(idx)

    def _on_queue_select(self, event):
        sel = self.queue_listbox.curselection()
        if sel:
            idx = sel[0]
            if 0 <= idx < len(self.pdf_paths):
                path = self.pdf_paths[idx]
                if path != self.current_path:
                    self._load_file(path)
                    self._highlight_queue_item(idx)

    def _prev_file(self):
        if not self.pdf_paths or not self.current_path:
            return
        for i, p in enumerate(self.pdf_paths):
            if Path(p).name == Path(self.current_path).name:
                prev_idx = (i - 1) % len(self.pdf_paths)
                self._load_file(self.pdf_paths[prev_idx])
                self._highlight_queue_item(prev_idx)
                return

    def _next_file(self):
        if not self.pdf_paths or not self.current_path:
            return
        for i, p in enumerate(self.pdf_paths):
            if Path(p).name == Path(self.current_path).name:
                next_idx = (i + 1) % len(self.pdf_paths)
                self._load_file(self.pdf_paths[next_idx])
                self._highlight_queue_item(next_idx)
                return

    def _clear_queue(self):
        self.pdf_paths = []
        self._file_queue = []
        self.queue_listbox.delete(0, tk.END)
        self.queue_count_label.config(text="0 个文件")
        self._clear_thumbnails()
        self.status_var.set("队列已清空")

    def _update_thumbnails(self):
        self._clear_thumbnails()
        if not self.pdf_paths:
            return
        if not hasattr(self, 'thumb_frame'):
            return
        for widget in self.thumb_frame.winfo_children():
            widget.destroy()
        for i, path in enumerate(self.pdf_paths[:50]):
            name = Path(path).name
            try:
                if HAS_PIL:
                    img = Image.open(path)
                    img.thumbnail((90, 120), Image.Resampling.LANCZOS)
                    photo = ImageTk.PhotoImage(img)
                    self._thumbnail_images.append(photo)
                    frame = ttk.Frame(self.thumb_frame)
                    frame.pack(fill=tk.X, pady=2)
                    label = ttk.Label(frame, image=photo)
                    label.pack()
                    ttk.Label(frame, text=name[:12], font=("SimSun", 7), anchor=tk.CENTER).pack(fill=tk.X)
                    label.bind('<Button-1>', lambda e, idx=i: self._on_thumb_click_internal(idx))
            except Exception:
                pass

    def _clear_thumbnails(self):
        self._thumbnail_images = []
        if hasattr(self, 'thumb_frame'):
            for widget in self.thumb_frame.winfo_children():
                widget.destroy()

    def _on_thumb_click_internal(self, idx):
        if 0 <= idx < len(self.pdf_paths):
            path = self.pdf_paths[idx]
            if path != self.current_path:
                self._load_file(path)
                self._highlight_queue_item(idx)

        # 旋转
    def _rotate_cw(self):
        self._rotation_angle = (self._rotation_angle + 90) % 360
        self._redraw_current_page()

    def _rotate_ccw(self):
        self._rotation_angle = (self._rotation_angle - 90) % 360
        self._redraw_current_page()

        # 批量处理
    def batch_process_all(self):
        if not self.pdf_paths:
            messagebox.showwarning("提示", "请先打开文件或文件夹")
            return
        if self._batch_running:
            self._batch_abort = True
            self.status_var.set("正在停止批量处理...")
            return
        self._batch_running = True
        self._batch_abort = False
        self.status_var.set(f"批量处理开始: 共 {len(self.pdf_paths)} 个文件")
        self.progress_var.set(0)
        all_codes = []
        all_code_info = {}

        def batch_thread():
            try:
                for i, path in enumerate(self.pdf_paths):
                    if self._batch_abort:
                        break
                    self.root.after(0, lambda p=i+1: self.status_var.set(
                        f"批量处理: {p}/{len(self.pdf_paths)} - {Path(path).name}"))
                    self.root.after(0, lambda p=i: self.progress_var.set((p / len(self.pdf_paths)) * 100))
                    try:
                        file_type = self._detect_file_type(path)
                        code_info = {}
                        if file_type == 'pdf' and HAS_FITZ:
                            with fitz.open(path) as doc:
                                for page_num in range(min(len(doc), 5)):
                                    page = doc.load_page(page_num)
                                    pix = page.get_pixmap(dpi=self._render_dpi)
                                    fd, img_path = tempfile.mkstemp(suffix='.png'); os.close(fd)
                                    pix.save(img_path)
                                    masked = mask_seals_pil(img_path)
                                    text, _ = ocr_image_standalone(masked)
                                    for p in [masked, img_path]:
                                        try: os.remove(p)
                                        except: pass
                                    cleaned = fullwidth_to_halfwidth(text)
                                    for c in CODE_PATTERN.findall(cleaned):
                                        norm = normalize_for_matching(c)
                                        if norm not in code_info:
                                            code_info[norm] = {'original': c, 'name': '', 'source': Path(path).name}
                        elif file_type == 'image':
                            masked = mask_seals_pil(path)
                            text, _ = ocr_image_standalone(masked)
                            try: os.remove(masked)
                            except: pass
                            cleaned = fullwidth_to_halfwidth(text)
                            for c in CODE_PATTERN.findall(cleaned):
                                norm = normalize_for_matching(c)
                                if norm not in code_info:
                                    code_info[norm] = {'original': c, 'name': '', 'source': Path(path).name}
                        elif file_type == 'cad':
                            if HAS_CAD:
                                img_path = render_cad_to_image(path)
                                if img_path:
                                    masked = mask_seals_pil(img_path)
                                    text, _ = ocr_image_standalone(masked)
                                    for p in [masked, img_path]:
                                        try: os.remove(p)
                                        except: pass
                                    cleaned = fullwidth_to_halfwidth(text)
                                    for c in CODE_PATTERN.findall(cleaned):
                                        norm = normalize_for_matching(c)
                                        if norm not in code_info:
                                            code_info[norm] = {'original': c, 'name': '', 'source': Path(path).name}
                        elif file_type == 'docx':
                            doc = Document(path)
                            full_text = '\n'.join([p.text for p in doc.paragraphs])
                            cleaned = fullwidth_to_halfwidth(full_text)
                            for c in CODE_PATTERN.findall(cleaned):
                                norm = normalize_for_matching(c)
                                if norm not in code_info:
                                    code_info[norm] = {'original': c, 'name': '', 'source': Path(path).name}
                        elif file_type == 'txt':
                            with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                                full_text = f.read()
                            cleaned = fullwidth_to_halfwidth(full_text)
                            for c in CODE_PATTERN.findall(cleaned):
                                norm = normalize_for_matching(c)
                                if norm not in code_info:
                                    code_info[norm] = {'original': c, 'name': '', 'source': Path(path).name}
                        for norm, info in code_info.items():
                            if norm not in all_code_info:
                                all_code_info[norm] = info
                            else:
                                existing = all_code_info[norm].get('source', '')
                                if Path(path).name not in existing:
                                    all_code_info[norm]['source'] = existing + ', ' + Path(path).name
                    except Exception as e:
                        print(f"Batch processing error for {path}: {e}")
                if self._batch_abort:
                    self.root.after(0, lambda: self.status_var.set("批量处理已中止"))
                    self.root.after(0, lambda: self.progress_var.set(0))
                else:
                    self.root.after(0, lambda: self._batch_finished(list(all_code_info.keys()), all_code_info))
            except Exception as e:
                self.root.after(0, lambda: self.status_var.set(f"批量处理出错: {e}"))
            finally:
                self._batch_running = False

        threading.Thread(target=batch_thread, daemon=True).start()

    def _batch_finished(self, all_codes, all_code_info):
        self.extracted_codes = all_codes
        self.extracted_code_info = all_code_info
        self.list_tree.delete(*self.list_tree.get_children())
        seen = set()
        for i, code in enumerate(self.extracted_codes, 1):
            if code in seen:
                continue
            seen.add(code)
            info = self.extracted_code_info.get(code, {})
            name = info.get('name', '')
            source = info.get('source', '')
            self.list_tree.insert('', tk.END, values=(i, info.get('original', code), name, source))
        self.status_var.set(
            f"批量处理完成! 共处理 {len(self.pdf_paths)} 个文件，识别到 {len(self.extracted_codes)} 个规范编号")
        self.progress_var.set(100)
        self.notebook.select(self.list_tree.master)
        messagebox.showinfo("批量处理完成",
                           f"处理文件: {len(self.pdf_paths)} 个\n识别规范: {len(self.extracted_codes)} 个")

        # 显示
    def show_page(self, idx):
        if idx < 0 or idx >= len(self.pdf_images):
            return
        self.pdf_canvas.delete('all')
        # 按需渲染：如果该页尚未渲染，先渲染
        if self.pdf_images[idx] is None and self._fitz_doc:
            self.status_var.set(f"正在渲染第 {idx+1} 页...")
            self.root.update_idletasks()
            img_path = self._render_page_to_image(idx)
            if img_path:
                self.pdf_images[idx] = img_path
                # 清理其它已渲染的临时页（只保留当前页，节省内存）
                for i in range(len(self.pdf_images)):
                    if i != idx and self.pdf_images[i] and self.pdf_images[i] != self.current_path:
                        try:
                            os.remove(self.pdf_images[i])
                        except Exception:
                            pass
                        self.pdf_images[i] = None
                self.status_var.set(f"第 {idx+1} / {len(self.pdf_images)} 页")
        img_path = self.pdf_images[idx]
        if img_path is None:
            return
        try:
            with Image.open(img_path) as img:
                if self._rotation_angle != 0:
                    img = img.rotate(self._rotation_angle, expand=True, resample=Image.Resampling.BICUBIC)
                self._current_base_image = img.copy()
                canvas_w = self.pdf_canvas.winfo_width() or 400
                canvas_h = self.pdf_canvas.winfo_height() or 600
                img_w, img_h = img.size
                if self._fit_mode.get() == 'fit_width':
                    scale = canvas_w / img_w
                    # fit_width 模式下，若缩放后高度超出画布则改用适应模式避免内容溢出
                    if img_h * scale > canvas_h:
                        scale = min(canvas_w / img_w, canvas_h / img_h)
                else:
                    scale = min(canvas_w / img_w, canvas_h / img_h)
                new_w, new_h = int(img_w * scale), int(img_h * scale)
                img_resized = img.resize((new_w, new_h), Image.Resampling.LANCZOS)
                self.current_img = ImageTk.PhotoImage(img_resized)
                center_x = canvas_w // 2 + getattr(self, '_pan_image_x', 0)
                center_y = canvas_h // 2 + getattr(self, '_pan_image_y', 0)
                self.current_image_item = self.pdf_canvas.create_image(center_x, center_y, image=self.current_img)
                self.page_var.set(f"第 {idx + 1} / {len(self.pdf_images)} 页")
                self.current_display_index = idx
                if self.selector:
                    self.selector.image_item_id = self.current_image_item
                if self.ocr_region:
                    self._draw_region_overlay(self.ocr_region, scale)
                self._draw_code_markers_for_page(idx, scale)
        except Exception as e:
            messagebox.showerror("图片错误", f"无法显示图片:\n{img_path}\n\n错误: {e}")
            self.status_var.set("图片显示失败")

    def _prev_page(self):
        if self.file_type not in ('pdf', 'image', 'cad') or not self.pdf_images:
            return
        idx = getattr(self, 'current_display_index', 0) - 1
        if idx < 0:
            idx = len(self.pdf_images) - 1
        self.show_page(idx)

    def _next_page(self):
        if self.file_type not in ('pdf', 'image', 'cad') or not self.pdf_images:
            return
        idx = getattr(self, 'current_display_index', 0) + 1
        if idx >= len(self.pdf_images):
            idx = 0
        self.show_page(idx)

    def _zoom_in(self):
        if not hasattr(self, '_zoom_level'):
            self._zoom_level = 1.0
        self._zoom_level = min(self._zoom_level * 1.2, 5.0)
        self._redraw_current_page()

    def _zoom_out(self):
        if not hasattr(self, '_zoom_level'):
            self._zoom_level = 1.0
        self._zoom_level = max(self._zoom_level / 1.2, 0.2)
        self._redraw_current_page()

    def _on_mouse_wheel(self, event):
        if not hasattr(self, '_zoom_level'):
            self._zoom_level = 1.0
        if event.delta > 0:
            self._zoom_level = min(self._zoom_level * 1.1, 5.0)
        else:
            self._zoom_level = max(self._zoom_level / 1.1, 0.2)
        self._redraw_current_page()

    def _reset_zoom(self):
        self._zoom_level = 1.0
        self._pan_image_x = 0
        self._pan_image_y = 0
        self._rotation_angle = 0
        self._redraw_current_page()

    def _on_pan_start(self, event):
        self._panning = True
        self._pan_start_x = event.x
        self._pan_start_y = event.y
        self.pdf_canvas.config(cursor="fleur")

    def _on_pan_drag(self, event):
        if not self._panning:
            return
        dx = event.x - self._pan_start_x
        dy = event.y - self._pan_start_y
        self._pan_image_x += dx
        self._pan_image_y += dy
        self._pan_start_x = event.x
        self._pan_start_y = event.y
        self._redraw_current_page()

    def _on_pan_end(self, event):
        self._panning = False
        self.pdf_canvas.config(cursor="")

    def _redraw_current_page(self, canvas_w=None, canvas_h=None):
        if not hasattr(self, '_current_base_image') or not self._current_base_image:
            return
        if not self.pdf_images:
            return
        from PIL import Image, ImageTk
        img = self._current_base_image
        img_w, img_h = img.size
        canvas_w = canvas_w or self.pdf_canvas.winfo_width() or 400
        canvas_h = canvas_h or self.pdf_canvas.winfo_height() or 600
        if self._fit_mode.get() == 'fit_width':
            base_scale = canvas_w / img_w
            # fit_width 模式下，若缩放后高度超出画布则改用适应模式
            if img_h * base_scale > canvas_h:
                base_scale = min(canvas_w / img_w, canvas_h / img_h)
        else:
            base_scale = min(canvas_w / img_w, canvas_h / img_h)
        scale = base_scale * getattr(self, '_zoom_level', 1.0)
        new_w, new_h = int(img_w * scale), int(img_h * scale)
        img_resized = img.resize((new_w, new_h), Image.Resampling.LANCZOS)
        self.pdf_canvas.delete('all')
        self.current_img = ImageTk.PhotoImage(img_resized)
        center_x = canvas_w // 2 + getattr(self, '_pan_image_x', 0)
        center_y = canvas_h // 2 + getattr(self, '_pan_image_y', 0)
        self.current_image_item = self.pdf_canvas.create_image(center_x, center_y, image=self.current_img)
        if self.ocr_region:
            self._draw_region_overlay(self.ocr_region, scale)
        if hasattr(self, 'current_display_index'):
            self._draw_code_markers_for_page(self.current_display_index, scale)
        self._highlight_rect_id = None
    def _on_canvas_resize(self, event):
        if hasattr(self, '_current_base_image') and self._current_base_image and self.pdf_images:
            self._redraw_current_page()

    def _start_periodic_redraw(self):
        self._last_canvas_size = (self.pdf_canvas.winfo_width(), self.pdf_canvas.winfo_height())
        self._periodic_redraw()

    def _periodic_redraw(self):
        if not self.root.winfo_exists():
            return
        if hasattr(self, '_current_base_image') and self._current_base_image and self.pdf_images:
            current_size = (self.pdf_canvas.winfo_width(), self.pdf_canvas.winfo_height())
            if current_size != getattr(self, '_last_canvas_size', None):
                self._last_canvas_size = current_size
                if current_size[0] > 10 and current_size[1] > 10:
                    self._redraw_current_page()
        self._redraw_after_id = self.root.after(200, self._periodic_redraw)

    def _draw_code_markers_for_page(self, page_idx, scale):
        if not hasattr(self, 'current_image_item'):
            return
        if hasattr(self, '_code_marker_ids'):
            for marker_id in self._code_marker_ids:
                self.pdf_canvas.delete(marker_id)
            self._code_marker_ids = []
        offset_x = 0
        offset_y = 0
        if hasattr(self, '_current_base_image') and self._current_base_image and scale:
            canvas_w = self.pdf_canvas.winfo_width() or 400
            canvas_h = self.pdf_canvas.winfo_height() or 600
            img_w, img_h = self._current_base_image.size
            new_w, new_h = int(img_w * scale), int(img_h * scale)
            offset_x = (canvas_w - new_w) // 2 + getattr(self, '_pan_image_x', 0)
            offset_y = (canvas_h - new_h) // 2 + getattr(self, '_pan_image_y', 0)
        page_codes = [loc for loc in self.code_locations if loc['page'] == page_idx]
        for loc in page_codes:
            x1, y1, x2, y2 = loc['bbox']
            if scale:
                x1, y1, x2, y2 = x1 * scale, y1 * scale, x2 * scale, y2 * scale
            x1 += offset_x
            y1 += offset_y
            x2 += offset_x
            y2 += offset_y
            rect_id = self.pdf_canvas.create_rectangle(x1, y1, x2, y2, outline='red', width=2, dash=(4, 2))
            self._code_marker_ids.append(rect_id)
            label_id = self.pdf_canvas.create_text(x1, y1 - 12, text=loc['code'], fill='red', anchor='sw', font=("SimSun", 9))
            self._code_marker_ids.append(label_id)

    def _draw_region_overlay(self, region, scale):
        if not hasattr(self, 'current_image_item'):
            return
        x1, y1, x2, y2 = region
        if scale:
            x1, y1, x2, y2 = x1 * scale, y1 * scale, x2 * scale, y2 * scale
        if hasattr(self, '_current_base_image') and self._current_base_image and scale:
            canvas_w = self.pdf_canvas.winfo_width() or 400
            canvas_h = self.pdf_canvas.winfo_height() or 600
            img_w, img_h = self._current_base_image.size
            new_w, new_h = int(img_w * scale), int(img_h * scale)
            offset_x = (canvas_w - new_w) // 2 + getattr(self, '_pan_image_x', 0)
            offset_y = (canvas_h - new_h) // 2 + getattr(self, '_pan_image_y', 0)
            x1 += offset_x
            y1 += offset_y
            x2 += offset_x
            y2 += offset_y
        if hasattr(self, '_region_overlay_id') and self._region_overlay_id:
            self.pdf_canvas.delete(self._region_overlay_id)
        self._region_overlay_id = self.pdf_canvas.create_rectangle(x1, y1, x2, y2, outline='red', width=2, dash=(4, 2))

    def start_selection(self):
        if not self.pdf_images:
            messagebox.showwarning("提示", "请先打开文件")
            return
        self.selection_mode = True
        self.status_var.set("请在预览图上拖拽选择识别区域")
        if self.selector:
            self.selector.enable()
            self.selector.image_item_id = getattr(self, 'current_image_item', None)

    def clear_region(self):
        self.ocr_region = None
        if hasattr(self, '_region_overlay_id') and self._region_overlay_id:
            self.pdf_canvas.delete(self._region_overlay_id)
            self._region_overlay_id = None
        self.region_var.set("识别区域：未设置（全页识别）")
        self.status_var.set("已清除识别区域，将使用全页识别")

    def _on_region_selected(self, region):
        if hasattr(self, '_current_base_image') and self._current_base_image:
            img_w, img_h = self._current_base_image.size
            canvas_w = self.pdf_canvas.winfo_width() or 400
            canvas_h = self.pdf_canvas.winfo_height() or 600
            scale = min(canvas_w / img_w, canvas_h / img_h)
            new_w, new_h = int(img_w * scale), int(img_h * scale)
            offset_x = (canvas_w - new_w) // 2
            offset_y = (canvas_h - new_h) // 2
            x1, y1, x2, y2 = region
            self.ocr_region = (
                max(0, (x1 - offset_x) / scale),
                max(0, (y1 - offset_y) / scale),
                min(img_w, (x2 - offset_x) / scale),
                min(img_h, (y2 - offset_y) / scale),
            )
        else:
            self.ocr_region = region
        self.selection_mode = False
        self.region_var.set(f"识别区域：({int(self.ocr_region[0])}, {int(self.ocr_region[1])}) -> ({int(self.ocr_region[2])}, {int(self.ocr_region[3])})")
        self.status_var.set("识别区域已设置")
        if hasattr(self, '_current_base_image') and self._current_base_image:
            img_w, img_h = self._current_base_image.size
            canvas_w = self.pdf_canvas.winfo_width() or 400
            canvas_h = self.pdf_canvas.winfo_height() or 600
            scale = min(canvas_w / img_w, canvas_h / img_h)
            self._draw_region_overlay(self.ocr_region, scale)

    def remove_selected_code(self, event=None):
        selected = self.list_tree.selection()
        if not selected:
            return
        for item in selected:
            code = self.list_tree.item(item, 'values')[1]
            norm = normalize_for_matching(code)
            for stored in list(self.extracted_codes):
                if normalize_for_matching(stored) == norm:
                    self.extracted_codes.remove(stored)
                    break
            self.list_tree.delete(item)
        for i, item in enumerate(self.list_tree.get_children(), 1):
            vals = self.list_tree.item(item, 'values')
            self.list_tree.item(item, values=(i, *vals[1:]))
        self.status_var.set(f"已移除选中项，剩余 {len(self.extracted_codes)} 个规范")
    def on_code_selected(self, event=None):
        selected = self.list_tree.selection()
        if not selected:
            return
        item = selected[0]
        values = self.list_tree.item(item, 'values')
        code = values[1]
        name = values[2] if len(values) > 2 else ''
        preview_name = f"{code} {name}".strip()
        if hasattr(self, '_preview_name_var'):
            self._preview_name_var.set(preview_name)
        if self.file_type == 'pdf':
            code_norm = normalize_for_matching(code)
            for loc in self.code_locations:
                if normalize_for_matching(loc['code']) == code_norm:
                    self.show_page(loc['page'])
                    break
        self._highlight_code_in_text(code)
        if self.file_type == 'pdf':
            self._highlight_standard_on_preview(code, name)

    def _crop_image_to_region(self, image_path, region):
        if region is None:
            return image_path
        try:
            img = Image.open(image_path)
            x1, y1, x2, y2 = region
            left = max(0, int(x1))
            top = max(0, int(y1))
            right = min(img.width, int(x2))
            bottom = min(img.height, int(y2))
            if right <= left or bottom <= top:
                return image_path
            cropped = img.crop((left, top, right, bottom))
            fd, out = tempfile.mkstemp(suffix=".png"); os.close(fd)
            try:
                cropped.save(out)
                return out
            except Exception:
                try:
                    os.unlink(out)
                except Exception:
                    pass
                raise
        except Exception as e:
            print(f"crop error: {e}")
            return image_path

    def _detect_and_split_columns(self, image_path):
        try:
            img = Image.open(image_path).convert('L')
            w, h = img.size
            if w < 300:
                return [image_path]
            analysis_w = 120
            analysis_h = max(1, int(h * analysis_w / w))
            analysis_img = img.resize((analysis_w, analysis_h), Image.Resampling.LANCZOS)
            binary = analysis_img.point(lambda p: 255 if p > 150 else 0)
            width, height = binary.size
            profile = [0] * width
            pixels = list(binary.getdata())
            for y in range(height):
                for x in range(width):
                    if pixels[y * width + x] == 0:
                        profile[x] += 1
            smoothed = []
            for i in range(width):
                start = max(0, i - 4)
                end = min(width, i + 5)
                smoothed.append(sum(profile[start:end]) / (end - start))
            mid_start = int(width * 0.25)
            mid_end = int(width * 0.75)
            mid_vals = smoothed[mid_start:mid_end]
            if not mid_vals:
                return [image_path]
            min_idx = mid_vals.index(min(mid_vals))
            split_x = mid_start + min_idx
            left_start = max(0, split_x - 20)
            right_end = min(width, split_x + 21)
            left_avg = sum(smoothed[left_start:split_x]) / max(1, split_x - left_start)
            right_avg = sum(smoothed[split_x + 1:right_end]) / max(1, right_end - split_x - 1)
            valley = smoothed[split_x]
            threshold = max(3, (left_avg + right_avg) * 0.25)
            if left_avg > threshold and right_avg > threshold and valley < threshold:
                scale = w / width
                split_original = max(1, int(split_x * scale))
                img_rgb = Image.open(image_path)
                left = img_rgb.crop((0, 0, split_original, h))
                right = img_rgb.crop((split_original, 0, w, h))
                fd, left_path = tempfile.mkstemp(suffix=".png"); os.close(fd)
                fd, right_path = tempfile.mkstemp(suffix=".png"); os.close(fd)
                try:
                    left.save(left_path)
                    right.save(right_path)
                    print(f"  Detected two-column layout, split at x={split_original}")
                    return [left_path, right_path]
                except Exception:
                    for p in [left_path, right_path]:
                        try:
                            os.unlink(p)
                        except Exception:
                            pass
                    raise
            return [image_path]
        except Exception as e:
            print(f"column detect error: {e}")
            return [image_path]

    def _split_pdf_page_to_columns(self, image_path):
        try:
            img = Image.open(image_path)
            w, h = img.size
            split_x = w // 2
            left = img.crop((0, 0, split_x, h))
            right = img.crop((split_x, 0, w, h))
            fd, left_path = tempfile.mkstemp(suffix=".png"); os.close(fd)
            fd, right_path = tempfile.mkstemp(suffix=".png"); os.close(fd)
            try:
                left.save(left_path)
                right.save(right_path)
                print(f"  Split A3 page at x={split_x}")
                return [left_path, right_path]
            except Exception:
                for p in [left_path, right_path]:
                    try:
                        os.unlink(p)
                    except Exception:
                        pass
                    raise
        except Exception as e:
            print(f"pdf column split error: {e}")
            return [image_path]

    def start_ocr(self):
        if self.current_path and self.file_type in ('pdf', 'image', 'cad'):
            self._ocr_current_file()
        else:
            self._ocr_from_text_dialog()

    def _ocr_current_file(self):
        if self.file_type in ('pdf', 'image', 'cad'):
            if not self.pdf_images:
                messagebox.showwarning("提示", "请先打开文件")
                return
            # 按需渲染模式下，如果有未渲染的页，先全部渲染
            if self._fitz_doc:
                self.status_var.set("正在为 OCR 准备所有页面...")
                self.root.update_idletasks()
                for i in range(len(self.pdf_images)):
                    if self.pdf_images[i] is None:
                        self.pdf_images[i] = self._render_page_to_image(i)
                    self.progress_var.set((i + 1) / len(self.pdf_images) * 50)
                    self.root.update_idletasks()
                self.progress_var.set(0)
        else:
            messagebox.showwarning("提示", "不支持的文件格式")
            return
        self.status_var.set("开始 OCR 识别...")
        self.progress_var.set(0)
        self.ocr_text.delete('1.0', tk.END)
        self.ocr_results = []
        self.extracted_codes = []
        self.code_locations = []
        self.list_tree.delete(*self.list_tree.get_children())
        if hasattr(self, '_executor') and self._executor is not None:
            try:
                self._executor.shutdown(wait=False)
            except Exception:
                pass
        self._executor = concurrent.futures.ThreadPoolExecutor(max_workers=4)
        self._ocr_queue = []
        self._ocr_done = False

        def do_ocr():
            try:
                page_code_blocks = {}
                total = len(self.pdf_images)
                for i, img_path in enumerate(self.pdf_images):
                    page_blocks = []
                    try:
                        cropped_path = self._crop_image_to_region(img_path, self.ocr_region)
                        crop_offsets = (0, 0)
                        if self.ocr_region and cropped_path != img_path:
                            crop_offsets = (int(self.ocr_region[0]), int(self.ocr_region[1]))
                        masked_path = mask_seals_pil(cropped_path)
                        column_paths = self._split_pdf_page_to_columns(masked_path)
                        page_texts = []
                        split_x = 0
                        if len(column_paths) == 2:
                            with Image.open(masked_path) as _img:
                                split_x = _img.width // 2
                        for col_idx, col_path in enumerate(column_paths):
                            try:
                                t, blocks = ocr_image_standalone(col_path)
                                page_texts.append(t)
                                for block_text, bbox in blocks:
                                    x1, y1, x2, y2 = bbox
                                    if col_idx == 1:
                                        x1 += split_x
                                        x2 += split_x
                                    if crop_offsets != (0, 0):
                                        x1 += crop_offsets[0]
                                        y1 += crop_offsets[1]
                                        x2 += crop_offsets[0]
                                        y2 += crop_offsets[1]
                                    page_blocks.append((block_text, (x1, y1, x2, y2)))
                            finally:
                                if col_path != masked_path and os.path.exists(col_path):
                                    try: os.remove(col_path)
                                    except: pass
                        text = '\n'.join(page_texts)
                        if masked_path != cropped_path and os.path.exists(masked_path):
                            try: os.remove(masked_path)
                            except: pass
                        if cropped_path != img_path and os.path.exists(cropped_path):
                            try: os.remove(cropped_path)
                            except: pass
                    except Exception as page_error:
                        text = f"OCR_PAGE_ERROR: {page_error}"
                        print(f"OCR page {i+1} error: {page_error}")
                    self.ocr_results.append(text)
                    self._ocr_queue.append(('page', i + 1, total, text))
                    page_cleaned = fullwidth_to_halfwidth(text)
                    for code in CODE_PATTERN.findall(page_cleaned):
                        normalized = code.upper().strip()
                        if normalized not in page_code_blocks:
                            page_code_blocks[normalized] = []
                        bbox = (0, 0, 0, 0)
                        for block_text, block_bbox in page_blocks:
                            if normalized.replace(' ', '') in block_text.replace(' ', '') or block_text.replace(' ', '') in normalized.replace(' ', ''):
                                bbox = block_bbox
                                break
                        page_code_blocks[normalized].append((i, bbox))
                self._ocr_queue.append(('status', '正在提取规范编号...'))
                all_text = '\n'.join(self.ocr_results)
                cleaned_text = fullwidth_to_halfwidth(all_text)
                raw_names = NAME_PATTERN.findall(cleaned_text)
                name_map = {}
                for raw_name in raw_names:
                    name_map[raw_name.strip()] = raw_name
                seen = set()
                self.extracted_codes = []
                self.extracted_code_info = {}
                for code in list(page_code_blocks.keys()):
                    if code in seen:
                        continue
                    seen.add(code)
                    self.extracted_codes.append(code)
                    matched_name = ''
                    for cname in name_map:
                        if code.replace(' ', '') in cname.replace(' ', '') or cname.replace(' ', '') in code.replace(' ', ''):
                            matched_name = name_map[cname]
                            break
                    self.extracted_code_info[normalize_for_matching(code)] = {'name': matched_name, 'original': code}
                    first_page, first_bbox = page_code_blocks[code][0]
                    self.code_locations.append({'code': code, 'page': first_page, 'bbox': first_bbox})
                self._ocr_queue.append(('codes', self.extracted_codes))
            except Exception as e:
                self._ocr_queue.append(('status', f'OCR 出错: {e}'))
                print(f"OCR fatal error: {e}")
            finally:
                self._ocr_done = True

        def process_queue():
            if not self._ocr_queue and not self._ocr_done:
                self.root.after(50, process_queue)
                return
            while self._ocr_queue:
                item = self._ocr_queue.pop(0)
                kind = item[0]
            if kind == 'page':
                _, page_no, total, text = item
                self.ocr_text.insert(tk.END, f"--- 第{page_no}页 ---\n{text}\n\n")
                self.ocr_text.see(tk.END)
                self.progress_var.set(page_no / total * 100)
                self.status_var.set(f"OCR 识别中: {page_no}/{total}")
                # 每页作为一个单独的 AI 气泡显示
                if self.ai_chat:
                    self.ai_chat.add_message("ai", f"📄 OCR 第{page_no}/{total}页\n\n{text}")
            elif kind == 'status':
                _, msg = item
                self.status_var.set(msg)
            elif kind == 'codes':
                codes = item[1]
                self.list_tree.delete(*self.list_tree.get_children())
                if codes:
                    for i, code in enumerate(codes, 1):
                        info = self.extracted_code_info.get(normalize_for_matching(code), {})
                        name = info.get('name', '')
                        self.list_tree.insert('', tk.END, values=(i, code, name, ''))
                    self.notebook.select(self.list_tree.master)
                    self.status_var.set(f"OCR 完成: 识别到 {len(codes)} 个规范编号")
                    self._push_ocr_to_ai()
                else:
                    sample = '\n'.join(self.ocr_results[:3])
                    self.list_tree.insert('', tk.END, values=(1, '【未识别到规范编号】'))
                    self.list_tree.insert('', tk.END, values=(2, '请查看 OCR 识别文本 确认内容'))
                    if sample.strip():
                        self.list_tree.insert('', tk.END, values=(3, sample[:120].replace('\n', ' ')))
                    self.notebook.select(self.list_tree.master)
                    self.status_var.set("OCR 完成，但未识别到规范编号")
                self.progress_var.set(100)
                if self.pdf_images:
                    self.show_page(self.current_display_index)
                    if not self._ocr_done or self._ocr_queue:
                        self.root.after(50, process_queue)
                    else:
                        self._ocr_queue = []
                        self._ocr_done = False
    
        self._executor.submit(do_ocr)
        self.root.after(50, process_queue)

    def _ocr_from_text_dialog(self):
        dialog = tk.Toplevel(self.root)
        dialog.title("输入文本进行 OCR")
        dialog.geometry("600x400")
        dialog.transient(self.root)
        dialog.grab_set()
        ttk.Label(dialog, text="请输入或粘贴需要检查的文本：").pack(anchor=tk.W, padx=10, pady=(10, 5))
        text_widget = tk.Text(dialog, wrap=tk.WORD, font=("SimSun", 10))
        text_widget.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        btn_frame = ttk.Frame(dialog)
        btn_frame.pack(side=tk.BOTTOM, fill=tk.X, padx=10, pady=10)

        def do_ocr_text():
            text = text_widget.get('1.0', tk.END).strip()
            if not text:
                messagebox.showwarning("提示", "请输入文本")
                return
            dialog.destroy()
            self.ocr_results = [text]
            self.ocr_text.delete('1.0', tk.END)
            self.ocr_text.insert(tk.END, text)
            self._extract_codes_from_text(text)
            self._push_ocr_to_ai()

        ttk.Button(btn_frame, text="开始 OCR", command=do_ocr_text).pack(side=tk.RIGHT)
        ttk.Button(btn_frame, text="取消", command=dialog.destroy).pack(side=tk.RIGHT, padx=(0, 5))

    def _extract_codes_from_text(self, text):
        raw_codes = CODE_PATTERN.findall(text)
        raw_names = NAME_PATTERN.findall(text)
        name_map = {}
        for raw_name in raw_names:
            cleaned = normalize_for_matching(raw_name).strip()
            if cleaned:
                name_map[cleaned] = raw_name
        for code in raw_codes:
            code_pos = text.find(code)
            if code_pos >= 0:
                after_text = text[code_pos + len(code):code_pos + len(code) + 80]
                name_match = re.search(r'[一-鿿]{2,30}', after_text)
                if name_match:
                    potential_name = name_match.group(0)
                    cleaned_name = normalize_for_matching(potential_name).strip()
                    if cleaned_name and len(cleaned_name) > 2:
                        name_map[cleaned_name] = potential_name
        seen = set()
        self.extracted_codes = []
        self.extracted_code_info = {}
        for code in raw_codes:
            normalized = normalize_for_matching(code)
            if normalized not in seen:
                seen.add(normalized)
                self.extracted_codes.append(code)
                matched_name = ''
                for cname in name_map:
                    if normalized.replace(' ', '') in cname.replace(' ', '') or cname.replace(' ', '') in normalized.replace(' ', ''):
                        matched_name = name_map[cname]
                        break
                self.extracted_code_info[normalized] = {'name': matched_name, 'original': code}
        for i, code in enumerate(self.extracted_codes, 1):
            info = self.extracted_code_info.get(normalize_for_matching(code), {})
            name = info.get('name', '')
            self.list_tree.insert('', tk.END, values=(i, code, name, ''))
        if self.extracted_codes:
            self.notebook.select(self.list_tree.master)
            self.status_var.set(f"提取完成: 识别到 {len(self.extracted_codes)} 个规范编号")

        # 检查规范
    def check_standards(self):
        if not self.extracted_codes:
            messagebox.showwarning("提示", "请先进行 OCR 识别并提取规范编号")
            return
        if not self._data_loaded or self.checker is None:
            messagebox.showwarning("提示", "标准数据库正在加载中，请稍候...")
            return
        self.status_var.set("检查规范中...")
        self.progress_var.set(0)
        self.check_tree.delete(*self.check_tree.get_children())
        self.check_results = []
        unique_codes = list(self.extracted_codes)
        total = len(unique_codes)
        for i, code in enumerate(unique_codes):
            info = self.extracted_code_info.get(normalize_for_matching(code), {})
            name = info.get('name', '')
            result = self.checker.check_code(code, name=name)
            self.check_results.append((code, result))
            status = result.get('status', '未找到')
            replacement = result.get('replacement_raw', '')
            matched = result.get('matched_name', result.get('matched_code', ''))
            display_code = code
            if matched:
                display_code = f"{code} -> {matched}"
            elif not result.get('found'):
                similar = self.checker.find_similar_codes(code, limit=2)
                if similar:
                    similar_str = '; '.join([f"{s[1]}《{s[2]}《"[:60] for s in similar])
                    display_code = f"{code} [相似:{similar_str}]"
            if result.get('found'):
                if '废止' in status or '作废' in status:
                    action = '需替换'
                else:
                    action = '现行'
                if result.get('dual_match'):
                    action += ' (双重确认)'
            else:
                action = '未查询到'
            matched_name = result.get('matched_name', result.get('matched_code', '')) or name
            self.check_tree.insert('', tk.END, text=str(i+1),
                                   values=(display_code, matched_name, status, replacement, action))
            self.progress_var.set((i + 1) / total * 100)
            self.root.update_idletasks()
        self.progress_var.set(100)
        self.status_var.set(f"检查完成: {len(unique_codes)} 个规范")
        self.notebook.select(self.check_tree.master)
        # 推送结果到 AI 聊天窗口
        if self.ai_chat is not None:
            self.ai_chat.send_standard_check(self.check_results)

    def export_doc(self):
        if not HAS_DOCX:
            messagebox.showwarning("提示", "需要安装 python-docx 库才能导出 Word 报告")
            return
        if not self.check_results:
            messagebox.showwarning("提示", "没有检查结果可导出")
            return
        path = filedialog.asksaveasfilename(
            title="保存 DOC 报告",
            defaultextension=".docx",
            filetypes=[("Word documents", "*.docx"), ("All files", "*.*")])
        if not path:
            return
        self.status_var.set("正在生成报告...")
        doc = Document()
        title = doc.add_heading('标准规范检查报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'文件: {os.path.basename(self.current_path) if self.current_path else "N/A"}')
        doc.add_paragraph(f'队列文件数: {len(self.pdf_paths)}')
        doc.add_paragraph()
        doc.add_heading('检查摘要', 1)
        total = len(self.check_results)
        found = sum(1 for _, r in self.check_results if r.get('found'))
        obsolete = sum(1 for _, r in self.check_results if '废止' in r.get('status', '') or '作废' in r.get('status', ''))
        doc.add_paragraph(f'共识别 {total} 个规范编号')
        doc.add_paragraph(f'数据库中查询到 {found} 个')
        doc.add_paragraph(f'其中废止/作废 {obsolete} 个')
        doc.add_paragraph()
        doc.add_heading('详细检查结果', 1)
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '序号'
        hdr_cells[1].text = '规范编号'
        hdr_cells[2].text = '规范名称'
        hdr_cells[3].text = '状态'
        hdr_cells[4].text = '替代情况'
        hdr_cells[5].text = '建议'
        for i, (code, result) in enumerate(self.check_results, 1):
            status = result.get('status', '未找到')
            replacement = result.get('replacement_raw', '')
            matched_name = result.get('matched_name', result.get('matched_code', ''))
            if result.get('found'):
                if '废止' in status or '作废' in status:
                    action = '需替换'
                else:
                    action = '现行'
            else:
                action = '未查询到'
            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = code
            row_cells[2].text = matched_name
            row_cells[3].text = status
            row_cells[4].text = replacement
            row_cells[5].text = action
        doc.save(path)
        self.progress_var.set(0)
        self.status_var.set(f"报告已保存: {path}")
        messagebox.showinfo("完成", f"报告已保存到:\n{path}")


    def export_excel(self):
        if not HAS_OPENPYXL:
            messagebox.showwarning("提示", "需要安装 openpyxl 库才能导出 Excel 报告")
            return
        if not self.check_results:
            messagebox.showwarning("提示", "没有检查结果可导出")
            return
        path = filedialog.asksaveasfilename(
            title="保存 Excel 报告",
            defaultextension=".xlsx",
            filetypes=[("Excel files", "*.xlsx"), ("All files", "*.*")])
        if not path:
            return
        self.status_var.set("正在生成 Excel 报告...")
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "标准规范检查报告"
    
        header_font = Font(bold=True, color="FFFFFF", size=11)
        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        header_align = Alignment(horizontal="center", vertical="center")
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin'))
    
        ws.merge_cells('A1:F1')
        ws['A1'] = '标准规范检查报告'
        ws['A1'].font = Font(bold=True, size=14)
        ws['A1'].alignment = Alignment(horizontal="center")
    
        total = len(self.check_results)
        found = sum(1 for _, r in self.check_results if r.get('found'))
        obsolete = sum(1 for _, r in self.check_results if '废止' in r.get('status', '') or '作废' in r.get('status', ''))
        ws['A3'] = f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}'
        ws['A4'] = f'文件: {os.path.basename(self.current_path) if self.current_path else "N/A"}'
        ws['A5'] = f'共识别 {total} 个规范编号，数据库中查询到 {found} 个，其中废止/作废 {obsolete} 个'
    
        headers = ['序号', '规范编号', '规范名称', '状态', '替代情况', '建议']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=7, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_align
            cell.border = thin_border
    
        for i, (code, result) in enumerate(self.check_results, 1):
            status = result.get('status', '未找到')
            replacement = result.get('replacement_raw', '')
            matched_name = result.get('matched_name', result.get('matched_code', ''))
            if result.get('found'):
                if '废止' in status or '作废' in status:
                    action = '需替换'
                else:
                    action = '现行'
            else:
                action = '未查询到'
            row = i + 7
            ws.cell(row=row, column=1, value=i).border = thin_border
            ws.cell(row=row, column=2, value=code).border = thin_border
            ws.cell(row=row, column=3, value=matched_name).border = thin_border
            ws.cell(row=row, column=4, value=status).border = thin_border
            ws.cell(row=row, column=5, value=replacement).border = thin_border
            ws.cell(row=row, column=6, value=action).border = thin_border
    
        ws.column_dimensions['A'].width = 8
        ws.column_dimensions['B'].width = 20
        ws.column_dimensions['C'].width = 40
        ws.column_dimensions['D'].width = 15
        ws.column_dimensions['E'].width = 30
        ws.column_dimensions['F'].width = 12
    
        wb.save(path)
        self.progress_var.set(0)
        self.status_var.set(f"Excel 报告已保存: {path}")
        messagebox.showinfo("完成", f"Excel 报告已保存到:\n{path}")
    def on_check_item_selected(self, event=None):
        selected = self.check_tree.selection()
        if not selected:
            return
        item = selected[0]
        values = self.check_tree.item(item, 'values')
        if not values:
            return
        display_code = values[0]
        name = values[1] if len(values) > 1 else ''
        preview_name = f"{display_code.split(' ')[0].split('->')[0].split('[')[0].strip()} {name}".strip()
        if hasattr(self, '_preview_name_var'):
            self._preview_name_var.set(preview_name)
        original_code = display_code.split(' ')[0].split('->')[0].split('[')[0].strip()
        code_norm = normalize_for_matching(original_code)
        if self.file_type == 'pdf':
            for loc in self.code_locations:
                if normalize_for_matching(loc['code']) == code_norm:
                    self.show_page(loc['page'])
                    self._highlight_code_location(loc)
                    break
            for item in self.list_tree.get_children():
                values = self.list_tree.item(item, 'values')
                if len(values) > 1 and normalize_for_matching(values[1]) == code_norm:
                    self.list_tree.selection_set(item)
                    self.list_tree.see(item)
                    break
        if self.file_type == 'pdf':
            self._highlight_standard_on_preview(original_code, name)

    def _highlight_code_location(self, loc):
        if not hasattr(self, '_current_base_image') or not self._current_base_image:
            return
        page_idx = loc.get('page', 0)
        if page_idx != getattr(self, 'current_display_index', -1):
            self.show_page(page_idx)
            self.root.update_idletasks()
        scale = getattr(self, '_zoom_level', 1.0)
        if hasattr(self, '_current_base_image') and self._current_base_image:
            canvas_w = self.pdf_canvas.winfo_width() or 400
            canvas_h = self.pdf_canvas.winfo_height() or 600
            img_w, img_h = self._current_base_image.size
            base_scale = min(canvas_w / img_w, canvas_h / img_h)
            scale = base_scale * scale
            offset_x = (canvas_w - int(img_w * scale)) // 2
            offset_y = (canvas_h - int(img_h * scale)) // 2
            x1, y1, x2, y2 = loc.get('bbox', (0, 0, 0, 0))
            if all(v == 0 for v in (x1, y1, x2, y2)):
                return
            x1, y1, x2, y2 = x1 * scale + offset_x, y1 * scale + offset_y, x2 * scale + offset_x, y2 * scale + offset_y
            if hasattr(self, '_highlight_rect_id') and self._highlight_rect_id:
                self.pdf_canvas.delete(self._highlight_rect_id)
            self._highlight_rect_id = self.pdf_canvas.create_rectangle(
                x1 - 2, y1 - 2, x2 + 2, y2 + 2, outline='red', width=3, dash=())
            self.root.after(3000, self._clear_highlight)

    def _clear_highlight(self):
        if hasattr(self, '_highlight_rect_id') and self._highlight_rect_id:
            self.pdf_canvas.delete(self._highlight_rect_id)
            self._highlight_rect_id = None

    def _highlight_code_in_text(self, code):
        if not hasattr(self, 'ocr_text'):
            return
        self.ocr_text.tag_remove('highlight', '1.0', tk.END)
        if not code:
            return
        start = '1.0'
        while True:
            pos = self.ocr_text.search(code, start, stopindex=tk.END, nocase=True)
            if not pos:
                break
            end = f"{pos}+{len(code)}c"
            self.ocr_text.tag_add('highlight', pos, end)
            start = end
        self.ocr_text.tag_config('highlight', background='yellow', foreground='red')

    def _highlight_standard_on_preview(self, code, name):
        if self.file_type != 'pdf' or not getattr(self, 'current_path', None):
            return
        if not hasattr(self, '_current_base_image') or not self._current_base_image:
            return
        try:
            with fitz.open(self.current_path) as doc:
                page_idx = getattr(self, 'current_display_index', 0)
                if page_idx < 0 or page_idx >= len(doc):
                    return
                page = doc.load_page(page_idx)
                rect = None
                for search_text in (code, name):
                    if not search_text:
                        continue
                    blocks = page.search_for(search_text)
                    if blocks:
                        rect = blocks[0]
                        break
            if not rect:
                return
            canvas_w = self.pdf_canvas.winfo_width() or 400
            canvas_h = self.pdf_canvas.winfo_height() or 600
            img_w, img_h = self._current_base_image.size
            dpi = self._render_dpi
            scale_factor = dpi / 72.0
            base_scale = min(canvas_w / img_w, canvas_h / img_h)
            zoom = getattr(self, '_zoom_level', 1.0)
            scale = base_scale * zoom
            offset_x = (canvas_w - int(img_w * scale)) // 2 + getattr(self, '_pan_image_x', 0)
            offset_y = (canvas_h - int(img_h * scale)) // 2 + getattr(self, '_pan_image_y', 0)
            x1 = rect.x0 * scale_factor * scale + offset_x
            y1 = rect.y0 * scale_factor * scale + offset_y
            x2 = rect.x1 * scale_factor * scale + offset_x
            y2 = rect.y1 * scale_factor * scale + offset_y
            if hasattr(self, '_highlight_rect_id') and self._highlight_rect_id:
                self.pdf_canvas.delete(self._highlight_rect_id)
            self._highlight_rect_id = self.pdf_canvas.create_rectangle(
                x1, y1, x2, y2, outline='red', width=3, dash=())
        except Exception as e:
            print(f"highlight error: {e}")

    def on_check_item_double_click(self, event=None):
        selected = self.check_tree.selection()
        if not selected:
            return
        item = selected[0]
        values = self.check_tree.item(item, 'values')
        if not values:
            return
        action = values[4] if len(values) > 4 else ''
        if action != '未查询到':
            return
        display_code = values[0]
        original_code = display_code.split(' ')[0].split('->')[0].split('[')[0].strip()
        name = ''
        if hasattr(self, 'extracted_code_info'):
            info = self.extracted_code_info.get(normalize_for_matching(original_code), {})
            name = info.get('name', '')
        dialog = StandardSearchDialog(self, self.checker, code=original_code, name=name) if self.checker else None
        if dialog:
            self.wait_window(dialog)

        # AI 聊天悬浮窗集成
    def _check_ai_config(self):
        """启动时检查 AI 配置是否已设置（首次使用弹出配置）"""
        if not _CONFIG_FILE.exists():
            # 首次使用，弹出配置对话框
            self._show_ai_config_dialog(force=True)

    def _show_ai_config_dialog(self, force=False):
        """显示 AI API 配置对话框（暗色主题）"""
        C = {'bg': "#1E293B", 'bg_dark': "#0F172A", 'card': "#334155",
             'text': "#E2E8F0", 'text_muted': "#94A3B8",
             'primary': "#3B82F6", 'primary_hover': "#2563EB",
             'success': "#22C55E", 'danger': "#EF4444"}
        config = _load_ai_config()
        dialog = tk.Toplevel(self.root)
        dialog.title("AI 助手配置")
        dialog.configure(bg=C['bg_dark'])
        dialog.geometry("520x380")
        dialog.transient(self.root)
        dialog.grab_set()

        # 标题
        tk.Label(dialog, text="🤖 AI 助手 API 配置", font=("Microsoft YaHei UI", 12, "bold"),
                 bg=C['bg_dark'], fg=C['text']).pack(anchor=tk.W, padx=16, pady=(12, 2))
        if force:
            tk.Label(dialog, text="首次使用请配置 AI API 信息（使用本地服务可直接保存）",
                     font=("Microsoft YaHei UI", 9), bg=C['bg_dark'], fg=C['danger']).pack(anchor=tk.W, padx=16, pady=(0, 8))

        frame = tk.Frame(dialog, bg=C['bg'], padx=16, pady=8)
        frame.pack(fill=tk.BOTH, expand=True)

        # API 地址
        tk.Label(frame, text="API 地址", font=("Microsoft YaHei UI", 9),
                 bg=C['bg'], fg=C['text']).pack(anchor=tk.W, pady=(4, 2))
        api_url_var = tk.StringVar(value=config.get("api_url", "http://localhost:3000/api/chat"))
        url_entry = tk.Entry(frame, textvariable=api_url_var, width=50,
                             font=("Microsoft YaHei UI", 9),
                             bg=C['card'], fg=C['text'],
                             insertbackground=C['text'],
                             borderwidth=0, highlightthickness=0)
        url_entry.pack(fill=tk.X, ipady=4)

        # API Key
        tk.Label(frame, text="API Key（可选，本地服务留空）", font=("Microsoft YaHei UI", 9),
                 bg=C['bg'], fg=C['text']).pack(anchor=tk.W, pady=(6, 2))
        key_row = tk.Frame(frame, bg=C['bg'])
        key_row.pack(fill=tk.X)
        api_key_var = tk.StringVar(value=config.get("api_key", ""))
        key_entry = tk.Entry(key_row, textvariable=api_key_var, width=45, show="*",
                             font=("Microsoft YaHei UI", 9),
                             bg=C['card'], fg=C['text'],
                             insertbackground=C['text'],
                             borderwidth=0, highlightthickness=0)
        key_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, ipady=4)
        def toggle_key_show():
            key_entry.config(show='' if key_entry.cget('show') == '*' else '*')
        btn_eye = tk.Label(key_row, text="👁", font=("Microsoft YaHei UI", 9),
                           bg=C['card'], fg=C['text'], cursor="hand2", padx=6)
        btn_eye.pack(side=tk.LEFT, padx=(4, 0))
        btn_eye.bind('<Button-1>', lambda e: toggle_key_show())

        # 模型 ID
        tk.Label(frame, text="模型 ID", font=("Microsoft YaHei UI", 9),
                 bg=C['bg'], fg=C['text']).pack(anchor=tk.W, pady=(6, 2))
        model_var = tk.StringVar(value=config.get("model", "glm-5.1"))
        # 暗色自定义下拉框（替代 ttk.Combobox）
        model_row = tk.Frame(frame, bg=C['bg'])
        model_row.pack(fill=tk.X, pady=(2, 0))
        model_entry = tk.Entry(model_row, textvariable=model_var, width=40,
                               font=("Microsoft YaHei UI", 9),
                               bg=C['card'], fg=C['text'],
                               insertbackground=C['text'],
                               borderwidth=0, highlightthickness=0)
        model_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, ipady=4)
        model_btn = tk.Label(model_row, text="▼", font=("Microsoft YaHei UI", 9),
                             bg=C['card'], fg=C['text_muted'], cursor="hand2", padx=8)
        model_btn.pack(side=tk.LEFT, padx=(0, 0))
        model_values = ('glm-4-flash', 'glm-4', 'glm-4-air', 'gpt-4o-mini', 'gpt-4o', 'deepseek-chat', 'qwen-plus')

        def _show_model_menu():
            menu = tk.Menu(model_row, tearoff=0,
                           bg=C['card'], fg=C['text'],
                           activebackground=C['primary'], activeforeground="#FFFFFF",
                           font=("Microsoft YaHei UI", 9))
            for val in model_values:
                menu.add_command(label=val, command=lambda v=val: model_var.set(v))
            menu.tk_popup(model_btn.winfo_rootx(), model_btn.winfo_rooty() + model_btn.winfo_height())

        model_btn.bind('<Button-1>', lambda e: _show_model_menu())

        # 快捷按钮
        local_frame = tk.Frame(frame, bg=C['bg'])
        local_frame.pack(fill=tk.X, pady=(6, 0))
        tk.Label(local_frame, text="快捷:", font=("Microsoft YaHei UI", 9),
                 bg=C['bg'], fg=C['text_muted']).pack(side=tk.LEFT)
        for txt, cb in [
            ("本地服务 (3000)", lambda: api_url_var.set("http://localhost:3000/api/chat")),
            ("智谱 GLM", lambda: [api_url_var.set("https://open.bigmodel.cn/api/paas/v4/chat/completions"), model_var.set("glm-4.9")]),
            ("DeepSeek", lambda: [api_url_var.set("https://api.deepseek.com/v1/chat/completions"), model_var.set("deepseek-chat")])
        ]:
            btn = tk.Label(local_frame, text=txt, font=("Microsoft YaHei UI", 8),
                           bg=C['card'], fg=C['text'], cursor="hand2", padx=6, pady=1)
            btn.pack(side=tk.LEFT, padx=2)
            btn.bind('<Button-1>', lambda e, c=cb: c())
            btn.bind('<Enter>', lambda e, b=btn: b.config(bg=C['primary']))
            btn.bind('<Leave>', lambda e, b=btn: b.config(bg=C['card']))

        # 状态标签
        status_lbl = tk.Label(frame, text="", font=("Microsoft YaHei UI", 9),
                              bg=C['bg'], fg=C['success'])
        status_lbl.pack(anchor=tk.W, pady=4)

        def test_connection():
            status_lbl.config(text="测试中...", fg=C['primary'])
            dialog.update()
            def do_test():
                try:
                    url = api_url_var.get().strip()
                    data = json.dumps({"message": "你好", "stream": False}).encode('utf-8')
                    req = urllib.request.Request(url, data=data,
                                                 headers={'Content-Type': 'application/json'}, method='POST')
                    resp = urllib.request.urlopen(req, timeout=10)
                    result = json.loads(resp.read().decode('utf-8'))
                    reply = result.get('reply', '') or result.get('content', '') or 'OK'
                    dialog.after(0, lambda: status_lbl.config(text="✅ 连接成功", fg=C['success']))
                except Exception as e:
                    dialog.after(0, lambda: status_lbl.config(text=f"❌ 连接失败: {str(e)[:80]}", fg=C['danger']))
            self._executor.submit(do_test)

        def save_config():
            new_config = {
                "api_url": api_url_var.get().strip(),
                "api_key": api_key_var.get().strip(),
                "model": model_var.get().strip()
            }
            if not new_config["api_url"]:
                messagebox.showwarning("提示", "请输入 API 地址", parent=dialog)
                return
            if _save_ai_config(new_config):
                self.ai_config = new_config
                try:
                    if self.ai_chat is not None:
                        self.ai_chat.config.update(self.ai_config)
                except Exception:
                    pass
                status_lbl.config(text="✅ 配置已保存", fg=C['success'])
                dialog.after(800, dialog.destroy)
            else:
                status_lbl.config(text="❌ 保存失败: 请检查文件权限", fg=C['danger'])
                messagebox.showerror("保存失败", "无法写入配置文件:\n" + str(_CONFIG_FILE) + "\n请检查文件权限或磁盘空间。", parent=dialog)

        # 底部按钮
        btn_frame = tk.Frame(dialog, bg=C['bg_dark'])
        btn_frame.pack(side=tk.BOTTOM, fill=tk.X, padx=16, pady=10)
        for txt, cmd, style in [
            ("测试连接", test_connection, C['card']),
            ("取消", dialog.destroy, C['card']),
            ("保存", save_config, C['primary']),
        ]:
            bg = style
            fg = "#FFFFFF" if bg == C['primary'] else C['text']
            btn = tk.Label(btn_frame, text=txt, font=("Microsoft YaHei UI", 9),
                           bg=bg, fg=fg, cursor="hand2", padx=14, pady=3)
            if txt == "保存":
                btn.pack(side=tk.RIGHT, padx=(4, 0))
            elif txt == "取消":
                btn.pack(side=tk.RIGHT, padx=4)
            else:
                btn.pack(side=tk.LEFT, padx=(0, 4))
            btn.bind('<Button-1>', lambda e, c=cmd: c())
            btn.bind('<Enter>', lambda e, b=btn, bg=bg: b.config(bg=C['primary_hover']))
            btn.bind('<Leave>', lambda e, b=btn, bg=bg: b.config(bg=bg))

        if force:
            dialog.protocol("WM_DELETE_WINDOW", lambda: None)  # 首次使用不能关闭

    def _toggle_ai_chat(self):
        if self.ai_chat is None:
            self.ai_chat = AIChatFloatingWindow(self.root, config=self.ai_config)
            self.ai_chat.show()
            # 如果有缓存的OCR结果，先推送
            if getattr(self, '_pending_ocr_results', None):
                self.ai_chat.set_ocr_results(self._pending_ocr_results)
                self._pending_ocr_results = None
            elif self.extracted_codes:
                results = []
                for code in self.extracted_codes:
                    info = self.extracted_code_info.get(normalize_for_matching(code), {})
                    check_result = None
                    for c, r in self.check_results:
                        if normalize_for_matching(c) == normalize_for_matching(code):
                            check_result = r
                            break
                    results.append({
                        'code': info.get('original', code),
                        'name': info.get('name', ''),
                        'source': info.get('source', ''),
                        'status': check_result.get('status', '') if check_result else '',
                        'found': check_result.get('found', False) if check_result else False
                    })
                self.ai_chat.set_ocr_results(results)

    def _push_ocr_to_ai(self):
        if not self.extracted_codes:
            return
        results = []
        for code in self.extracted_codes:
            info = self.extracted_code_info.get(normalize_for_matching(code), {})
            results.append({
                'code': info.get('original', code),
                'name': info.get('name', ''),
                'source': info.get('source', ''),
            })
        if self.ai_chat is None:
            # 缓存结果，等AI窗口打开后再推送
            self._pending_ocr_results = results
            return
        self._pending_ocr_results = None
        self.ai_chat.set_ocr_results(results)

    def convert_pdf_to_images(self):
        if not self.current_path or self.file_type != 'pdf' or not HAS_FITZ:
            return
        self.status_var.set("正在转换 PDF...")
        self.progress_var.set(0)
        self.pdf_images = []
        try:
            with fitz.open(self.current_path) as doc:
                total = len(doc)
                for page_num in range(total):
                    page = doc.load_page(page_num)
                    pix = page.get_pixmap(dpi=self._render_dpi)
                    fd, img_path = tempfile.mkstemp(suffix='.png'); os.close(fd)
                    pix.save(img_path)
                    self.pdf_images.append(img_path)
                    self.progress_var.set((page_num + 1) / total * 100)
                    self.root.update_idletasks()
            self.status_var.set(f"PDF 已转换: {len(self.pdf_images)} 页")
            self.page_var.set(f"第 1 / {len(self.pdf_images)} 页")
            self.progress_var.set(0)
            if self.pdf_images:
                self.show_page(0)
        except Exception as e:
            messagebox.showerror("PDF 错误", f"无法打开或转换 PDF 文件:\n{self.current_path}\n\n错误: {e}")
            self.status_var.set("PDF 加载失败")

    def _render_text_to_canvas(self, text, title):
        """将文本内容渲染为图片并显示在预览画布上"""
        if not HAS_PIL:
            return
        self.status_var.set(f"正在渲染文本内容: {title}...")
        try:
            lines = text.split('\n')
            if not lines or all(not l.strip() for l in lines):
                lines = ["（文件内容为空）"]
            # 自动换行（按 50 个中文字符宽度）
            wrapped = []
            max_chars = 50
            for line in lines:
                while len(line) > max_chars:
                    wrapped.append(line[:max_chars])
                    line = line[max_chars:]
                if line:
                    wrapped.append(line)
            lines = wrapped
            line_h = 28
            margin = 40
            w = max(600, margin * 2 + max_chars * 22)
            h = margin * 2 + len(lines) * line_h
            # 限制最大高度，避免过大
            if h > 1600:
                h = 1600
                line_h = max(16, (h - margin * 2) // len(lines))
            img = Image.new("RGB", (w, h), (255, 255, 255))
            draw = ImageDraw.Draw(img)
            # 尝试用系统字体，回退用位图
            font_path = None
            for fp in [
                "C:/Windows/Fonts/simsun.ttc",
                "C:/Windows/Fonts/simhei.ttf",
                "C:/Windows/Fonts/msyh.ttc",
            ]:
                if Path(fp).exists():
                    font_path = fp
                    break
            if font_path:
                try:
                    title_font = ImageFont.truetype(font_path, 20)
                    text_font = ImageFont.truetype(font_path, 16)
                except Exception:
                    title_font = ImageFont.load_default()
                    text_font = ImageFont.load_default()
            else:
                text_font = ImageFont.load_default()
                title_font = ImageFont.load_default()
                draw.text((margin, 12), title, fill=(0, 51, 102), font=title_font)
                draw.line([(margin, 40), (w - margin, 40)], fill=(0, 120, 200), width=2)
                cy = 56
                for ln in lines:
                    if cy + line_h > h - margin:
                        draw.text((margin, cy), "...（内容截断，超出预览区域）", fill=(128, 128, 128), font=text_font)
                        break
                    draw.text((margin, cy), ln, fill=(0, 0, 0), font=text_font)
                    cy += line_h
                fd, tmp = tempfile.mkstemp(suffix='.png'); os.close(fd)
                img.save(tmp)
                self.pdf_images = [tmp]
                if self.pdf_images:
                    self.show_page(0)
                self.status_var.set(f"已加载文本内容: {title}")
                self.page_var.set("文本预览")
        except Exception as e:
            print(f"text render error: {e}")
            self.status_var.set("文本渲染失败")

    def extract_text_file(self):
        if not self.current_path:
            return
        self._close_fitz_doc()
        self.status_var.set("正在提取文本...")
        self.progress_var.set(0)
        self.ocr_results = []
        self.pdf_images = []
        self.code_locations = []
        self.extracted_codes = []
        self.list_tree.delete(*self.list_tree.get_children())
        self.check_tree.delete(*self.check_tree.get_children())
        self.pdf_canvas.delete('all')
        full_text = ""
        rendered = False
        try:
            if self.file_type == 'docx' and HAS_DOCX:
                doc = Document(self.current_path)
                full_text = '\n'.join([p.text for p in doc.paragraphs])
                self.ocr_results = [full_text]
                title_text = f"Word 文档: {Path(self.current_path).name}"
                self._render_text_to_canvas(full_text, title_text)
                rendered = True
            elif self.file_type == 'txt':
                with open(self.current_path, 'r', encoding='utf-8', errors='ignore') as f:
                    full_text = f.read()
                    self.ocr_results = [full_text]
                    title_text = f"文本文件: {Path(self.current_path).name}"
                    self._render_text_to_canvas(full_text, title_text)
                    rendered = True
            else:
                messagebox.showwarning("提示", "不支持的文件格式")
                return
            if not rendered:
                self.page_var.set("文本预览")
                self.progress_var.set(100)
                self.status_var.set("文本提取完成")
                self.ocr_text.delete('1.0', tk.END)
                self.ocr_text.insert(tk.END, full_text)
                self._extract_codes_from_text(full_text)
        except Exception as e:
            messagebox.showerror("错误", f"读取文件失败: {e}")
            self.status_var.set("读取文件失败")
    def run(self):
        self._start_periodic_redraw()
        self.root.protocol("WM_DELETE_WINDOW", self._on_exit)
        self.root.mainloop()
    def _on_exit(self):
        # 取消所有 pending after 回调
        if hasattr(self, '_redraw_after_id'):
            try:
                self.root.after_cancel(self._redraw_after_id)
            except Exception:
                pass
        # 关闭线程池
        if hasattr(self, "_executor"):
            self._executor.shutdown(wait=False)
        # 清理临时图片文件
        import os
        for img in getattr(self, 'pdf_images', []):
            if img and os.path.exists(img):
                try:
                    os.unlink(img)
                except Exception:
                    pass
        # 关闭标准数据库连接
        if hasattr(self, 'checker') and self.checker is not None:
            try:
                self.checker.close()
            except Exception:
                pass
        self.root.destroy()



def main():
    print("Starting 工程助手 LDAssistant v2...")
    app = App()
    app.run()
    print("Application exited.")


if __name__ == "__main__":
    main()
