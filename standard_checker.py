#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
工程助手 LDAssistant
功能：
1. 上传PDF/WORD/TXT文件（支持多文件）
2. 选择识别区域（拖拽矩形，后续页面按同一区域识别）
3. OCR识别文字（自动排除公章等圆形印章）
4. 检查规范是否最新版/作废
5. 生成DOC报告
"""
import sqlite3
import os
import sys
import json
import re
import subprocess
import tempfile
import threading
from pathlib import Path
from datetime import datetime

import tkinter as tk
from tkinter import ttk, filedialog, messagebox

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    from PIL import Image, ImageDraw, ImageFilter, ImageCms
    HAS_PIL = True
except Exception:
    HAS_PIL = False

try:
    from standard_db import StandardChecker as SQLiteStandardChecker, normalize_for_matching as db_normalize_for_matching
    USE_SQLITE = True
except Exception:
    USE_SQLITE = False

# Fix blurry text on high-DPI Windows displays
try:
    import ctypes
    ctypes.windll.shcore.SetProcessDpiAwareness(2)
except Exception:
    try:
        ctypes.windll.user32.SetProcessDPIAware()
    except Exception:
        pass

# Paths - support both portable and original locations
_APP_DIR = Path(__file__).parent.resolve()

def _find_portable_paths():
    ocr_dir = _APP_DIR / "ocr"
    data_file = _APP_DIR / "data" / "all_standards_merged_20260629_092235.json"
    if ocr_dir.exists() and data_file.exists():
        return ocr_dir / "PaddleOCR-json.exe", data_file, ocr_dir
    return None, None, None

_PADDLE_OCR_EXE, _DATA_FILE, _OCR_DIR = _find_portable_paths()
if _PADDLE_OCR_EXE is None or _DATA_FILE is None:
    UMI_OCR_DIR = Path(r"D:/Program Files/图片文字识别/UmiOCR-data/plugins/win7_x64_PaddleOCR-json")
    PADDLE_OCR_EXE = UMI_OCR_DIR / "PaddleOCR-json.exe"
    OCR_DIR = UMI_OCR_DIR
    DATA_DIR = Path(r"J:/WorkBuddy-work/csres-standards")
    DATA_FILE = DATA_DIR / "all_standards_merged_20260629_092235.json"
else:
    PADDLE_OCR_EXE = _PADDLE_OCR_EXE
    OCR_DIR = _OCR_DIR
    DATA_FILE = _DATA_FILE

# Standard code pattern
# Supports GB/T, DB14/T, DB3206/T, CJJ, GB51038, T/CCAA, etc.
CODE_PATTERN = re.compile(r'\b(?:[A-Z]{1,5}[0-9]*(?:/[A-Z]{1,10})?)\s*\d+(?:\.\d+)?-\d{4}\b', re.IGNORECASE)
# Standard name pattern: Chinese text after code
NAME_PATTERN = re.compile(r'(?:[A-Z]{1,5}(?:/[A-Z]{1,2})?)\s*\d+(?:\.\d+)?-\d{4}\s+([\u4e00-\u9fff]{2,60})')

# Status keywords
OBsolete_KEYWORDS = ['废止', '作废', '代替', '被代替', '被...代替']


def fullwidth_to_halfwidth(text):
    """全角转半角"""
    result = []
    for ch in text:
        code = ord(ch)
        if 0xFF01 <= code <= 0xFF5E:
            result.append(chr(code - 0xFEE0))
        elif code == 0x3000:
            result.append(' ')
        else:
            result.append(ch)
    return ''.join(result)


def normalize_for_matching(text):
    """统一格式用于匹配：全角转半角、中文符号转英文符号、去除空格、修正OCR错误"""
    if not text:
        return ''
    result = fullwidth_to_halfwidth(text)
    # 常见中文符号转英文
    punct_map = {
        '\u3002': '.',   # 。
        '\u3001': ',',   # 、
        '\u301C': '~',   # ～
        '\u2014': '-',   # —
        '\u2013': '-',   # –
        '\u2026': '...', # …
        '\u201C': '"',   # "
        '\u201D': '"',   # "
        '\u2018': "'",   # '
        '\u2019': "'",   # '
        '\u00D7': 'x',   # ×
    }
    for cn, en in punct_map.items():
        result = result.replace(cn, en)
    # 去除空格
    result = re.sub(r'\s+', '', result)
    # OCR常见错误修正
    # CJJJ -> CJJ (OCR多识别了一个J)
    result = re.sub(r'CJJJ', 'CJJ', result, flags=re.IGNORECASE)
    # DGJ -> DG/TJ (OCR漏识别了斜杠和T)
    result = re.sub(r'DGJ(?=\d)', 'DG/TJ', result, flags=re.IGNORECASE)
    return result

class StandardChecker:
    """标准规范检查器：优先使用 SQLite + FTS5，回退到 JSON 内存索引。"""
    def __init__(self):
        self.data = []
        self.code_index = {}
        self.name_index = {}
        self._sqlite_checker = None
        if USE_SQLITE:
            try:
                self._sqlite_checker = SQLiteStandardChecker()
                print("[StandardChecker] 已启用 SQLite + FTS5 加速")
            except Exception as e:
                print(f"[StandardChecker] SQLite 初始化失败，回退到 JSON: {e}")
                self._sqlite_checker = None
        if self._sqlite_checker is None:
            self.load_data()

    def load_data(self):
        if not DATA_FILE.exists():
            print(f"Data file not found: {DATA_FILE}")
            return
        print(f"Loading data from {DATA_FILE}...")
        with open(DATA_FILE, 'r', encoding='utf-8') as f:
            self.data = json.load(f)
        for r in self.data:
            code = normalize_for_matching(r.get('code', ''))
            if code:
                self.code_index[code] = r
            name = r.get('name', '').strip()
            if name:
                norm_name = normalize_for_matching(name)
                self.name_index[norm_name] = r
        print(f"Loaded {len(self.data)} records, indexed {len(self.code_index)} codes, {len(self.name_index)} names")

    def check_code(self, code, name=''):
        if self._sqlite_checker is not None:
            return self._sqlite_checker.check_code(code, name=name)

        normalized = normalize_for_matching(code)
        result = {'found': False, 'status': '未找到', 'replacement_raw': '', 'publisher': '', 'implement_date': ''}
        
        # 1. Exact code match
        if normalized in self.code_index:
            r = self.code_index[normalized]
            result.update({
                'found': True,
                'status': r.get('status', ''),
                'replacement_raw': r.get('replacement_raw', ''),
                'publisher': r.get('publisher', ''),
                'implement_date': r.get('implement_date', ''),
                'matched_name': r.get('name', ''),
            })
            if name:
                norm_name = normalize_for_matching(name).strip()
                db_name = normalize_for_matching(r.get('name', '')).strip()
                if norm_name and db_name and (norm_name in db_name or db_name in norm_name):
                    result['dual_match'] = True
            return result
        
        # 2. Name match (combined name+code checking)
        if name:
            norm_name = normalize_for_matching(name).strip()
            if norm_name and norm_name in self.name_index:
                r = self.name_index[norm_name]
                result.update({
                    'found': True,
                    'status': r.get('status', ''),
                    'replacement_raw': r.get('replacement_raw', ''),
                    'publisher': r.get('publisher', ''),
                    'implement_date': r.get('implement_date', ''),
                    'matched_name': r.get('name', ''),
                })
                if code:
                    norm_code = normalize_for_matching(code).strip()
                    db_code = normalize_for_matching(r.get('code', '')).strip()
                    if norm_code and db_code and (norm_code in db_code or db_code in norm_code):
                        result['dual_match'] = True
                return result
            # Partial name match: search for names containing the query
            if norm_name and len(norm_name) >= 4:
                for k, v in self.name_index.items():
                    if norm_name in k or k in norm_name:
                        result.update({
                            'found': True,
                            'status': v.get('status', ''),
                            'replacement_raw': v.get('replacement_raw', ''),
                            'publisher': v.get('publisher', ''),
                            'implement_date': v.get('implement_date', ''),
                            'matched_name': v.get('name', ''),
                        })
                        if code:
                            norm_code = normalize_for_matching(code).strip()
                            db_code = normalize_for_matching(v.get('code', '')).strip()
                            if norm_code and db_code and (norm_code in db_code or db_code in norm_code):
                                result['dual_match'] = True
                        return result
        
        # 3. Partial/fuzzy code match
        best_match = None
        best_score = 0
        for k, v in self.code_index.items():
            # Exact substring match
            if normalized in k or k in normalized:
                score = len(normalized) / max(len(k), len(normalized))
                if score > best_score:
                    best_score = score
                    best_match = v
            # Fuzzy match: character-by-character similarity
            elif len(normalized) > 3 and len(k) > 3:
                matches = 0
                n_idx = 0
                k_idx = 0
                while n_idx < len(normalized) and k_idx < len(k):
                    if normalized[n_idx] == k[k_idx]:
                        matches += 1
                        n_idx += 1
                        k_idx += 1
                    else:
                        k_idx += 1
                similarity = matches / max(len(normalized), len(k))
                if similarity > 0.8 and similarity > best_score:
                    best_score = similarity
                    best_match = v
        
        if best_match:
            result.update({
                'found': True,
                'status': best_match.get('status', ''),
                'replacement_raw': best_match.get('replacement_raw', ''),
                'publisher': best_match.get('publisher', ''),
                'implement_date': best_match.get('implement_date', ''),
                'matched_code': best_match.get('code', ''),
                'matched_name': best_match.get('name', ''),
            })
            if name:
                norm_name = normalize_for_matching(name).strip()
                db_name = normalize_for_matching(best_match.get('name', '')).strip()
                if norm_name and db_name and (norm_name in db_name or db_name in norm_name):
                    result['dual_match'] = True
            return result
        
        return result

    def find_similar_codes(self, code, limit=5):
        """Find similar codes in database for debugging/recommendation."""
        if self._sqlite_checker is not None:
            # 在 SQLite 模式下，优先用 FTS5 做名称/发布单位检索；
            # 若输入更像 code，则回退到 LIKE 检索，避免 FTS5 对 / 等字符报错。
            norm_code = db_normalize_for_matching(code)
            raw_code = code.strip()
            rows = []
            try:
                cur = self._sqlite_checker.conn.cursor()
                cur.execute(
                    """
                    SELECT code, name, status FROM standards_fts
                    WHERE standards_fts MATCH ?
                    LIMIT ?
                    """,
                    (norm_code, limit)
                )
                rows = cur.fetchall()
            except sqlite3.OperationalError:
                rows = []

            if not rows and raw_code:
                try:
                    cur = self._sqlite_checker.conn.cursor()
                    cur.execute(
                        """
                        SELECT code, name, status FROM standards
                        WHERE code LIKE ? OR name LIKE ?
                        LIMIT ?
                        """,
                        (f'%{raw_code}%', f'%{raw_code}%', limit)
                    )
                    rows = cur.fetchall()
                except sqlite3.OperationalError:
                    rows = []

            results = []
            for r in rows:
                results.append((r['code'], r['code'], r['name'], 'sqlite'))
            return results

        normalized = normalize_for_matching(code)
        results = []
        for k, v in self.code_index.items():
            if normalized in k or k in normalized:
                results.append((k, v.get('code', ''), v.get('name', ''), 'substring'))
            elif len(normalized) > 3 and len(k) > 3:
                matches = 0
                n_idx = 0
                k_idx = 0
                while n_idx < len(normalized) and k_idx < len(k):
                    if normalized[n_idx] == k[k_idx]:
                        matches += 1
                        n_idx += 1
                        k_idx += 1
                    else:
                        k_idx += 1
                similarity = matches / max(len(normalized), len(k))
                if similarity > 0.6:
                    results.append((k, v.get('code', ''), v.get('name', ''), f'similar:{similarity:.2f}'))
        results.sort(key=lambda x: x[3], reverse=True)
        return results[:limit]


    def ocr_image(self, image_path):
        cmd = [
            str(PADDLE_OCR_EXE),
            f"-image_path={image_path}",
        ]
        try:
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=30, cwd=str(OCR_DIR))
            # Strip ANSI escape codes
            import re
            ansi_escape = re.compile(r'\x1b\[[0-9;]*m')
            # Parse JSON from output
            for line in result.stdout.split('\n'):
                line = ansi_escape.sub('', line).strip()
                if line.startswith('{'):
                    try:
                        ocr_result = json.loads(line)
                        blocks = []
                        for item in ocr_result.get('data', []):
                            text = item.get('text', '')
                            box = item.get('box', [])
                            if box and len(box) == 4:
                                xs = [p[0] for p in box]
                                ys = [p[1] for p in box]
                                bbox = (min(xs), min(ys), max(xs), max(ys))
                            else:
                                bbox = (0, 0, 0, 0)
                            blocks.append((text, bbox))
                        text = ' '.join([b[0] for b in blocks])
                        return text, blocks
                    except Exception as e:
                        print(f"OCR parse error: {e}, line: {line[:200]}")
                        pass
            # If no JSON found, return raw stdout for debugging
            cleaned = ansi_escape.sub('', result.stdout).strip()
            return cleaned, []
        except Exception as e:
            return f"OCR_ERROR: {e}", []

    def close(self):
        if getattr(self, '_sqlite_checker', None) is not None:
            self._sqlite_checker.close()


class RegionSelector:
    """拖拽选择识别区域的辅助类"""
    def __init__(self, canvas, image_item_id, on_selected):
        self.canvas = canvas
        self.image_item_id = image_item_id
        self.on_selected = on_selected
        self.start_x = None
        self.start_y = None
        self.rect_id = None
        self.active = False

    def enable(self):
        self.active = True
        self.canvas.config(cursor="cross")
        self.canvas.bind('<ButtonPress-1>', self.on_press)
        self.canvas.bind('<B1-Motion>', self.on_drag)
        self.canvas.bind('<ButtonRelease-1>', self.on_release)

    def disable(self):
        self.active = False
        self.canvas.config(cursor="")
        self.canvas.unbind('<ButtonPress-1>')
        self.canvas.unbind('<B1-Motion>')
        self.canvas.unbind('<ButtonRelease-1>')
        if self.rect_id:
            self.canvas.delete(self.rect_id)
            self.rect_id = None

    def on_press(self, event):
        if not self.active:
            return
        self.start_x = self.canvas.canvasx(event.x)
        self.start_y = self.canvas.canvasy(event.y)
        if self.rect_id:
            self.canvas.delete(self.rect_id)
        self.rect_id = self.canvas.create_rectangle(
            self.start_x, self.start_y, self.start_x, self.start_y,
            outline='red', width=2, dash=(4, 2)
        )

    def on_drag(self, event):
        if not self.active or self.rect_id is None:
            return
        cur_x = self.canvas.canvasx(event.x)
        cur_y = self.canvas.canvasy(event.y)
        self.canvas.coords(self.rect_id, self.start_x, self.start_y, cur_x, cur_y)

    def on_release(self, event):
        if not self.active or self.rect_id is None:
            return
        end_x = self.canvas.canvasx(event.x)
        end_y = self.canvas.canvasy(event.y)
        x1 = min(self.start_x, end_x)
        y1 = min(self.start_y, end_y)
        x2 = max(self.start_x, end_x)
        y2 = max(self.start_y, end_y)
        if abs(x2 - x1) < 10 or abs(y2 - y1) < 10:
            self.canvas.delete(self.rect_id)
            self.rect_id = None
            return
        if self.on_selected:
            self.on_selected((x1, y1, x2, y2))
        self.disable()


def mask_seals_pil(image_path, out_path=None):
    """简单公章过滤：检测图像中的红色圆形区域并遮盖，减少 OCR 误识别"""
    if not HAS_PIL:
        return image_path
    try:
        img = Image.open(image_path).convert("RGB")
        w, h = img.size
        # 若图片过大先缩放，加快检测
        max_side = 1600
        if max(w, h) > max_side:
            scale = max_side / max(w, h)
            img_small = img.resize((int(w * scale), int(h * scale)), Image.Resampling.LANCZOS)
        else:
            img_small = img
            scale = 1.0

        # 转 HSV，检测红色区域
        hsv = img_small.convert("HSV")
        pixels = hsv.load()
        mask = Image.new("L", img_small.size, 0)
        mask_pixels = mask.load()
        for y in range(img_small.size[1]):
            for x in range(img_small.size[0]):
                h_val, s_val, v_val = pixels[x, y]
                # 红色大致在 Hue 0-10 或 170-180
                if (h_val < 12 or h_val > 160) and s_val > 60 and v_val > 60:
                    mask_pixels[x, y] = 255

        # 简单膨胀，覆盖印章边缘
        mask = mask.filter(ImageFilter.MaxFilter(3))

        # 计算红色区域占比，只有占比超过 8% 才遮盖，避免误伤文档中的红线、红字
        red_count = sum(1 for p in mask.getdata() if p > 128)
        total = mask.width * mask.height
        should_mask = red_count > total * 0.08

        if should_mask:
            # 将小图 mask 放大回原图尺寸后再遮盖
            if scale != 1.0:
                mask = mask.resize((w, h), Image.Resampling.NEAREST)
            # 用 mask 把红色区域替换为白色
            white = Image.new("RGB", (w, h), (255, 255, 255))
            img = Image.composite(white, img, mask)

        if out_path:
            img.save(out_path)
            return out_path
        tmp_path = tempfile.mktemp(suffix='.png')
        img.save(tmp_path)
        return tmp_path
    except Exception as e:
        print(f"mask_seals error: {e}")
        return image_path


class App:
    def __init__(self):
        self.checker = StandardChecker()
        self.pdf_paths = []
        self.current_path = None
        self.file_type = None  # 'pdf', 'docx', 'txt'
        self.pdf_images = []
        self.ocr_results = []
        self.extracted_codes = []
        self.extracted_code_info = {}  # code -> {name, original}
        self.code_locations = []  # list of dicts: page_index, bbox, code
        self.check_results = []

        # Region selection state
        self.ocr_region = None
        self.selector = None
        self.selection_mode = False
        self.current_display_index = 0
        self._left_text_input = None
        self._file_text_widget = None
        self._file_text_scroll = None
        self._file_preview_frame = None
        self._text_input_frame = None
        self._left_mode_var = None

        # Pan state for preview canvas
        self._pan_start_x = 0
        self._pan_start_y = 0
        self._pan_image_x = 0
        self._pan_image_y = 0
        self._panning = False
        self._preview_name_text_id = None

        self.root = tk.Tk()
        self._name_index = {}

        self._left_mode_var = tk.StringVar(value='text')
        self.root.title("工程助手 LDAssistant")
        self.root.geometry("1280x820")
        self.root.minsize(1024, 640)

        self._left_mode_var = tk.StringVar(value='text')

        self._setup_style()
        self.setup_ui()

    def _setup_style(self):
        style = ttk.Style()
        try:
            style.theme_use("vista")
        except Exception:
            try:
                style.theme_use("xpnative")
            except Exception:
                style.theme_use("default")

        default_font = ("SimSun", 10)
        header_font = ("SimSun", 11, "bold")
        title_font = ("SimSun", 16, "bold")

        style.configure(".", font=default_font)
        style.configure("TLabel", font=default_font, padding=4)
        style.configure("TButton", font=default_font, padding=6)
        style.configure("TFrame", padding=8)
        style.configure("Header.TLabel", font=header_font)
        style.configure("Title.TLabel", font=title_font, foreground="#1f1f1f")
        style.configure("Primary.TButton", font=("SimSun", 10, "bold"))
        style.configure("Action.TButton", padding=8)
        style.configure("Treeview", rowheight=26, font=default_font)
        style.configure("Treeview.Heading", font=header_font)
        style.configure("Status.TLabel", background="#f0f0f0", relief="sunken", anchor="w", padding=6)

    def setup_ui(self):
        # Top header
        header = ttk.Frame(self.root)
        header.pack(side=tk.TOP, fill=tk.X)
        ttk.Label(header, text="工程助手 LDAssistant", style="Title.TLabel").pack(side=tk.LEFT, padx=16, pady=10)
        ttk.Label(header, text="OCR 识别 -> 规范列表 -> 状态检查 -> 导出报告", foreground="#666666").pack(side=tk.LEFT, padx=8, pady=10)

        # Logo on the right side
        if HAS_PIL:
            try:
                from PIL import Image, ImageTk
                logo_path = _APP_DIR / "LDA.png"
                if logo_path.exists():
                    logo_img = Image.open(logo_path)
                    # Resize to fit header height
                    header_height = 40
                    ratio = header_height / max(logo_img.height, 1)
                    new_size = (max(1, int(logo_img.width * ratio)), max(1, int(logo_img.height * ratio)))
                    logo_img = logo_img.resize(new_size, Image.Resampling.LANCZOS)
                    self._logo_photo = ImageTk.PhotoImage(logo_img)
                    logo_label = ttk.Label(header, image=self._logo_photo)
                    logo_label.pack(side=tk.RIGHT, padx=16, pady=6)
            except Exception as e:
                print(f"Failed to load logo: {e}")

        # Main container
        main_container = ttk.Frame(self.root)
        main_container.pack(side=tk.TOP, fill=tk.BOTH, expand=True)

        # Left pane: PDF preview / text input
        left_pane = ttk.Frame(main_container)
        left_pane.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(12, 6), pady=(0, 12))

        preview_header = ttk.Frame(left_pane)
        preview_header.pack(side=tk.TOP, fill=tk.X)
        ttk.Label(preview_header, text="文件预览", style="Header.TLabel").pack(side=tk.LEFT)

        # Mode toggle: preview file OR paste text
        mode_frame = ttk.Frame(left_pane)
        mode_frame.pack(side=tk.TOP, fill=tk.X, pady=(4, 0))
        self._left_mode_var = tk.StringVar(value='text')
        ttk.Radiobutton(mode_frame, text="粘贴文本", variable=self._left_mode_var, value='text', command=self._on_left_mode_changed).pack(side=tk.LEFT, padx=(0, 8))
        ttk.Radiobutton(mode_frame, text="文件预览", variable=self._left_mode_var, value='file', command=self._on_left_mode_changed).pack(side=tk.LEFT)

        # Container for the two modes
        left_content = ttk.Frame(left_pane)
        left_content.pack(side=tk.TOP, fill=tk.BOTH, expand=True, pady=(6, 0))

        # Text input mode
        self._text_input_frame = ttk.Frame(left_content)
        self._left_text_input = tk.Text(self._text_input_frame, wrap=tk.WORD, font=("SimSun", 10))
        text_scroll = ttk.Scrollbar(self._text_input_frame, orient=tk.VERTICAL, command=self._left_text_input.yview)
        self._left_text_input.configure(yscrollcommand=text_scroll.set)
        text_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self._left_text_input.pack(fill=tk.BOTH, expand=True)
        self._text_input_frame.pack(fill=tk.BOTH, expand=True)

        # File preview mode
        self._file_preview_frame = ttk.Frame(left_content)
        self.pdf_canvas = tk.Canvas(self._file_preview_frame, bg="#f3f3f3", highlightthickness=1, highlightbackground="#dcdcdc")
        self.pdf_canvas.pack(side=tk.TOP, fill=tk.BOTH, expand=True)
        self.pdf_canvas.bind('<Configure>', self._on_canvas_resize)
        self.pdf_canvas.bind('<MouseWheel>', self._on_mouse_wheel)
        self.pdf_canvas.bind('<ButtonPress-2>', self._on_pan_start)
        self.pdf_canvas.bind('<B2-Motion>', self._on_pan_drag)
        self.pdf_canvas.bind('<ButtonRelease-2>', self._on_pan_end)
        self._resize_after_id = None

        preview_footer = ttk.Frame(self._file_preview_frame)
        preview_footer.pack(side=tk.TOP, fill=tk.X, pady=(6, 0))
        self.page_var = tk.StringVar(value="第 0 / 0 页")
        ttk.Label(preview_footer, textvariable=self.page_var).pack(side=tk.LEFT)
        self._preview_name_var = tk.StringVar(value="")
        ttk.Label(preview_footer, textvariable=self._preview_name_var, foreground="#c00000", font=("SimSun", 10, "bold")).pack(side=tk.LEFT, padx=(12, 0))
        ttk.Button(preview_footer, text="上一页", command=self._prev_page, width=6).pack(side=tk.RIGHT, padx=(4, 0))
        ttk.Button(preview_footer, text="下一页", command=self._next_page, width=6).pack(side=tk.RIGHT, padx=(4, 0))
        ttk.Button(preview_footer, text="放大", command=self._zoom_in, width=6).pack(side=tk.RIGHT, padx=(4, 0))
        ttk.Button(preview_footer, text="缩小", command=self._zoom_out, width=6).pack(side=tk.RIGHT, padx=(4, 0))
        ttk.Button(preview_footer, text="重置", command=self._reset_zoom, width=6).pack(side=tk.RIGHT, padx=(4, 0))

        # Right pane
        right_pane = ttk.Frame(main_container)
        right_pane.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=(6, 12), pady=(0, 12))

        # Actions
        action_frame = ttk.Frame(right_pane)
        action_frame.pack(side=tk.TOP, fill=tk.X, pady=(0, 8))

        ttk.Button(action_frame, text="打开文件", command=self.open_file, style="Action.TButton").pack(side=tk.LEFT, padx=(0, 6))
        ttk.Button(action_frame, text="选择识别区域", command=self.start_selection, style="Action.TButton").pack(side=tk.LEFT, padx=(0, 6))
        ttk.Button(action_frame, text="清除区域", command=self.clear_region, style="Action.TButton").pack(side=tk.LEFT, padx=(0, 6))
        ttk.Button(action_frame, text="开始 OCR", command=self.start_ocr, style="Primary.TButton").pack(side=tk.LEFT, padx=(0, 6))
        ttk.Button(action_frame, text="检查规范", command=self.check_standards, style="Primary.TButton").pack(side=tk.LEFT, padx=(0, 6))
        ttk.Button(action_frame, text="导出报告", command=self.export_doc, style="Action.TButton").pack(side=tk.LEFT)

        self.region_var = tk.StringVar(value="识别区域：未设置（全页识别）")
        ttk.Label(right_pane, textvariable=self.region_var, foreground="#555555").pack(side=tk.TOP, anchor=tk.W)

        # Progress bar
        self.progress_var = tk.DoubleVar(value=0.0)
        self.progress_bar = ttk.Progressbar(right_pane, variable=self.progress_var, maximum=100)
        self.progress_bar.pack(side=tk.TOP, fill=tk.X, pady=(4, 8))

        # Notebook
        self.notebook = ttk.Notebook(right_pane)
        self.notebook.pack(side=tk.TOP, fill=tk.BOTH, expand=True)

        # OCR text tab
        ocr_frame = ttk.Frame(self.notebook)
        self.notebook.add(ocr_frame, text="OCR 识别文本")
        self.ocr_text = tk.Text(ocr_frame, wrap=tk.WORD, font=("SimSun", 10))
        ocr_scroll = ttk.Scrollbar(ocr_frame, orient=tk.VERTICAL, command=self.ocr_text.yview)
        self.ocr_text.configure(yscrollcommand=ocr_scroll.set)
        ocr_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.ocr_text.pack(fill=tk.BOTH, expand=True)

        # Extracted standards list tab
        list_frame = ttk.Frame(self.notebook)
        self.notebook.add(list_frame, text="识别到的规范列表")
        list_toolbar = ttk.Frame(list_frame)
        list_toolbar.pack(side=tk.TOP, fill=tk.X)
        ttk.Label(list_toolbar, text="双击可移除误识别项，单击可定位到 PDF").pack(side=tk.LEFT)
        list_columns = ('no', 'code', 'name')
        self.list_tree = ttk.Treeview(list_frame, columns=list_columns, show='headings', selectmode='extended')
        self.list_tree.heading('no', text='序号')
        self.list_tree.heading('code', text='规范编号')
        self.list_tree.heading('name', text='规范名称')
        self.list_tree.column('no', width=60, anchor=tk.CENTER)
        self.list_tree.column('code', width=160, anchor=tk.W)
        self.list_tree.column('name', width=260, anchor=tk.W)
        list_scroll = ttk.Scrollbar(list_frame, orient=tk.VERTICAL, command=self.list_tree.yview)
        self.list_tree.configure(yscrollcommand=list_scroll.set)
        list_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.list_tree.pack(fill=tk.BOTH, expand=True)
        self.list_tree.bind('<Double-Button-1>', self.remove_selected_code)
        self.list_tree.bind('<<TreeviewSelect>>', self.on_code_selected)

        # Check results tab
        check_frame = ttk.Frame(self.notebook)
        self.notebook.add(check_frame, text="规范检查结果")
        columns = ('code', 'name', 'status', 'replacement', 'action')
        self.check_tree = ttk.Treeview(check_frame, columns=columns, show='tree headings', selectmode='extended')
        self.check_tree.heading('#0', text='序号')
        self.check_tree.heading('code', text='规范编号')
        self.check_tree.heading('name', text='规范名称')
        self.check_tree.heading('status', text='状态')
        self.check_tree.heading('replacement', text='替代情况')
        self.check_tree.heading('action', text='建议')
        self.check_tree.column('#0', width=60, anchor=tk.CENTER)
        self.check_tree.column('code', width=160)
        self.check_tree.column('name', width=280)
        self.check_tree.column('status', width=90)
        self.check_tree.column('replacement', width=200)
        self.check_tree.column('action', width=100)
        check_scroll = ttk.Scrollbar(check_frame, orient=tk.VERTICAL, command=self.check_tree.yview)
        self.check_tree.configure(yscrollcommand=check_scroll.set)
        check_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.check_tree.pack(fill=tk.BOTH, expand=True)
        self.check_tree.bind('<Double-Button-1>', self.on_check_item_double_click)
        self.check_tree.bind('<<TreeviewSelect>>', self.on_check_item_selected)

        # Status bar
        self.status_var = tk.StringVar(value="就绪")
        statusbar = ttk.Label(self.root, textvariable=self.status_var, style="Status.TLabel")
        statusbar.pack(side=tk.BOTTOM, fill=tk.X)

        # Selection helper
        self.selector = RegionSelector(self.pdf_canvas, None, self._on_region_selected)

    def _on_left_mode_changed(self):
        mode = self._left_mode_var.get()
        if mode == 'text':
            self._text_input_frame.pack(fill=tk.BOTH, expand=True)
            self._file_preview_frame.forget()
        else:
            self._file_preview_frame.pack(fill=tk.BOTH, expand=True)
            self._text_input_frame.forget()

    def _get_active_text(self):
        """Get text from whichever mode is active."""
        mode = self._left_mode_var.get()
        if mode == 'text':
            return self._left_text_input.get('1.0', tk.END).strip()
        else:
            # File mode: use combined OCR results
            return '\n'.join(self.ocr_results).strip()

    def open_file(self):
        paths = filedialog.askopenfilenames(
            title="选择文件（PDF/WORD/TXT）",
            filetypes=[
                ("PDF files", "*.pdf"),
                ("Word files", "*.docx"),
                ("Text files", "*.txt"),
                ("All files", "*.*")
            ]
        )
        if not paths:
            return
        self.pdf_paths = list(paths)
        self.current_path = self.pdf_paths[0]
        ext = Path(self.current_path).suffix.lower()
        if ext == '.pdf':
            self.file_type = 'pdf'
        elif ext == '.docx':
            self.file_type = 'docx'
        elif ext == '.txt':
            self.file_type = 'txt'
        else:
            self.file_type = 'unknown'

        self.status_var.set(f"已打开 {len(self.pdf_paths)} 个文件，当前: {Path(self.current_path).name}")
        self._left_mode_var.set('file')
        self._on_left_mode_changed()
        if self.file_type == 'pdf':
            self.convert_pdf_to_images()
        else:
            self.extract_text_file()

    def convert_pdf_to_images(self):
        if not self.current_path or self.file_type != 'pdf':
            return
        self.status_var.set("正在转换 PDF...")
        self.progress_var.set(0)
        self.pdf_images = []

        doc = fitz.open(self.current_path)
        total = len(doc)
        for page_num in range(total):
            page = doc.load_page(page_num)
            pix = page.get_pixmap(dpi=200)
            img_path = tempfile.mktemp(suffix='.png')
            pix.save(img_path)
            self.pdf_images.append(img_path)
            self.progress_var.set((page_num + 1) / total * 100)
            self.root.update_idletasks()

        doc.close()
        self.status_var.set(f"PDF 已转换: {len(self.pdf_images)} 页")
        self.page_var.set(f"第 1 / {len(self.pdf_images)} 页")
        self.progress_var.set(0)

        if self.pdf_images:
            self.show_page(0)

    def extract_text_file(self):
        """Extract text from DOCX or TXT file."""
        if not self.current_path:
            return
        self.status_var.set("正在提取文本...")
        self.progress_var.set(0)
        self.ocr_results = []
        self.pdf_images = []
        self.code_locations = []
        self.extracted_codes = []
        self.list_tree.delete(*self.list_tree.get_children())
        self.check_tree.delete(*self.check_tree.get_children())
        self.pdf_canvas.delete('all')

        try:
            if self.file_type == 'docx':
                doc = Document(self.current_path)
                full_text = '\n'.join([p.text for p in doc.paragraphs])
                self.ocr_results = [full_text]
            elif self.file_type == 'txt':
                with open(self.current_path, 'r', encoding='utf-8', errors='ignore') as f:
                    full_text = f.read()
                self.ocr_results = [full_text]
            else:
                messagebox.showwarning("提示", "不支持的文件格式")
                return

            self.page_var.set("文本文件")
            self.progress_var.set(100)
            self.status_var.set("文本提取完成")

            # Show text in OCR tab
            self.ocr_text.delete('1.0', tk.END)
            self.ocr_text.insert(tk.END, full_text)

            # Store as ocr_results for consistent handling
            self.ocr_results = [full_text]

            # Extract standard codes
            self._extract_codes_from_text(full_text)

        except Exception as e:
            messagebox.showerror("错误", f"读取文件失败: {e}")
            self.status_var.set("读取文件失败")

    def _extract_codes_from_text(self, text):
        """Extract standard codes and names from text and populate list."""
        raw_codes = CODE_PATTERN.findall(text)
        raw_names = NAME_PATTERN.findall(text)
        name_map = {}
        for raw_name in raw_names:
            cleaned = normalize_for_matching(raw_name).strip()
            if cleaned:
                name_map[cleaned] = raw_name

        # Also extract names by looking at text after each code
        for code in raw_codes:
            # Find the code position in text
            code_pos = text.find(code)
            if code_pos >= 0:
                # Look at text after the code (up to 80 chars)
                after_text = text[code_pos + len(code):code_pos + len(code) + 80]
                # Extract Chinese text (the name)
                name_match = re.search(r'[一-鿿]{2,30}', after_text)
                if name_match:
                    potential_name = name_match.group(0)
                    cleaned_name = normalize_for_matching(potential_name).strip()
                    if cleaned_name and len(cleaned_name) > 2:
                        name_map[cleaned_name] = potential_name

        seen = set()
        self.extracted_codes = []
        self.extracted_code_info = {}
        for code in raw_codes:
            normalized = normalize_for_matching(code)
            if normalized not in seen:
                seen.add(normalized)
                self.extracted_codes.append(code)
                # Try to find matching name
                matched_name = ''
                for cname in name_map:
                    if normalized.replace(' ', '') in cname.replace(' ', '') or cname.replace(' ', '') in normalized.replace(' ', ''):
                        matched_name = name_map[cname]
                        break
                self.extracted_code_info[normalized] = {
                    'name': matched_name,
                    'original': code,
                }

        for i, code in enumerate(self.extracted_codes, 1):
            info = self.extracted_code_info.get(normalize_for_matching(code), {})
            name = info.get('name', '')
            self.list_tree.insert('', tk.END, values=(i, code, name))

        if self.extracted_codes:
            self.notebook.select(self.list_tree.master)
        self.status_var.set(f"提取完成: 识别到 {len(self.extracted_codes)} 个规范编号")

    def show_page(self, idx):
        if idx < 0 or idx >= len(self.pdf_images):
            return
        self.pdf_canvas.delete('all')
        img_path = self.pdf_images[idx]
        from PIL import Image, ImageTk
        img = Image.open(img_path)
        self._current_base_image = img
        canvas_w = self.pdf_canvas.winfo_width() or 400
        canvas_h = self.pdf_canvas.winfo_height() or 600
        img_w, img_h = img.size
        scale = min(canvas_w / img_w, canvas_h / img_h)
        new_w, new_h = int(img_w * scale), int(img_h * scale)
        img_resized = img.resize((new_w, new_h), Image.Resampling.LANCZOS)
        self.current_img = ImageTk.PhotoImage(img_resized)
        center_x = canvas_w // 2 + getattr(self, '_pan_image_x', 0)
        center_y = canvas_h // 2 + getattr(self, '_pan_image_y', 0)
        self.current_image_item = self.pdf_canvas.create_image(center_x, center_y, image=self.current_img)
        self.page_var.set(f"第 {idx + 1} / {len(self.pdf_images)} 页")
        self.current_display_index = idx

        # Update selector target
        if self.selector:
            self.selector.image_item_id = self.current_image_item

        # Draw existing region if any
        if self.ocr_region:
            self._draw_region_overlay(self.ocr_region, scale)

        # Draw code location markers for this page
        self._draw_code_markers_for_page(idx, scale)

    def _prev_page(self):
        if self.file_type != 'pdf' or not self.pdf_images:
            return
        idx = getattr(self, 'current_display_index', 0) - 1
        if idx < 0:
            idx = len(self.pdf_images) - 1
        self.show_page(idx)

    def _next_page(self):
        if self.file_type != 'pdf' or not self.pdf_images:
            return
        idx = getattr(self, 'current_display_index', 0) + 1
        if idx >= len(self.pdf_images):
            idx = 0
        self.show_page(idx)

    def _zoom_in(self):
        """Zoom in on PDF preview."""
        if not hasattr(self, '_zoom_level'):
            self._zoom_level = 1.0
        self._zoom_level = min(self._zoom_level * 1.2, 5.0)
        self._redraw_current_page()

    def _zoom_out(self):
        """Zoom out on PDF preview."""
        if not hasattr(self, '_zoom_level'):
            self._zoom_level = 1.0
        self._zoom_level = max(self._zoom_level / 1.2, 0.2)
        self._redraw_current_page()

    def _on_mouse_wheel(self, event):
        """Handle mouse wheel zoom."""
        if not hasattr(self, '_zoom_level'):
            self._zoom_level = 1.0
        if event.delta > 0:
            self._zoom_level = min(self._zoom_level * 1.1, 5.0)
        else:
            self._zoom_level = max(self._zoom_level / 1.1, 0.2)
        self._redraw_current_page()


    def _reset_zoom(self):
        """Reset zoom to default."""
        self._zoom_level = 1.0
        self._pan_image_x = 0
        self._pan_image_y = 0
        self._redraw_current_page()

    def _on_pan_start(self, event):
        """Start panning the preview image with middle mouse button."""
        self._panning = True
        self._pan_start_x = event.x
        self._pan_start_y = event.y
        self.pdf_canvas.config(cursor="fleur")

    def _on_pan_drag(self, event):
        """Pan the preview image while dragging."""
        if not self._panning:
            return
        dx = event.x - self._pan_start_x
        dy = event.y - self._pan_start_y
        self._pan_image_x += dx
        self._pan_image_y += dy
        self._pan_start_x = event.x
        self._pan_start_y = event.y
        self._redraw_current_page()

    def _on_pan_end(self, event):
        """End panning."""
        self._panning = False
        self.pdf_canvas.config(cursor="")

    def _draw_code_markers_for_page(self, page_idx, scale):
        """Draw markers on canvas for standard codes found on this page."""
        if not hasattr(self, 'current_image_item'):
            return
        # Remove old markers
        if hasattr(self, '_code_marker_ids'):
            for marker_id in self._code_marker_ids:
                self.pdf_canvas.delete(marker_id)
        self._code_marker_ids = []

        # Account for centered image offset on canvas
        offset_x = 0
        offset_y = 0
        if hasattr(self, '_current_base_image') and self._current_base_image and scale:
            canvas_w = self.pdf_canvas.winfo_width() or 400
            canvas_h = self.pdf_canvas.winfo_height() or 600
            img_w, img_h = self._current_base_image.size
            new_w, new_h = int(img_w * scale), int(img_h * scale)
            offset_x = (canvas_w - new_w) // 2 + getattr(self, '_pan_image_x', 0)
            offset_y = (canvas_h - new_h) // 2 + getattr(self, '_pan_image_y', 0)

        # Find codes on this page
        page_codes = [loc for loc in self.code_locations if loc['page'] == page_idx]
        for loc in page_codes:
            x1, y1, x2, y2 = loc['bbox']
            if scale:
                x1, y1, x2, y2 = x1 * scale, y1 * scale, x2 * scale, y2 * scale
            x1 += offset_x
            y1 += offset_y
            x2 += offset_x
            y2 += offset_y
            rect_id = self.pdf_canvas.create_rectangle(
                x1, y1, x2, y2,
                outline='red', width=2, dash=(4, 2)
            )
            self._code_marker_ids.append(rect_id)
            # Add label
            label_id = self.pdf_canvas.create_text(
                x1, y1 - 12, text=loc['code'],
                fill='red', anchor='sw', font=("SimSun", 9)
            )
            self._code_marker_ids.append(label_id)

    def _draw_region_overlay(self, region, scale):
        if not hasattr(self, 'current_image_item'):
            return
        x1, y1, x2, y2 = region
        if scale:
            x1, y1, x2, y2 = x1 * scale, y1 * scale, x2 * scale, y2 * scale

        # Account for centered image offset on canvas
        if hasattr(self, '_current_base_image') and self._current_base_image and scale:
            canvas_w = self.pdf_canvas.winfo_width() or 400
            canvas_h = self.pdf_canvas.winfo_height() or 600
            img_w, img_h = self._current_base_image.size
            new_w, new_h = int(img_w * scale), int(img_h * scale)
            offset_x = (canvas_w - new_w) // 2 + getattr(self, '_pan_image_x', 0)
            offset_y = (canvas_h - new_h) // 2 + getattr(self, '_pan_image_y', 0)
            x1 += offset_x
            y1 += offset_y
            x2 += offset_x
            y2 += offset_y

        if hasattr(self, '_region_overlay_id') and self._region_overlay_id:
            self.pdf_canvas.delete(self._region_overlay_id)
        self._region_overlay_id = self.pdf_canvas.create_rectangle(
            x1, y1, x2, y2, outline='red', width=2, dash=(4, 2)
        )

    def _on_canvas_resize(self, event):
        """Redraw current PDF page when canvas is resized."""
        if hasattr(self, '_current_base_image') and self._current_base_image and self.pdf_images:
            self._redraw_current_page()

    def _start_periodic_redraw(self):
        """Start periodic check for canvas resize."""
        self._last_canvas_size = (self.pdf_canvas.winfo_width(), self.pdf_canvas.winfo_height())
        self._periodic_redraw()

    def _periodic_redraw(self):
        """Periodically check if canvas size changed and redraw."""
        if hasattr(self, '_current_base_image') and self._current_base_image and self.pdf_images:
            current_size = (self.pdf_canvas.winfo_width(), self.pdf_canvas.winfo_height())
            if current_size != getattr(self, '_last_canvas_size', None):
                self._last_canvas_size = current_size
                if current_size[0] > 10 and current_size[1] > 10:
                    self._redraw_current_page()
        self.root.after(200, self._periodic_redraw)

    def _redraw_current_page(self, canvas_w=None, canvas_h=None):
        """Redraw current page with current zoom level and optional canvas size."""
        if not hasattr(self, '_current_base_image') or not self._current_base_image:
            return
        if not self.pdf_images:
            return
        
        from PIL import Image, ImageTk
        img = self._current_base_image
        img_w, img_h = img.size
        canvas_w = canvas_w or self.pdf_canvas.winfo_width() or 400
        canvas_h = canvas_h or self.pdf_canvas.winfo_height() or 600
        
        # Calculate scale with zoom level
        base_scale = min(canvas_w / img_w, canvas_h / img_h)
        scale = base_scale * getattr(self, '_zoom_level', 1.0)
        
        new_w, new_h = int(img_w * scale), int(img_h * scale)
        img_resized = img.resize((new_w, new_h), Image.Resampling.LANCZOS)
        
        self.pdf_canvas.delete('all')
        self.current_img = ImageTk.PhotoImage(img_resized)
        center_x = canvas_w // 2 + getattr(self, '_pan_image_x', 0)
        center_y = canvas_h // 2 + getattr(self, '_pan_image_y', 0)
        self.current_image_item = self.pdf_canvas.create_image(
            center_x, center_y, image=self.current_img
        )
        
        # Redraw region overlay if exists
        if self.ocr_region:
            self._draw_region_overlay(self.ocr_region, scale)
        
        # Redraw code markers for current page
        if hasattr(self, 'current_display_index'):
            self._draw_code_markers_for_page(self.current_display_index, scale)
        
        # Highlight was cleared by delete('all')
        self._highlight_rect_id = None

    def start_selection(self):
        if not self.pdf_images:
            messagebox.showwarning("提示", "请先打开PDF文件")
            return
        self.selection_mode = True
        self.status_var.set("请在预览图上拖拽选择识别区域，右键或再次点击按钮取消")
        if self.selector:
            self.selector.enable()
            self.selector.image_item_id = getattr(self, 'current_image_item', None)

    def clear_region(self):
        self.ocr_region = None
        if hasattr(self, '_region_overlay_id') and self._region_overlay_id:
            self.pdf_canvas.delete(self._region_overlay_id)
            self._region_overlay_id = None
        self.region_var.set("识别区域：未设置（全页识别）")
        self.status_var.set("已清除识别区域，将使用全页识别")

    def _on_region_selected(self, region):
        # Convert canvas drag coordinates back to original image coordinates
        if hasattr(self, '_current_base_image') and self._current_base_image:
            img_w, img_h = self._current_base_image.size
            canvas_w = self.pdf_canvas.winfo_width() or 400
            canvas_h = self.pdf_canvas.winfo_height() or 600
            scale = min(canvas_w / img_w, canvas_h / img_h)
            new_w, new_h = int(img_w * scale), int(img_h * scale)
            offset_x = (canvas_w - new_w) // 2
            offset_y = (canvas_h - new_h) // 2

            x1, y1, x2, y2 = region
            self.ocr_region = (
                max(0, (x1 - offset_x) / scale),
                max(0, (y1 - offset_y) / scale),
                min(img_w, (x2 - offset_x) / scale),
                min(img_h, (y2 - offset_y) / scale),
            )
        else:
            self.ocr_region = region

        self.selection_mode = False
        self.region_var.set(f"识别区域：({int(self.ocr_region[0])}, {int(self.ocr_region[1])}) -> ({int(self.ocr_region[2])}, {int(self.ocr_region[3])})")
        self.status_var.set("识别区域已设置，后续页面将按此区域识别")
        # Redraw overlay in original image coordinates
        if hasattr(self, '_current_base_image') and self._current_base_image:
            img_w, img_h = self._current_base_image.size
            canvas_w = self.pdf_canvas.winfo_width() or 400
            canvas_h = self.pdf_canvas.winfo_height() or 600
            scale = min(canvas_w / img_w, canvas_h / img_h)
            self._draw_region_overlay(self.ocr_region, scale)

    def remove_selected_code(self, event=None):
        selected = self.list_tree.selection()
        if not selected:
            return
        for item in selected:
            code = self.list_tree.item(item, 'values')[1]
            norm = normalize_for_matching(code)
            if norm in self.extracted_codes:
                self.extracted_codes.remove(norm)
            self.list_tree.delete(item)
        for i, item in enumerate(self.list_tree.get_children(), 1):
            self.list_tree.item(item, values=(i, self.list_tree.item(item, 'values')[1]))
        self.status_var.set(f"已移除选中项，剩余 {len(self.extracted_codes)} 个规范")

    def on_code_selected(self, event=None):
        """When user selects a code in the list, navigate to its page and highlight."""
        selected = self.list_tree.selection()
        if not selected:
            return
        item = selected[0]
        values = self.list_tree.item(item, 'values')
        code = values[1]
        name = values[2] if len(values) > 2 else ''
        
        # Update preview name label
        preview_name = f"{code} {name}".strip()
        if hasattr(self, '_preview_name_var'):
            self._preview_name_var.set(preview_name)
        
        # Navigate to page only in PDF mode
        if self.file_type == 'pdf':
            code_norm = normalize_for_matching(code)
            for loc in self.code_locations:
                if normalize_for_matching(loc['code']) == code_norm:
                    self.show_page(loc['page'])
                    break

        # Highlight in text regardless of mode
        self._highlight_code_in_text(code)
        
        # Highlight on preview for PDF
        if self.file_type == 'pdf':
            self._highlight_standard_on_preview(code, name)

    def _crop_image_to_region(self, image_path, region):
        """按选定区域裁剪图片，返回临时文件路径"""
        if region is None:
            return image_path
        try:
            img = Image.open(image_path)
            x1, y1, x2, y2 = region
            left = max(0, int(x1))
            top = max(0, int(y1))
            right = min(img.width, int(x2))
            bottom = min(img.height, int(y2))
            if right <= left or bottom <= top:
                return image_path
            cropped = img.crop((left, top, right, bottom))
            out = tempfile.mktemp(suffix='.png')
            cropped.save(out)
            return out
        except Exception as e:
            print(f"crop error: {e}")
            return image_path

    def _detect_and_split_columns(self, image_path):
        """Detect two-column layout and split image if needed. Returns list of image paths."""
        try:
            img = Image.open(image_path).convert('L')
            w, h = img.size
            if w < 300:
                return [image_path]

            # Downscale for fast analysis
            analysis_w = 120
            analysis_h = max(1, int(h * analysis_w / w))
            analysis_img = img.resize((analysis_w, analysis_h), Image.Resampling.LANCZOS)
            # Threshold: text pixels are dark
            binary = analysis_img.point(lambda p: 255 if p > 150 else 0)

            width, height = binary.size
            profile = [0] * width
            pixels = list(binary.getdata())
            for y in range(height):
                for x in range(width):
                    if pixels[y * width + x] == 0:
                        profile[x] += 1

            # Smooth profile
            smoothed = []
            for i in range(width):
                start = max(0, i - 4)
                end = min(width, i + 5)
                smoothed.append(sum(profile[start:end]) / (end - start))

            # Find valley in middle 50%
            mid_start = int(width * 0.25)
            mid_end = int(width * 0.75)
            mid_vals = smoothed[mid_start:mid_end]
            if not mid_vals:
                return [image_path]

            min_idx = mid_vals.index(min(mid_vals))
            split_x = mid_start + min_idx

            # Check if valley is significant compared to surrounding density
            left_start = max(0, split_x - 20)
            right_end = min(width, split_x + 21)
            left_avg = sum(smoothed[left_start:split_x]) / max(1, split_x - left_start)
            right_avg = sum(smoothed[split_x + 1:right_end]) / max(1, right_end - split_x - 1)
            valley = smoothed[split_x]
            threshold = max(3, (left_avg + right_avg) * 0.25)

            if left_avg > threshold and right_avg > threshold and valley < threshold:
                scale = w / width
                split_original = max(1, int(split_x * scale))
                img_rgb = Image.open(image_path)
                left = img_rgb.crop((0, 0, split_original, h))
                right = img_rgb.crop((split_original, 0, w, h))
                left_path = tempfile.mktemp(suffix='.png')
                right_path = tempfile.mktemp(suffix='.png')
                left.save(left_path)
                right.save(right_path)
                print(f"  Detected two-column layout, split at x={split_original}")
                return [left_path, right_path]

            return [image_path]
        except Exception as e:
            print(f"column detect error: {e}")
            return [image_path]

    def _split_pdf_page_to_columns(self, image_path):
        """Split PDF page image into two columns for A3 two-column layout."""
        try:
            img = Image.open(image_path)
            w, h = img.size
            # Always split at middle for A3 two-column format
            split_x = w // 2
            left = img.crop((0, 0, split_x, h))
            right = img.crop((split_x, 0, w, h))
            left_path = tempfile.mktemp(suffix='.png')
            right_path = tempfile.mktemp(suffix='.png')
            left.save(left_path)
            right.save(right_path)
            print(f"  Split A3 page at x={split_x}")
            return [left_path, right_path]
        except Exception as e:
            print(f"pdf column split error: {e}")
            return [image_path]

    def start_ocr(self):
        mode = self._left_mode_var.get()

        if mode == 'text':
            # Text input mode: extract codes from pasted text directly
            text = self._get_active_text()
            if not text:
                messagebox.showwarning("提示", "请在左侧文本框中粘贴需要检查的内容")
                return

            self.status_var.set("正在从文本中提取规范编号...")
            self.progress_var.set(0)
            self.extracted_codes = []
            self.code_locations = []
            self.check_results = []
            self.ocr_results = [text]
            self.list_tree.delete(*self.list_tree.get_children())
            self.check_tree.delete(*self.check_tree.get_children())

            # Show text in OCR tab
            self.ocr_text.delete('1.0', tk.END)
            self.ocr_text.insert(tk.END, text)

            # Extract codes
            self._extract_codes_from_text(text)
            self.progress_var.set(100)
            self.status_var.set(f"文本提取完成: 识别到 {len(self.extracted_codes)} 个规范编号")
            self.notebook.select(self.list_tree.master)
            return

        # File mode: original OCR logic
        if self.file_type == 'pdf':
            if not self.pdf_images:
                messagebox.showwarning("提示", "请先打开PDF文件")
                return
        elif self.file_type in ('docx', 'txt'):
            if not self.current_path:
                messagebox.showwarning("提示", "请先打开文件")
                return
        else:
            messagebox.showwarning("提示", "不支持的文件格式")
            return

        self.status_var.set("开始 OCR 识别...")
        self.progress_var.set(0)
        self.ocr_text.delete('1.0', tk.END)
        self.ocr_results = []
        self.extracted_codes = []
        self.code_locations = []
        self.list_tree.delete(*self.list_tree.get_children())

        self._ocr_queue = []
        self._ocr_done = False

        def do_ocr():
            try:
                page_code_blocks = {}  # normalized_code -> [(page_idx, bbox), ...]

                if self.file_type == 'pdf':
                    total = len(self.pdf_images)
                    for i, img_path in enumerate(self.pdf_images):
                        page_blocks = []
                        try:
                            cropped_path = self._crop_image_to_region(img_path, self.ocr_region)
                            crop_offsets = (0, 0)
                            if self.ocr_region and cropped_path != img_path:
                                crop_offsets = (int(self.ocr_region[0]), int(self.ocr_region[1]))
                            masked_path = mask_seals_pil(cropped_path)
                            column_paths = []
                            try:
                                column_paths = self._split_pdf_page_to_columns(masked_path)
                                page_texts = []
                                split_x = 0
                                if len(column_paths) == 2:
                                    with Image.open(masked_path) as _img:
                                        split_x = _img.width // 2
                                for col_idx, col_path in enumerate(column_paths):
                                    try:
                                        t, blocks = self.checker.ocr_image(col_path)
                                        page_texts.append(t)
                                        for block_text, bbox in blocks:
                                            x1, y1, x2, y2 = bbox
                                            if col_idx == 1:
                                                x1 += split_x
                                                x2 += split_x
                                            if crop_offsets != (0, 0):
                                                x1 += crop_offsets[0]
                                                y1 += crop_offsets[1]
                                                x2 += crop_offsets[0]
                                                y2 += crop_offsets[1]
                                            page_blocks.append((block_text, (x1, y1, x2, y2)))
                                    finally:
                                        if col_path != masked_path and os.path.exists(col_path):
                                            try:
                                                os.remove(col_path)
                                            except Exception:
                                                pass
                                text = '\n'.join(page_texts)
                            finally:
                                if masked_path != cropped_path and os.path.exists(masked_path):
                                    try:
                                        os.remove(masked_path)
                                    except Exception:
                                        pass
                                if cropped_path != img_path and os.path.exists(cropped_path):
                                    try:
                                        os.remove(cropped_path)
                                    except Exception:
                                        pass
                        except Exception as page_error:
                            text = f"OCR_PAGE_ERROR: {page_error}"
                            print(f"OCR page {i+1} error: {page_error}")

                        self.ocr_results.append(text)
                        self._ocr_queue.append(('page', i + 1, total, text))

                        page_cleaned = fullwidth_to_halfwidth(text)
                        page_codes = CODE_PATTERN.findall(page_cleaned)
                        for code in page_codes:
                            normalized = code.upper().strip()
                            if normalized not in page_code_blocks:
                                page_code_blocks[normalized] = []
                            bbox = (0, 0, 0, 0)
                            for block_text, block_bbox in page_blocks:
                                if normalized.replace(' ', '') in block_text.replace(' ', '') or block_text.replace(' ', '') in normalized.replace(' ', ''):
                                    bbox = block_bbox
                                    break
                            page_code_blocks[normalized].append((i, bbox))
                else:
                    # Non-PDF files: use stored text or re-read file
                    text = ''
                    if self.current_path and self.file_type in ('docx', 'txt'):
                        try:
                            if self.file_type == 'docx':
                                doc = Document(self.current_path)
                                text = '\n'.join([p.text for p in doc.paragraphs])
                            else:
                                with open(self.current_path, 'r', encoding='utf-8', errors='ignore') as f:
                                    text = f.read()
                        except Exception:
                            text = '\n'.join(self.ocr_results)
                    else:
                        text = '\n'.join(self.ocr_results)
                    self._ocr_queue.append(('page', 1, 1, text))
                    cleaned = fullwidth_to_halfwidth(text)
                    raw_codes = CODE_PATTERN.findall(cleaned)
                    for code in raw_codes:
                        normalized = code.upper().strip()
                        if normalized not in page_code_blocks:
                            page_code_blocks[normalized] = []
                        page_code_blocks[normalized].append((0, (0, 0, 0, 0)))

                self._ocr_queue.append(('status', '正在提取规范编号...'))
                all_text = '\n'.join(self.ocr_results)
                cleaned_text = fullwidth_to_halfwidth(all_text)
                raw_names = NAME_PATTERN.findall(cleaned_text)
                name_map = {}
                for raw_name in raw_names:
                    name_map[raw_name.strip()] = raw_name

                seen = set()
                self.extracted_codes = []
                self.extracted_code_info = {}
                for code in list(page_code_blocks.keys()):
                    if code in seen:
                        continue
                    seen.add(code)
                    self.extracted_codes.append(code)
                    matched_name = ''
                    for cname in name_map:
                        if code.replace(' ', '') in cname.replace(' ', '') or cname.replace(' ', '') in code.replace(' ', ''):
                            matched_name = name_map[cname]
                            break
                    self.extracted_code_info[code] = {
                        'name': matched_name,
                        'original': code,
                    }
                    first_page, first_bbox = page_code_blocks[code][0]
                    self.code_locations.append({
                        'code': code,
                        'page': first_page,
                        'bbox': first_bbox,
                    })

                self._ocr_queue.append(('codes', self.extracted_codes))

            except Exception as e:
                self._ocr_queue.append(('status', f'OCR 出错: {e}'))
                print(f"OCR fatal error: {e}")
            finally:
                self._ocr_done = True

        def process_queue():
            if not self._ocr_queue and not self._ocr_done:
                self.root.after(50, process_queue)
                return

            while self._ocr_queue:
                item = self._ocr_queue.pop(0)
                kind = item[0]
                if kind == 'page':
                    _, page_no, total, text = item
                    self.ocr_text.insert(tk.END, f"--- 第{page_no}页 ---\n{text}\n\n")
                    self.ocr_text.see(tk.END)
                    self.progress_var.set(page_no / total * 100)
                    self.status_var.set(f"OCR 识别中: {page_no}/{total}")
                elif kind == 'status':
                    _, msg = item
                    self.status_var.set(msg)
                elif kind == 'codes':
                    codes = item[1]
                    self.list_tree.delete(*self.list_tree.get_children())
                    if codes:
                        for i, code in enumerate(codes, 1):
                            info = self.extracted_code_info.get(normalize_for_matching(code), {})
                            name = info.get('name', '')
                            self.list_tree.insert('', tk.END, values=(i, code, name))
                        self.notebook.select(self.list_tree.master)
                        self.status_var.set(f"OCR 完成: 识别到 {len(codes)} 个规范编号")
                    else:
                        # Show preview text so user can see what OCR actually got
                        sample = '\n'.join(self.ocr_results[:3])
                        self.list_tree.insert('', tk.END, values=(1, '【未识别到规范编号】'))
                        self.list_tree.insert('', tk.END, values=(2, '请查看左侧 OCR 识别文本 确认内容'))
                        if sample.strip():
                            self.list_tree.insert('', tk.END, values=(3, sample[:120].replace('\n', ' ')))
                        self.notebook.select(self.list_tree.master)
                        self.status_var.set("OCR 完成，但未识别到规范编号，请查看 OCR 文本")

                    self.progress_var.set(100)

                    # Redraw current page to show markers
                    if self.file_type == 'pdf' and self.pdf_images:
                        self.show_page(self.current_display_index)

            if not self._ocr_done or self._ocr_queue:
                self.root.after(50, process_queue)
            else:
                self._ocr_queue = []
                self._ocr_done = False

        threading.Thread(target=do_ocr, daemon=True).start()
        self.root.after(50, process_queue)

    def check_standards(self):
        if not self.extracted_codes:
            messagebox.showwarning("提示", "请先进行 OCR 识别并提取规范编号")
            return
        self.status_var.set("检查规范中...")
        self.progress_var.set(0)
        self.check_tree.delete(*self.check_tree.get_children())
        self.check_results = []

        unique_codes = list(self.extracted_codes)
        total = len(unique_codes)
        for i, code in enumerate(unique_codes):
            info = self.extracted_code_info.get(normalize_for_matching(code), {})
            name = info.get('name', '')
            result = self.checker.check_code(code, name=name)
            self.check_results.append((code, result))

            status = result.get('status', '未找到')
            replacement = result.get('replacement_raw', '')
            matched = result.get('matched_name', result.get('matched_code', ''))
            display_code = code
            if matched:
                display_code = f"{code} -> {matched}"
            elif not result.get('found'):
                # Show similar codes for debugging
                similar = self.checker.find_similar_codes(code, limit=2)
                if similar:
                    similar_str = '; '.join([f"{s[1]}《{s[2]}《"[:60] for s in similar])
                    display_code = f"{code} [相似:{similar_str}]"
            if result.get('found'):
                if '废止' in status or '作废' in status:
                    action = '需替换'
                else:
                    action = '现行'
                if result.get('dual_match'):
                    action += ' (双重确认)'
            else:
                action = '未查询到'

            matched_name = result.get('matched_name', result.get('matched_code', '')) or name
            self.check_tree.insert('', tk.END, text=str(i+1),
                                   values=(display_code, matched_name, status, replacement, action))
            self.progress_var.set((i + 1) / total * 100)
            self.root.update_idletasks()

        self.progress_var.set(100)
        self.status_var.set(f"检查完成: {len(unique_codes)} 个规范")
        self.notebook.select(self.check_tree.master)

    def export_doc(self):
        if not self.check_results:
            messagebox.showwarning("提示", "没有检查结果可导出")
            return

        path = filedialog.asksaveasfilename(
            title="保存 DOC 报告",
            defaultextension=".docx",
            filetypes=[("Word documents", "*.docx"), ("All files", "*.*")]
        )
        if not path:
            return

        self.status_var.set("正在生成报告...")
        doc = Document()

        title = doc.add_heading('标准规范检查报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'文件: {os.path.basename(self.current_path) if self.current_path else "N/A"}')
        doc.add_paragraph()

        doc.add_heading('检查摘要', 1)
        total = len(self.check_results)
        found = sum(1 for _, r in self.check_results if r.get('found'))
        obsolete = sum(1 for _, r in self.check_results if '废止' in r.get('status', '') or '作废' in r.get('status', ''))
        doc.add_paragraph(f'共识别 {total} 个规范编号')
        doc.add_paragraph(f'数据库中查询到 {found} 个')
        doc.add_paragraph(f'其中废止/作废 {obsolete} 个')
        doc.add_paragraph()

        doc.add_heading('详细检查结果', 1)
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '序号'
        hdr_cells[1].text = '规范编号'
        hdr_cells[2].text = '规范名称'
        hdr_cells[3].text = '状态'
        hdr_cells[4].text = '替代情况'
        hdr_cells[5].text = '建议'

        for i, (code, result) in enumerate(self.check_results, 1):
            status = result.get('status', '未找到')
            replacement = result.get('replacement_raw', '')
            matched_name = result.get('matched_name', result.get('matched_code', ''))
            if result.get('found'):
                if '废止' in status or '作废' in status:
                    action = '需替换'
                else:
                    action = '现行'
            else:
                action = '未查询到'

            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = code
            row_cells[2].text = matched_name
            row_cells[3].text = status
            row_cells[4].text = replacement
            row_cells[5].text = action

        doc.save(path)
        self.progress_var.set(0)
        self.status_var.set(f"报告已保存: {path}")
        messagebox.showinfo("完成", f"报告已保存到:\n{path}")

    def on_check_item_selected(self, event=None):
        """When user selects a check result item, navigate to its page and highlight if possible."""
        selected = self.check_tree.selection()
        if not selected:
            return
        item = selected[0]
        values = self.check_tree.item(item, 'values')
        if not values:
            return
        display_code = values[0]
        name = values[1] if len(values) > 1 else ''
        
        # Update preview name label
        preview_name = f"{display_code.split(' ')[0].split('->')[0].split('[')[0].strip()} {name}".strip()
        if hasattr(self, '_preview_name_var'):
            self._preview_name_var.set(preview_name)
        
        # Extract original code from display (may contain " -> " or " [相似:")
        original_code = display_code.split(' ')[0].split('->')[0].split('[')[0].strip()
        
# Normalize code for matching (used in all modes)
        code_norm = normalize_for_matching(original_code)

        # Navigate to page only in PDF mode
        if self.file_type == 'pdf':
            for loc in self.code_locations:
                if normalize_for_matching(loc['code']) == code_norm:
                    self.show_page(loc['page'])
                    # Highlight the code location
                    self._highlight_code_location(loc)
                    break
        for item in self.list_tree.get_children():
            values = self.list_tree.item(item, 'values')
            if len(values) > 1 and normalize_for_matching(values[1]) == code_norm:
                self.list_tree.selection_set(item)
                self.list_tree.see(item)
                break
        
        # Highlight on preview for PDF
        if self.file_type == 'pdf':
            self._highlight_standard_on_preview(original_code, name)
    
    def _highlight_code_location(self, loc):
        """Highlight a specific code location on the preview."""
        if not hasattr(self, '_current_base_image') or not self._current_base_image:
            return
        page_idx = loc.get('page', 0)
        if page_idx != getattr(self, 'current_display_index', -1):
            self.show_page(page_idx)
            self.root.update_idletasks()
        
        # Draw a prominent red highlight box
        scale = getattr(self, '_zoom_level', 1.0)
        if hasattr(self, '_current_base_image') and self._current_base_image:
            canvas_w = self.pdf_canvas.winfo_width() or 400
            canvas_h = self.pdf_canvas.winfo_height() or 600
            img_w, img_h = self._current_base_image.size
            base_scale = min(canvas_w / img_w, canvas_h / img_h)
            scale = base_scale * scale
            offset_x = (canvas_w - int(img_w * scale)) // 2
            offset_y = (canvas_h - int(img_h * scale)) // 2
        
        x1, y1, x2, y2 = loc.get('bbox', (0, 0, 0, 0))
        if all(v == 0 for v in (x1, y1, x2, y2)):
            return
        
        x1, y1, x2, y2 = x1 * scale + offset_x, y1 * scale + offset_y, x2 * scale + offset_x, y2 * scale + offset_y
        
        # Remove old highlight
        if hasattr(self, '_highlight_rect_id') and self._highlight_rect_id:
            self.pdf_canvas.delete(self._highlight_rect_id)
        
        self._highlight_rect_id = self.pdf_canvas.create_rectangle(
            x1 - 2, y1 - 2, x2 + 2, y2 + 2,
            outline='red', width=3, dash=()
        )
        
        # Auto remove highlight after 3 seconds
        self.root.after(3000, self._clear_highlight)
    
    def _clear_highlight(self):
        if hasattr(self, '_highlight_rect_id') and self._highlight_rect_id:
            self.pdf_canvas.delete(self._highlight_rect_id)
            self._highlight_rect_id = None
    
    def _highlight_code_in_text(self, code):
        """Highlight the selected standard code in the OCR text widget."""
        if not hasattr(self, 'ocr_text'):
            return
        self.ocr_text.tag_remove('highlight', '1.0', tk.END)
        if not code:
            return
        start = '1.0'
        while True:
            pos = self.ocr_text.search(code, start, stopindex=tk.END, nocase=True)
            if not pos:
                break
            end = f"{pos}+{len(code)}c"
            self.ocr_text.tag_add('highlight', pos, end)
            start = end
        self.ocr_text.tag_config('highlight', background='yellow', foreground='red')
    
    def _highlight_standard_on_preview(self, code, name):
        """Highlight the standard code or name on the PDF preview canvas using fitz text search."""
        if self.file_type != 'pdf' or not getattr(self, 'current_path', None):
            return
        if not hasattr(self, '_current_base_image') or not self._current_base_image:
            return
        try:
            doc = fitz.open(self.current_path)
            page_idx = getattr(self, 'current_display_index', 0)
            if page_idx < 0 or page_idx >= len(doc):
                doc.close()
                return
            page = doc.load_page(page_idx)
            
            rect = None
            for search_text in (code, name):
                if not search_text:
                    continue
                blocks = page.search_for(search_text)
                if blocks:
                    rect = blocks[0]
                    break
            doc.close()
            
            if not rect:
                return
            
            canvas_w = self.pdf_canvas.winfo_width() or 400
            canvas_h = self.pdf_canvas.winfo_height() or 600
            img_w, img_h = self._current_base_image.size
            dpi = 200
            scale_factor = dpi / 72.0
            base_scale = min(canvas_w / img_w, canvas_h / img_h)
            zoom = getattr(self, '_zoom_level', 1.0)
            scale = base_scale * zoom
            offset_x = (canvas_w - int(img_w * scale)) // 2 + getattr(self, '_pan_image_x', 0)
            offset_y = (canvas_h - int(img_h * scale)) // 2 + getattr(self, '_pan_image_y', 0)
            
            x1 = rect.x0 * scale_factor * scale + offset_x
            y1 = rect.y0 * scale_factor * scale + offset_y
            x2 = rect.x1 * scale_factor * scale + offset_x
            y2 = rect.y1 * scale_factor * scale + offset_y
            
            if hasattr(self, '_highlight_rect_id') and self._highlight_rect_id:
                self.pdf_canvas.delete(self._highlight_rect_id)
            
            self._highlight_rect_id = self.pdf_canvas.create_rectangle(
                x1, y1, x2, y2,
                outline='red', width=3, dash=()
            )
        except Exception as e:
            print(f"highlight error: {e}")
    
    def on_check_item_double_click(self, event=None):
        """Double-click on check result item - if not found, show search dialog."""
        selected = self.check_tree.selection()
        if not selected:
            return
        item = selected[0]
        values = self.check_tree.item(item, 'values')
        if not values:
            return
        
        action = values[4] if len(values) > 4 else ''
        if action != '未查询到':
            return
        
        display_code = values[0]
        original_code = display_code.split(' ')[0].split('->')[0].split('[')[0].strip()
        
        # Get associated name if available
        name = ''
        if hasattr(self, 'extracted_code_info'):
            info = self.extracted_code_info.get(normalize_for_matching(original_code), {})
            name = info.get('name', '')
        
        dialog = StandardSearchDialog(self, self.checker, code=original_code, name=name)
        self.wait_window(dialog)

    def run(self):
        self._start_periodic_redraw()
        self.root.mainloop()


class StandardSearchDialog(tk.Toplevel):
    """规范搜索与推荐弹窗"""
    def __init__(self, parent, checker, code='', name=''):
        super().__init__(parent)
        self.checker = checker
        self.code = code
        self.name = name
        self.title("规范搜索与推荐")
        self.geometry("750x550")
        self.minsize(600, 450)
        self.transient(parent)
        self.grab_set()
        
        self._setup_ui()
        self._search_recommend()
        
    def _setup_ui(self):
        # Top search area
        search_frame = ttk.Frame(self, padding=10)
        search_frame.pack(side=tk.TOP, fill=tk.X)
        
        ttk.Label(search_frame, text="搜索规范:").pack(side=tk.LEFT, padx=(0, 5))
        self.search_var = tk.StringVar()
        search_entry = ttk.Entry(search_frame, textvariable=self.search_var, width=50)
        search_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 5))
        search_entry.bind('<Return>', lambda e: self._do_search())
        
        ttk.Button(search_frame, text="搜索", command=self._do_search).pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(search_frame, text="推荐相近", command=self._search_recommend).pack(side=tk.LEFT)
        
        # Info label
        if self.code:
            self.info_label = ttk.Label(self, text=f"当前规范: {self.code}" + (f"  《{self.name}》" if self.name else ""), foreground="#555555")
        else:
            self.info_label = ttk.Label(self, text="请输入编号或名称搜索规范", foreground="#555555")
        self.info_label.pack(side=tk.TOP, anchor=tk.W, padx=10, pady=(0, 5))
        
        # Results
        results_frame = ttk.Frame(self, padding=(10, 0, 10, 10))
        results_frame.pack(side=tk.TOP, fill=tk.BOTH, expand=True)
        
        columns = ('code', 'name', 'status', 'action')
        self.result_tree = ttk.Treeview(results_frame, columns=columns, show='headings', height=15)
        self.result_tree.heading('code', text='规范编号')
        self.result_tree.heading('name', text='规范名称')
        self.result_tree.heading('status', text='状态')
        self.result_tree.heading('action', text='操作')
        self.result_tree.column('code', width=160)
        self.result_tree.column('name', width=350)
        self.result_tree.column('status', width=100)
        self.result_tree.column('action', width=120)
        self.result_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        scroll = ttk.Scrollbar(results_frame, orient=tk.VERTICAL, command=self.result_tree.yview)
        self.result_tree.configure(yscrollcommand=scroll.set)
        scroll.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Bottom buttons
        btn_frame = ttk.Frame(self, padding=(10, 0, 10, 10))
        btn_frame.pack(side=tk.BOTTOM, fill=tk.X)
        ttk.Button(btn_frame, text="关闭", command=self.destroy).pack(side=tk.RIGHT)
        
    def _do_search(self):
        query = self.search_var.get().strip()
        if not query:
            messagebox.showwarning("提示", "请输入搜索内容")
            return
        self._perform_search(query)
        
    def _search_recommend(self):
        if self.code:
            self._perform_search(self.code)
        
    def _perform_search(self, query):
        self.result_tree.delete(*self.result_tree.get_children())
        norm_query = normalize_for_matching(query)
        
        results = []
        for k, v in self.checker.code_index.items():
            code = v.get('code', '')
            name = v.get('name', '')
            norm_code = normalize_for_matching(code)
            norm_name = normalize_for_matching(name)
            
            score = 0
            if norm_query in norm_code or norm_code in norm_query:
                score = max(score, len(norm_query) / max(len(norm_code), 1))
            if norm_query in norm_name or norm_name in norm_query:
                score = max(score, len(norm_query) / max(len(norm_name), 1))
            if score > 0:
                results.append((score, v))
        
        results.sort(key=lambda x: x[0], reverse=True)
        results = results[:50]
        
        for _, v in results:
            code = v.get('code', '')
            name = v.get('name', '')
            status = v.get('status', '')
            is_obsolete = '废止' in status or '作废' in status
            
            item_id = self.result_tree.insert('', tk.END, values=(code, name, status, ''))
            if is_obsolete:
                self.result_tree.item(item_id, tags=('obsolete',))
            else:
                self.result_tree.item(item_id, tags=('active',))
        
        self.result_tree.tag_configure('obsolete', foreground='red')
        self.result_tree.tag_configure('active', foreground='green')
        
        self.result_tree.bind('<Double-Button-1>', self._on_result_double_click)
        
    def _on_result_double_click(self, event):
        item = self.result_tree.selection()
        if not item:
            return
        values = self.result_tree.item(item[0], 'values')
        code, name, status = values[0], values[1], values[2]
        
        # Create copy dialog
        copy_win = tk.Toplevel(self)
        copy_win.title("复制规范信息")
        copy_win.geometry("500x200")
        copy_win.transient(self)
        copy_win.grab_set()
        
        ttk.Label(copy_win, text=f"规范编号: {code}", font=("SimSun", 11)).pack(anchor=tk.W, padx=10, pady=(10, 5))
        ttk.Label(copy_win, text=f"规范名称: {name}", font=("SimSun", 11)).pack(anchor=tk.W, padx=10, pady=5)
        ttk.Label(copy_win, text=f"状态: {status}", foreground="red" if '废止' in status or '作废' in status else "green").pack(anchor=tk.W, padx=10, pady=5)
        
        btn_frame = ttk.Frame(copy_win, padding=10)
        btn_frame.pack(side=tk.BOTTOM, fill=tk.X)
        
        def copy_code():
            self.clipboard_clear()
            self.clipboard_append(code)
            messagebox.showinfo("完成", "规范编号已复制", parent=copy_win)
        
        def copy_name():
            self.clipboard_clear()
            self.clipboard_append(name)
            messagebox.showinfo("完成", "规范名称已复制", parent=copy_win)
        
        ttk.Button(btn_frame, text=f"复制编号: {code[:30]}...", command=copy_code).pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(btn_frame, text=f"复制名称: {name[:30]}...", command=copy_name).pack(side=tk.LEFT)
        ttk.Button(btn_frame, text="关闭", command=copy_win.destroy).pack(side=tk.RIGHT)


def main():
    print("Starting 工程助手 LDAssistant...")
    app = App()
    app.run()
    print("Application exited.")


if __name__ == "__main__":
    main()
